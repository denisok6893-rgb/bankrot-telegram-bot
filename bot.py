import asyncio
import json
import logging
import os
import sqlite3
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple
from dotenv import load_dotenv
load_dotenv()

from bankrot_bot.logging_setup import setup_logging
#from bankrot_bot.services.gigachat import gigachat_chat

logger = logging.getLogger(__name__)


from bankrot_bot.services.blocks import (
    build_creditors_header_block,
    build_creditors_block,
    sum_creditors_total,
    build_vehicle_block,
    build_attachments_list,
)
from bankrot_bot.services.public_docs import (
    get_categories,
    get_docs_in_category,
    get_document,
    CATEGORY_TITLES,
)
from bankrot_bot.services.case_financials import (
    get_case_parties,
    add_case_party,
    delete_case_party,
    calculate_parties_totals,
    format_parties_for_doc,
    get_case_assets,
    add_case_asset,
    delete_case_asset,
    calculate_assets_total,
    parse_amount_input,
    normalize_amount_to_string,
    string_to_decimal,
)
from bankrot_bot.services.docx_forms import (
    render_creditors_list,
    render_inventory,
)

import aiohttp
setup_logging()
from aiogram import Bot, Dispatcher, F
from aiogram.filters import Command, CommandStart, StateFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import CallbackQuery, FSInputFile, BufferedInputFile, Message, InlineKeyboardMarkup, ReplyKeyboardMarkup
from aiogram.utils.keyboard import InlineKeyboardBuilder
from docx import Document
from bankrot_bot.config import load_settings

# Database and handlers
from bankrot_bot.database import init_db as init_pg_db
from bankrot_bot.handlers import cases as cases_handlers
#from handlers import probability as probability_handlers
from handlers import newcase_fsm

from bankrot_bot.keyboards.menus import (
    main_menu_kb,
    start_ikb,
    home_ikb,
    profile_ikb,
    cases_list_ikb,
    case_card_ikb,
    docs_home_ikb,
    help_ikb,
    help_item_ikb,
    docs_menu_ikb,
    case_files_ikb,
    case_archive_ikb,
    cases_menu_ikb,
    my_cases_ikb,
    docs_catalog_ikb,
    docs_category_ikb,
    docs_item_ikb,
    case_parties_ikb,
    party_view_ikb,
    case_assets_ikb,
    asset_view_ikb,
)

class CaseCreate(StatesGroup):
    code_name = State()
    case_number = State()
    court = State()
    judge = State()
    fin_manager = State()
class ProfileFill(StatesGroup):
    full_name = State()
    role = State()
    address = State()
    phone = State()
    email = State()
class CaseEdit(StatesGroup):
    value = State()
class CaseCardFill(StatesGroup):
    waiting_value = State()
class CreditorsFill(StatesGroup):
    name = State()
    inn = State()
    ogrn = State()
    address = State()
    debt_rubles = State()
    debt_kopeks = State()
    note = State()
    creditors_text = State()  # For bulk creditor text input

class AddParty(StatesGroup):
    """FSM –¥–ª—è –¥–æ–±–∞–≤–ª–µ–Ω–∏—è –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞/–¥–æ–ª–∂–Ω–∏–∫–∞."""
    name = State()
    amount = State()
    basis = State()

class AddAsset(StatesGroup):
    """FSM –¥–ª—è –¥–æ–±–∞–≤–ª–µ–Ω–∏—è –∏–º—É—â–µ—Å—Ç–≤–∞."""
    kind = State()
    description = State()
    value = State()

# =========================
# env
# =========================


def build_gender_forms(gender: str | None) -> dict:
    """
    –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Å–ª–æ–≤–∞ –≤ –Ω—É–∂–Ω–æ–º —Ä–æ–¥–µ –¥–ª—è –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–æ–≤ —à–∞–±–ª–æ–Ω–∞:
    {{debtor_having_word}}, {{debtor_registered_word}}, {{debtor_living_word}},
    {{debtor_not_registered_word}}, {{debtor_insolvent_word}}
    """
    g = (gender or "").strip().lower()
    if g == "female":
        return {
            "debtor_having_word": "–∏–º–µ—é—â–∞—è",
            "debtor_registered_word": "–∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–Ω–∞—è",
            "debtor_living_word": "–ø—Ä–æ–∂–∏–≤–∞—é—â–∞—è",
            "debtor_not_registered_word": "–Ω–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∞",
            "debtor_insolvent_word": "–Ω–µ—Å–æ—Å—Ç–æ—è—Ç–µ–ª—å–Ω–æ–π",
        }
    # –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é male
    return {
        "debtor_having_word": "–∏–º–µ—é—â–∏–π",
        "debtor_registered_word": "–∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–Ω—ã–π",
        "debtor_living_word": "–ø—Ä–æ–∂–∏–≤–∞—é—â–∏–π",
        "debtor_not_registered_word": "–Ω–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω",
        "debtor_insolvent_word": "–Ω–µ—Å–æ—Å—Ç–æ—è—Ç–µ–ª—å–Ω—ã–º",
    }


def build_debtor_last_name_initials(card: dict) -> str:
    """
    –ò–∑ '–ò–≤–∞–Ω–æ–≤ –ò–≤–∞–Ω –ò–≤–∞–Ω–æ–≤–∏—á' –¥–µ–ª–∞–µ—Ç '–ò–≤–∞–Ω–æ–≤ –ò. –ò.'
    –ï—Å–ª–∏ –§–ò–û –ø—É—Å—Ç–æ–µ/–Ω–µ–ø–æ–ª–Ω–æ–µ ‚Äî –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç –∫–∞–∫ –µ—Å—Ç—å.
    """
    full_name = (card.get("debtor_full_name") or "").strip()
    parts = [p for p in full_name.split() if p]
    if len(parts) >= 2:
        last = parts[0]
        first_i = parts[1][0].upper() + "."
        patro_i = (parts[2][0].upper() + ".") if len(parts) >= 3 and parts[2] else ""
        return (last + " " + first_i + (" " + patro_i if patro_i else "")).strip()
    return full_name


def build_family_status_block(card: dict) -> str:
    """
    –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Ç–µ–∫—Å—Ç–æ–≤—ã–π –±–ª–æ–∫ –æ —Å–µ–º–µ–π–Ω–æ–º –ø–æ–ª–æ–∂–µ–Ω–∏–∏/–¥–µ—Ç—è—Ö –¥–ª—è {{family_status_block}}.
    –ü–æ–ª—è –æ–∂–∏–¥–∞—é—Ç—Å—è: marital_status, spouse_full_name, has_minor_children, children_count,
    marriage_certificate_number, marriage_certificate_date
    """
    marital_status = (card.get("marital_status") or "").strip()
    spouse_full_name = (card.get("spouse_full_name") or "").strip()
    has_minor_children = card.get("has_minor_children")
    children_count = card.get("children_count")
    cert_no = (card.get("marriage_certificate_number") or "").strip()
    cert_date = (card.get("marriage_certificate_date") or "").strip()

    lines: list[str] = []

    if marital_status == "married":
        line = "–°–æ—Å—Ç–æ–∏—Ç –≤ –±—Ä–∞–∫–µ"
        if spouse_full_name:
            line += f" —Å {spouse_full_name}"
        line += "."
        lines.append(line)

        if cert_no:
            cert_line = f"–°–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–æ –æ –∑–∞–∫–ª—é—á–µ–Ω–∏–∏ –±—Ä–∞–∫–∞ ‚Ññ {cert_no}"
            if cert_date:
                cert_line += f" –æ—Ç {cert_date}"
            cert_line += "."
            lines.append(cert_line)

    elif marital_status == "single":
        lines.append("–í –±—Ä–∞–∫–µ –Ω–µ —Å–æ—Å—Ç–æ–∏—Ç.")

    if has_minor_children is True:
        cnt = ""
        if children_count not in (None, ""):
            cnt = f" ({children_count} —Ä–µ–±—ë–Ω–æ–∫(–¥–µ—Ç–µ–π))"
        lines.append(f"–ò–º–µ–µ—Ç –Ω–µ—Å–æ–≤–µ—Ä—à–µ–Ω–Ω–æ–ª–µ—Ç–Ω–∏—Ö –¥–µ—Ç–µ–π{cnt}.")
    elif has_minor_children is False:
        lines.append("–ù–µ—Å–æ–≤–µ—Ä—à–µ–Ω–Ω–æ–ª–µ—Ç–Ω–∏—Ö –¥–µ—Ç–µ–π –Ω–µ—Ç.")

    return "\n".join(lines)


def _old_build_creditors_block(creditors: list[dict] | None) -> str:
    """
    –ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ—Ç 2 —Ñ–æ—Ä–º–∞—Ç–∞ –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤:

    1) –ù–æ–≤—ã–π (–æ–ø—Ä–æ—Å–Ω–∏–∫):
       {
         "name": "...",
         "inn": "...", "ogrn": "...",
         "debt_rubles": "...", "debt_kopeks": "...",
         "note": "–û–ö–ë/–¥–æ–≥–æ–≤–æ—Ä/–∏ —Ç.–ø."
       }

    2) –°—Ç–∞—Ä—ã–π:
       {
         "name": "...",
         "obligations": [{"amount_rubles":123,"amount_kopeks":45,"source":"–û–ö–ë"}]
       }
    """
    if not isinstance(creditors, list) or not creditors:
        return ""

    def _digits(s: str) -> str:
        return "".join(ch for ch in str(s) if ch.isdigit())

    lines: list[str] = []

    for i, c in enumerate(creditors, start=1):
        if not isinstance(c, dict):
            continue

        name = str((c.get("name") or "–ö—Ä–µ–¥–∏—Ç–æ—Ä")).strip()

        # --- –∏–¥–µ–Ω—Ç–∏—Ñ–∏–∫–∞—Ç–æ—Ä—ã (–¥–ª—è –Ω–æ–≤–æ–≥–æ —Ñ–æ—Ä–º–∞—Ç–∞) ---
        inn = str(c.get("inn") or "").strip()
        ogrn = str(c.get("ogrn") or "").strip()
        ids = []
        if inn:
            ids.append(f"–ò–ù–ù {inn}")
        if ogrn:
            ids.append(f"–û–ì–†–ù {ogrn}")
        name_with_ids = name + (f" ({', '.join(ids)})" if ids else "")

        # --- –Ω–æ–≤—ã–π —Ñ–æ—Ä–º–∞—Ç —Å—É–º–º—ã ---
        debt_r = c.get("debt_rubles")
        debt_k = c.get("debt_kopeks")
        note = str(c.get("note") or "").strip()

        money_new = ""
        if debt_r not in (None, "", "-"):
            dr = _digits(debt_r)
            if dr != "":
                money_new = f"{int(dr)} —Ä—É–±."
        if debt_k not in (None, "", "-"):
            dk = _digits(debt_k)
            if dk != "":
                money_new = (money_new + " " if money_new else "") + f"{int(dk):02d} –∫–æ–ø."

        if money_new and note:
            line_new = f"{i}) {name_with_ids} ‚Äî {money_new} ({note})"
        elif money_new:
            line_new = f"{i}) {name_with_ids} ‚Äî {money_new}"
        elif note:
            line_new = f"{i}) {name_with_ids} ‚Äî ({note})"
        else:
            line_new = f"{i}) {name_with_ids}"

        # --- —Å—Ç–∞—Ä—ã–π —Ñ–æ—Ä–º–∞—Ç obligations –∏–º–µ–µ—Ç –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç, –µ—Å–ª–∏ –æ–Ω —Ä–µ–∞–ª—å–Ω–æ –∑–∞–ø–æ–ª–Ω–µ–Ω ---
        obs = c.get("obligations")
        if isinstance(obs, list) and any(isinstance(x, dict) for x in obs):
            obs_txt: list[str] = []
            for ob in obs:
                if not isinstance(ob, dict):
                    continue
                r = ob.get("amount_rubles")
                k = ob.get("amount_kopeks")
                src = (ob.get("source") or "").strip()

                money_parts: list[str] = []
                if r is not None and str(r).strip() != "":
                    money_parts.append(f"{int(r)} —Ä—É–±.")
                if k is not None and str(k).strip() != "":
                    money_parts.append(f"{int(k):02d} –∫–æ–ø.")
                money = " ".join(money_parts).strip()

                if money and src:
                    obs_txt.append(f"{money} ({src})")
                elif money:
                    obs_txt.append(money)
                elif src:
                    obs_txt.append(f"({src})")

            if obs_txt:
                lines.append(f"{i}) {name} ‚Äî " + "; ".join(obs_txt))
            else:
                lines.append(f"{i}) {name}")
        else:
            lines.append(line_new)

    return "\n".join(lines)

def _old_sum_creditors_total(creditors: list[dict] | None) -> tuple[int, int]:
    """
    –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç (rubles, kopeks) –∫–∞–∫ —Å—É–º–º—É –ø–æ –≤—Å–µ–º –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞–º.
    –ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ—Ç –æ–±–∞ —Ñ–æ—Ä–º–∞—Ç–∞:
      - obligations: [{amount_rubles, amount_kopeks, ...}]
      - debt_rubles / debt_kopeks
    """
    if not isinstance(creditors, list) or not creditors:
        return (0, 0)

    def _to_int(x) -> int:
        if x is None:
            return 0
        s = "".join(ch for ch in str(x) if ch.isdigit())
        return int(s) if s else 0

    total_k = 0

    for c in creditors:
        if not isinstance(c, dict):
            continue

        obs = c.get("obligations")
        if isinstance(obs, list) and any(isinstance(o, dict) for o in obs):
            for ob in obs:
                if not isinstance(ob, dict):
                    continue
                r = _to_int(ob.get("amount_rubles"))
                k = _to_int(ob.get("amount_kopeks"))
                total_k += r * 100 + k
            continue

        # –Ω–æ–≤—ã–π —Ñ–æ—Ä–º–∞—Ç
        r = _to_int(c.get("debt_rubles"))
        k = _to_int(c.get("debt_kopeks"))
        total_k += r * 100 + k

    return (total_k // 100, total_k % 100)

def _old_build_vehicle_block(card: dict) -> str:
    """
    –ï—Å–ª–∏ –∞–≤—Ç–æ –Ω–µ—Ç ‚Äî '–û—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç'.
    –ï—Å–ª–∏ –µ—Å—Ç—å —Å–ø–∏—Å–æ–∫ vehicles –∏–ª–∏ vehicle ‚Äî –ø–µ—á–∞—Ç–∞–µ–º —Å–ø–∏—Å–∫–æ–º.
    """
    vehicles: list[dict] = []

    vlist = card.get("vehicles")
    if isinstance(vlist, list):
        vehicles.extend([v for v in vlist if isinstance(v, dict)])

    one = card.get("vehicle")
    if isinstance(one, dict):
        vehicles.append(one)

    if not vehicles:
        return "–¢—Ä–∞–Ω—Å–ø–æ—Ä—Ç–Ω—ã–µ —Å—Ä–µ–¥—Å—Ç–≤–∞ –æ—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç."

    lines: list[str] = []
    for i, v in enumerate(vehicles, start=1):
        brand_model = (v.get("brand_model") or "").strip()
        plate = (v.get("plate_number") or "").strip()
        vin = (v.get("vin") or "").strip()
        year = (v.get("year") or "").strip()
        parts = [p for p in [brand_model, plate, vin, year] if p]
        desc = "; ".join(parts) if parts else "–ê–≤—Ç–æ–º–æ–±–∏–ª—å"
        lines.append(f"{i}) {desc}")

    return "\n".join(lines)


def _old_build_attachments_list(card: dict) -> str:
    items: list[str] = []
    if card.get("passport_series") and card.get("passport_number"):
        items.append("–ö–æ–ø–∏—è –ø–∞—Å–ø–æ—Ä—Ç–∞ –≥—Ä–∞–∂–¥–∞–Ω–∏–Ω–∞ –†–§.")
    if card.get("debtor_inn"):
        items.append("–ö–æ–ø–∏—è –ò–ù–ù.")
    if card.get("debtor_snils"):
        items.append("–ö–æ–ø–∏—è –°–ù–ò–õ–°.")
    if card.get("creditors"):
        items.append("–î–æ–∫—É–º–µ–Ω—Ç—ã, –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–∞—é—â–∏–µ –∑–∞–¥–æ–ª–∂–µ–Ω–Ω–æ—Å—Ç—å –ø–µ—Ä–µ–¥ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞–º–∏.")

    if not items:
        return ""
    return "\n".join(f"{i}) {x}" for i, x in enumerate(items, start=1))

def _set_paragraph_text_keep_style(paragraph, new_text: str) -> None:
    """
    –ù–∞–¥—ë–∂–Ω–∞—è –∑–∞–º–µ–Ω–∞ —Ç–µ–∫—Å—Ç–∞ –≤ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–µ: –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä—ã –º–æ–≥—É—Ç –±—ã—Ç—å —Ä–∞–∑–æ—Ä–≤–∞–Ω—ã –ø–æ runs.
    –°–æ—Ö—Ä–∞–Ω—è–µ–º —Å—Ç–∏–ª—å –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞, –Ω–æ runs –ø–µ—Ä–µ—Å–æ–∑–¥–∞—ë–º.
    """
    if paragraph.runs:
        for r in paragraph.runs:
            r.text = ""
    paragraph.add_run(new_text)


def _replace_placeholders_strong(doc: Document, mapping: Dict[str, Any]) -> None:
    """
    –ó–∞–º–µ–Ω–∞ –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–æ–≤ —Ñ–æ—Ä–º–∞—Ç–∞ {{key}} –ø–æ –ø–æ–ª–Ω–æ–º—É —Ç–µ–∫—Å—Ç—É –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤ –∏ —è—á–µ–µ–∫ —Ç–∞–±–ª–∏—Ü.
    mapping: –∫–ª—é—á–∏ –ë–ï–ó —Ñ–∏–≥—É—Ä–Ω—ã—Ö —Å–∫–æ–±–æ–∫, –Ω–∞–ø—Ä–∏–º–µ—Ä: {"court_name": "..." }
    """
    def apply_to_paragraph(p):
        text = p.text
        if not text or "{{" not in text:
            return
        changed = False
        for k, v in mapping.items():
            placeholder = f"{{{{{k}}}}}"
            if placeholder in text:
                text = text.replace(placeholder, "" if v is None else str(v))
                changed = True
        if changed:
            _set_paragraph_text_keep_style(p, text)

    for p in doc.paragraphs:
        apply_to_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_to_paragraph(p)
                for nested in cell.tables:
                    # —Ä–µ–∫—É—Ä—Å–∏–≤–Ω–æ –¥–ª—è –≤–ª–æ–∂–µ–Ω–Ω—ã—Ö —Ç–∞–±–ª–∏—Ü
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for np in ncell.paragraphs:
                                apply_to_paragraph(np)

def _old_build_online_hearing_docx(case_row: Tuple) -> Path:
    """
    –ì–µ–Ω–µ—Ä–∞—Ü–∏—è —Ö–æ–¥–∞—Ç–∞–π—Å—Ç–≤–∞ –æ –í–ö–° (–æ–Ω–ª–∞–π–Ω-–∑–∞—Å–µ–¥–∞–Ω–∏–µ).
    –î–µ–ª–∞–µ—Ç –ø—Ä–æ—Å—Ç–æ–π DOCX –±–µ–∑ —à–∞–±–ª–æ–Ω–∞, —á—Ç–æ–±—ã –≥–∞—Ä–∞–Ω—Ç–∏—Ä–æ–≤–∞–Ω–Ω–æ –Ω–µ –ø–∞–¥–∞—Ç—å.
    """
    (
        cid,
        owner_user_id,
        code_name,
        case_number,
        court,
        judge,
        fin_manager,
        stage,
        notes,
        created_at,
        updated_at,
    ) = case_row

    doc = Document()

    # –®–∞–ø–∫–∞: –∫–æ–º—É –∏ –æ—Ç –∫–æ–≥–æ
    doc.add_paragraph("–í –ê—Ä–±–∏—Ç—Ä–∞–∂–Ω—ã–π —Å—É–¥")
    doc.add_paragraph(court or "–Ω–µ —É–∫–∞–∑–∞–Ω–æ")
    doc.add_paragraph("")

    prof = get_profile(owner_user_id)
    if prof:
        _, full_name, role, address, phone, email, *_ = prof
        doc.add_paragraph("–û—Ç: " + (full_name or "–Ω–µ —É–∫–∞–∑–∞–Ω–æ"))
        if role:
            doc.add_paragraph("–°—Ç–∞—Ç—É—Å: " + role)
        if address:
            doc.add_paragraph("–ê–¥—Ä–µ—Å: " + address)
        if phone:
            doc.add_paragraph("–¢–µ–ª–µ—Ñ–æ–Ω: " + phone)
        if email:
            doc.add_paragraph("Email: " + email)
    else:
        doc.add_paragraph("–û—Ç: –Ω–µ —É–∫–∞–∑–∞–Ω–æ")

    doc.add_paragraph("")
    doc.add_paragraph("–•–û–î–ê–¢–ê–ô–°–¢–í–û")
    doc.add_paragraph("–æ –ø—Ä–æ–≤–µ–¥–µ–Ω–∏–∏ —Å—É–¥–µ–±–Ω–æ–≥–æ –∑–∞—Å–µ–¥–∞–Ω–∏—è —Å –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ–º –í–ö–°")
    doc.add_paragraph("")

    if case_number:
        doc.add_paragraph(f"–î–µ–ª–æ ‚Ññ {case_number}")
    if judge:
        doc.add_paragraph(f"–°—É–¥—å—è: {judge}")
    if fin_manager:
        doc.add_paragraph(f"–§–∏–Ω–∞–Ω—Å–æ–≤—ã–π —É–ø—Ä–∞–≤–ª—è—é—â–∏–π: {fin_manager}")

    doc.add_paragraph("")
    doc.add_paragraph("–ü—Ä–æ—à—É –æ–±–µ—Å–ø–µ—á–∏—Ç—å —É—á–∞—Å—Ç–∏–µ –≤ —Å—É–¥–µ–±–Ω–æ–º –∑–∞—Å–µ–¥–∞–Ω–∏–∏ —Å –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ–º —Å–∏—Å—Ç–µ–º—ã –≤–∏–¥–µ–æ–∫–æ–Ω—Ñ–µ—Ä–µ–Ω—Ü-—Å–≤—è–∑–∏.")
    doc.add_paragraph("")
    doc.add_paragraph("–î–∞—Ç–∞: " + datetime.now().strftime("%d.%m.%Y"))
    doc.add_paragraph("")
    doc.add_paragraph("–ü–æ–¥–ø–∏—Å—å: ____________")

    fname = f"online_hearing_case_{cid}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    case_dir = GENERATED_DIR / "cases" / str(cid)
    case_dir.mkdir(parents=True, exist_ok=True)
    out_path = case_dir / fname
    doc.save(out_path)
    return out_path


async def build_bankruptcy_petition_doc(case_row: Tuple, card: dict) -> Path:
    """
    –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –∑–∞—è–≤–ª–µ–Ω–∏—è –æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–µ –ø–æ —à–∞–±–ª–æ–Ω—É.
    –ü–æ–¥—Å—Ç–∞–Ω–æ–≤–∫–∞ —Å—Ç—Ä–æ–≥–æ –ø–æ 23 –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–∞–º —à–∞–±–ª–æ–Ω–∞ + –¥–µ—Ñ–æ–ª—Ç—ã –¥–ª—è –ø—É—Å—Ç—ã—Ö –¥–∞–Ω–Ω—ã—Ö.

    –ù–û–í–û–ï: –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç–Ω–æ –∏—Å–ø–æ–ª—å–∑—É–µ–º –¥–∞–Ω–Ω—ã–µ –∏–∑ case_parties (–µ—Å–ª–∏ –µ—Å—Ç—å).
    """
    cid = case_row[0]

    template_path = Path("templates/petitions/bankruptcy_petition.docx")
    doc = Document(template_path)

    # –ü–æ–ø—ã—Ç–∫–∞ –∑–∞–≥—Ä—É–∑–∏—Ç—å –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤ –∏–∑ –Ω–æ–≤—ã—Ö —Ç–∞–±–ª–∏—Ü
    creditors_from_db = []
    try:
        from bankrot_bot.database import get_session

        async with get_session() as session:
            parties = await get_case_parties(session, cid, role="creditor")
            if parties:
                creditors_from_db = format_parties_for_doc(parties, role="creditor")
    except Exception as e:
        logger.warning(f"Failed to load creditors from DB for case {cid}: {e}")

    # --- –¥–µ—Ñ–æ–ª—Ç—ã ---
    def _txt(v: Any) -> str:
        v = "" if v is None else str(v).strip()
        return v if v else "–Ω–µ —É–∫–∞–∑–∞–Ω–æ"

    def _money_rubles(v: Any) -> str:
        v = "" if v is None else str(v).strip()
        return v if v else "0"

    def _money_kopeks(v: Any) -> str:
        if v is None or str(v).strip() == "":
            return "00"
        try:
            return f"{int(str(v).strip()):02d}"
        except (ValueError, TypeError):
            s = str(v).strip()
            digits = "".join(ch for ch in s if ch.isdigit())
            if digits == "":
                return "00"
            try:
                return f"{int(digits):02d}"
            except (ValueError, TypeError):
                return "00"

    # --- –∏—Å—Ö–æ–¥–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ ---
    court_name = _txt(card.get("court_name") or (case_row[4] if len(case_row) > 4 else None))
    court_address = _txt(card.get("court_address"))

    financial_manager_info = _txt(
        card.get("financial_manager_info") or (case_row[6] if len(case_row) > 6 else None)
    )

    debtor_full_name = _txt(card.get("debtor_full_name"))
    debtor_address = _txt(card.get("debtor_address"))
    debtor_birth_date = _txt(card.get("debtor_birth_date"))
    debtor_inn = _txt(card.get("debtor_inn"))
    debtor_snils = _txt(card.get("debtor_snils"))
    debtor_phone = _txt(card.get("debtor_phone"))

    passport_series = (card.get("passport_series") or "").strip()
    passport_number = (card.get("passport_number") or "").strip()
    debtor_passport = _txt(f"{passport_series} {passport_number}".strip())

    debtor_passport_issued_by = _txt(card.get("passport_issued_by"))
    debtor_passport_date = _txt(card.get("passport_date"))
    debtor_passport_code = _txt(card.get("passport_code"))

    raw_marital = card.get("marital_status")
    raw_marital = ("" if raw_marital is None else str(raw_marital)).strip().lower()

    debtor_address = (debtor_address or "").strip()
    while ",," in debtor_address:
        debtor_address = debtor_address.replace(",,", ",")
    debtor_address = debtor_address.rstrip(" ,")

    marital_map = {
        "married": "–°–æ—Å—Ç–æ–∏—Ç –≤ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–Ω–æ–º –±—Ä–∞–∫–µ.",
        "single": "–í –±—Ä–∞–∫–µ –Ω–µ —Å–æ—Å—Ç–æ–∏—Ç.",
        "divorced": "–ë—Ä–∞–∫ —Ä–∞—Å—Ç–æ—Ä–≥–Ω—É—Ç.",
        "widowed": "–í–¥–æ–≤–µ—Ü/–≤–¥–æ–≤–∞.",
    }

    # –ï—Å–ª–∏ —É–∂–µ –≤–≤–µ–¥—ë–Ω –Ω–æ—Ä–º–∞–ª—å–Ω—ã–π —Ä—É—Å—Å–∫–∏–π —Ç–µ–∫—Å—Ç ‚Äî –æ—Å—Ç–∞–≤–ª—è–µ–º –∫–∞–∫ –µ—Å—Ç—å
    if raw_marital in marital_map:
        marital_status = marital_map[raw_marital]
    else:
        # –µ—Å–ª–∏ —Å—Ç—Ä–æ–∫–∞ –ø—É—Å—Ç–∞—è -> –¥–µ—Ñ–æ–ª—Ç "–Ω–µ —É–∫–∞–∑–∞–Ω–æ"
        # –µ—Å–ª–∏ —Å—Ç—Ä–æ–∫–∞ –Ω–µ –ø—É—Å—Ç–∞—è (–≤ —Ç.—á. —Ä—É—Å—Å–∫–∏–π —Ç–µ–∫—Å—Ç) -> –∏—Å–ø–æ–ª—å–∑—É–µ–º –∫–∞–∫ –µ—Å—Ç—å
        marital_status = _txt(raw_marital)

    certificate_number = card.get("certificate_number") or card.get("marriage_certificate_number")
    certificate_date = card.get("certificate_date") or card.get("marriage_certificate_date")
    certificate_number = _txt(certificate_number)
    certificate_date = _txt(certificate_date)

    # –ù–û–í–ê–Ø –õ–û–ì–ò–ö–ê: –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç - creditors_from_db, –ø–æ—Ç–æ–º card creditors
    if creditors_from_db:
        creditors = creditors_from_db
    else:
        creditors = card.get("creditors") if isinstance(card.get("creditors"), list) else []

    auto_r, auto_k = sum_creditors_total(creditors)
    if auto_r or auto_k:
        total_debt_rubles = str(auto_r)
        total_debt_kopeks = f"{auto_k:02d}"
    else:
        total_debt_rubles = _money_rubles(card.get("total_debt_rubles"))
        total_debt_kopeks = _money_kopeks(card.get("total_debt_kopeks"))

    deposit_deferral_request = card.get("deposit_deferral_request") or ""
    # attachments_list –ø–æ —É—Ç–≤–µ—Ä–∂–¥—ë–Ω–Ω—ã–º –¥–µ—Ñ–æ–ª—Ç–∞–º –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –ø—É—Å—Ç—ã–º, –µ—Å–ª–∏ –Ω–µ—Ç –¥–∞–Ω–Ω—ã—Ö
    attachments_list = ""
    try:
        built_attachments = build_attachments_list(card)
        if built_attachments and str(built_attachments).strip():
            attachments_list = str(built_attachments)
    except (KeyError, TypeError, AttributeError) as e:
        logger.warning(f"Failed to build attachments list: {e}")
        attachments_list = ""

    # creditors_block: creditors_text –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç–Ω–æ, –∏–Ω–∞—á–µ —Å–ø–∏—Å–æ–∫, –∏–Ω–∞—á–µ –Ω–µ–π—Ç—Ä–∞–ª—å–Ω—ã–π —Ç–µ–∫—Å—Ç
    creditors_text = card.get("creditors_text")
    creditors_text = str(creditors_text).strip() if creditors_text is not None else ""
    creditors = card.get("creditors") if isinstance(card.get("creditors"), list) else []

    if creditors_text:
        creditors_block = creditors_text
    elif creditors:
        creditors_block = build_creditors_block(creditors)
    else:
        creditors_block = "–°–≤–µ–¥–µ–Ω–∏—è –æ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞—Ö –Ω–µ –ø—Ä–µ–¥—Å—Ç–∞–≤–ª–µ–Ω—ã."

    # creditors_header_block: –∫–æ—Ä–æ—Ç–∫–∏–π —Å–ø–∏—Å–æ–∫ –¥–ª—è —à–∞–ø–∫–∏ (–∏–∑ —Ç–æ–≥–æ –∂–µ –∏—Å—Ç–æ—á–Ω–∏–∫–∞, —á—Ç–æ –∏ creditors_block)
    if creditors:
        creditors_header_block = build_creditors_header_block(creditors)
    else:
        creditors_header_block = "–°–≤–µ–¥–µ–Ω–∏—è –æ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞—Ö –Ω–µ –ø—Ä–µ–¥—Å—Ç–∞–≤–ª–µ–Ω—ã."

    # vehicle_block: –¥–µ—Ñ–æ–ª—Ç –ø—Ä–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤–∏–∏
    vehicle_block = ""
    try:
        vehicle_block = build_vehicle_block(card) or ""
    except (KeyError, TypeError, AttributeError) as e:
        logger.warning(f"Failed to build vehicle block: {e}")
        vehicle_block = ""
    if not str(vehicle_block).strip():
        vehicle_block = "–¢—Ä–∞–Ω—Å–ø–æ—Ä—Ç–Ω—ã–µ —Å—Ä–µ–¥—Å—Ç–≤–∞: –æ—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç."

    # --- —Å—Ç–∞—Ç—É—Å –ò–ü (—É–º–Ω–∞—è –ª–æ–≥–∏–∫–∞: —Å–ø—Ä–∞–≤–∫–∞ –∏–ª–∏ –ï–ì–†–ò–ü) ---
    ip_cert_number = (card.get("ip_certificate_number") or "").strip()
    ip_cert_date = (card.get("ip_certificate_date") or "").strip()

    if ip_cert_number and ip_cert_date:
        ip_status_text = (
            "–Ω–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω –≤ –∫–∞—á–µ—Å—Ç–≤–µ –∏–Ω–¥–∏–≤–∏–¥—É–∞–ª—å–Ω–æ–≥–æ –ø—Ä–µ–¥–ø—Ä–∏–Ω–∏–º–∞—Ç–µ–ª—è, "
            f"—á—Ç–æ –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–∞–µ—Ç—Å—è —Å–ø—Ä–∞–≤–∫–æ–π ‚Ññ {ip_cert_number} –æ—Ç {ip_cert_date}."
        )
    else:
        ip_status_text = (
            "–Ω–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω –≤ –∫–∞—á–µ—Å—Ç–≤–µ –∏–Ω–¥–∏–≤–∏–¥—É–∞–ª—å–Ω–æ–≥–æ –ø—Ä–µ–¥–ø—Ä–∏–Ω–∏–º–∞—Ç–µ–ª—è, "
            "—á—Ç–æ –ø–æ–¥—Ç–≤–µ—Ä–∂–¥–∞–µ—Ç—Å—è —Å–≤–µ–¥–µ–Ω–∏—è–º–∏ –∏–∑ –ï–ì–†–ò–ü"
        )

    # –Ω–æ—Ä–º–∞–ª–∏–∑–∞—Ü–∏—è, —á—Ç–æ–±—ã –Ω–µ –±—ã–ª–æ '–ï–ì–†–ò–ü..'
    ip_status_text = (ip_status_text or "").strip()
    while ".." in ip_status_text:
        ip_status_text = ip_status_text.replace("..", ".")

    mapping = {
        "attachments_list": attachments_list,
        "certificate_date": certificate_date,
        "certificate_number": certificate_number,
        "court_address": court_address,
        "court_name": court_name,

        # –ö—Ä–µ–¥–∏—Ç–æ—Ä—ã: —à–∞–ø–∫–∞ + –æ—Å–Ω–æ–≤–Ω–æ–π –±–ª–æ–∫
        "creditors_block": creditors_block,
        "creditors_header_block": creditors_header_block,

        "date": datetime.now().strftime("%d.%m.%Y"),

        "debtor_address": debtor_address,
        "debtor_birth_date": debtor_birth_date,
        "debtor_full_name": debtor_full_name,

        # –í —à–∞–±–ª–æ–Ω–µ –µ—Å—Ç—å –∏ –æ–±—ã—á–Ω—ã–µ, –∏ *_or_absent
        "debtor_inn": debtor_inn if debtor_inn != "–Ω–µ —É–∫–∞–∑–∞–Ω–æ" else "",
        "debtor_inn_or_absent": debtor_inn if debtor_inn != "–Ω–µ —É–∫–∞–∑–∞–Ω–æ" else "–æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç",

        "debtor_snils": debtor_snils if debtor_snils != "–Ω–µ —É–∫–∞–∑–∞–Ω–æ" else "",
        "debtor_snils_or_absent": debtor_snils if debtor_snils != "–Ω–µ —É–∫–∞–∑–∞–Ω–æ" else "–æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç",

        "debtor_phone_or_absent": debtor_phone if debtor_phone != "–Ω–µ —É–∫–∞–∑–∞–Ω–æ" else "–æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç",

        # –ü–∞—Å–ø–æ—Ä—Ç: –∫–ª—é—á–∏ –¥–æ–ª–∂–Ω—ã —Å–æ–≤–ø–∞–¥–∞—Ç—å —Å –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–∞–º–∏ —à–∞–±–ª–æ–Ω–∞
        "passport_series": passport_series or "",
        "passport_number": passport_number or "",
        "passport_issued_by": debtor_passport_issued_by if debtor_passport_issued_by != "–Ω–µ —É–∫–∞–∑–∞–Ω–æ" else "",
        "passport_date": debtor_passport_date if debtor_passport_date != "–Ω–µ —É–∫–∞–∑–∞–Ω–æ" else "",
        "passport_code": debtor_passport_code if debtor_passport_code != "–Ω–µ —É–∫–∞–∑–∞–Ω–æ" else "",

        # –≠—Ç–∏ –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä—ã –µ—Å—Ç—å –≤ —à–∞–±–ª–æ–Ω–µ (—Ç—ã –∏—Ö –ø–æ–∫–∞–∑—ã–≤–∞–ª –≤ —Å–ø–∏—Å–∫–µ)
        "debtor_last_name_initials": build_debtor_last_name_initials(card),

        "financial_manager_info": financial_manager_info,
        "family_status_block": build_family_status_block(card),
        "ip_status_text": ip_status_text,

        "marital_status": marital_status,

        "total_debt_kopeks": total_debt_kopeks,
        "total_debt_rubles": total_debt_rubles,

        "vehicle_block": vehicle_block,

        "deposit_deferral_request": deposit_deferral_request,
    }

    # –≥–µ–Ω–¥–µ—Ä–Ω—ã–µ —Ñ–æ—Ä–º—ã (debtor_having_word, debtor_registered_word, debtor_living_word,
    # debtor_not_registered_word, debtor_insolvent_word)
    try:
        gender_forms = build_gender_forms(card.get("debtor_gender"))
        if isinstance(gender_forms, dict):
            mapping.update(gender_forms)
    except (KeyError, TypeError, AttributeError) as e:
        # –µ—Å–ª–∏ –ø–æ–ª –Ω–µ –∑–∞–ø–æ–ª–Ω–µ–Ω –∏–ª–∏ —Ñ—É–Ω–∫—Ü–∏—è —É–ø–∞–ª–∞ ‚Äî —Å—Ç–∞–≤–∏–º –Ω–µ–π—Ç—Ä–∞–ª—å–Ω—ã–µ –∑–Ω–∞—á–µ–Ω–∏—è
        logger.warning(f"Failed to build gender forms: {e}")
        mapping.update(
            {
                "debtor_having_word": "–∏–º–µ—é—â–∏–π(–∞—è)",
                "debtor_registered_word": "–∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω–Ω—ã–π(–∞—è)",
                "debtor_living_word": "–ø—Ä–æ–∂–∏–≤–∞—é—â–∏–π(–∞—è)",
                "debtor_not_registered_word": "–Ω–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω(–∞)",
                "debtor_insolvent_word": "–Ω–µ–ø–ª–∞—Ç—ë–∂–µ—Å–ø–æ—Å–æ–±–Ω—ã–π(–∞—è)",
            }
        )


    _replace_placeholders_strong(doc, mapping)
    # –≤—Ç–æ—Ä–æ–π –ø—Ä–æ—Ö–æ–¥ ‚Äî –¥–æ–±–∏–≤–∞–µ–º –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä—ã, —Ä–∞–∑–æ—Ä–≤–∞–Ω–Ω—ã–µ Word –ø–æ runs
    for p in doc.paragraphs:
        for run in p.runs:
            if "{{" in run.text:
                for k, v in mapping.items():
                    run.text = run.text.replace(f"{{{{{k}}}}}", "" if v is None else str(v))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if "{{" in run.text:
                            for k, v in mapping.items():
                                run.text = run.text.replace(f"{{{{{k}}}}}", "" if v is None else str(v))


    # –ö–æ–Ω—Ç—Ä–æ–ª—å: –Ω–µ –¥–æ–ª–∂–Ω–æ –æ—Å—Ç–∞—Ç—å—Å—è {{...}}
    def _has_unreplaced_placeholders(d: Document) -> bool:
        for p in d.paragraphs:
            if "{{" in (p.text or ""):
                return True
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{{" in (p.text or ""):
                            return True
        return False

    if _has_unreplaced_placeholders(doc):
        import re

        # –¥–∏–∞–≥–Ω–æ—Å—Ç–∏—Ä—É–µ–º, —á—Ç–æ –∏–º–µ–Ω–Ω–æ –æ—Å—Ç–∞–ª–æ—Å—å –≤ –¥–æ–∫—É–º–µ–Ω—Ç–µ
        left = set()

        def _scan_paragraph_for_left(p):
            txt = p.text or ""
            for m in re.findall(r"\{\{[^}]+\}\}", txt):
                left.add(m)

        for p in doc.paragraphs:
            _scan_paragraph_for_left(p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _scan_paragraph_for_left(p)
                    for nested in cell.tables:
                        for nrow in nested.rows:
                            for ncell in nrow.cells:
                                for p in ncell.paragraphs:
                                    _scan_paragraph_for_left(p)

        import logging
        logging.exception("UNREPLACED_PLACEHOLDERS: %s", sorted(left))

        raise ValueError("–í –¥–æ–∫—É–º–µ–Ω—Ç–µ –æ—Å—Ç–∞–ª–∏—Å—å –Ω–µ –∑–∞–º–µ–Ω—ë–Ω–Ω—ã–µ –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä—ã –≤–∏–¥–∞ {{...}}")

    fname = f"bankruptcy_petition_case_{cid}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    case_dir = GENERATED_DIR / "cases" / str(cid)
    case_dir.mkdir(parents=True, exist_ok=True)
    out_path = case_dir / fname
    doc.save(out_path)
    return out_path


async def _selected_case_id(state: FSMContext) -> int | None:
    data = await state.get_data()
    try:
        return int(data.get("docs_case_id"))
    except (TypeError, ValueError):
        return None

settings = load_settings()

BOT_TOKEN = settings["BOT_TOKEN"]
AUTH_KEY = settings["GIGACHAT_AUTH_KEY"]
SCOPE = settings["GIGACHAT_SCOPE"]
MODEL = settings["GIGACHAT_MODEL"]

RAW_ALLOWED = settings["RAW_ALLOWED"]
RAW_ADMINS = settings["RAW_ADMINS"]
GENERATED_DIR = settings["GENERATED_DIR"]

DB_PATH = settings["DB_PATH"]

# Initialize cases_db module with database path
from bankrot_bot.services.cases_db import (
    init_cases_db,
    list_cases,
    get_case,
    create_case,
    update_case_fields,
    update_case_meta,
    get_case_card,
    validate_case_card,
    get_profile,
    upsert_profile,
    migrate_case_cards_table,
    CASE_CARD_REQUIRED_FIELDS,
    CASE_CARDS_ALLOWED_COLUMNS,
)
init_cases_db(DB_PATH)

def _parse_ids(s: str) -> set[int]:
    out = set()
    for x in (s.split(",") if s else []):
        x = x.strip()
        if x.isdigit():
            out.add(int(x))
    return out


ALLOWED_USERS = _parse_ids(RAW_ALLOWED)
ADMIN_USERS = _parse_ids(RAW_ADMINS)

# Import authorization functions from shared module to break circular imports
from bankrot_bot.shared import is_allowed, is_admin, init_allowed_users


# =========================
# sqlite (cases)
# =========================
def init_db() -> None:
    Path(DB_PATH).parent.mkdir(parents=True, exist_ok=True)
    with sqlite3.connect(DB_PATH) as con:
        con.execute("PRAGMA journal_mode=WAL;")

        # ===== cases =====
        con.execute("""
        CREATE TABLE IF NOT EXISTS cases (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            owner_user_id INTEGER NOT NULL,
            code_name TEXT NOT NULL,
            case_number TEXT,
            court TEXT,
            judge TEXT,
            fin_manager TEXT,
            stage TEXT,
            notes TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );
        """)

        # ===== profiles (–¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤) =====
        con.execute("""
        CREATE TABLE IF NOT EXISTS profiles (
            owner_user_id INTEGER PRIMARY KEY,
            full_name TEXT,
            role TEXT,
            address TEXT,
            phone TEXT,
            email TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        );
        """)

        con.execute(
            """
            CREATE TABLE IF NOT EXISTS case_cards (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                owner_user_id INTEGER NOT NULL,
                case_id INTEGER NOT NULL,
                data TEXT,
                court_name TEXT,
                court_address TEXT,
                judge_name TEXT,
                debtor_full_name TEXT,
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(owner_user_id, case_id)
            );
            """
        )

        migrate_case_cards_table(con)
        con.commit()


def upsert_case_card(owner_user_id: int, case_id: int, data: dict[str, Any]) -> None:
    migrate_case_cards_table()
    with sqlite3.connect(DB_PATH) as con:
        cur = con.cursor()
        cur.execute("PRAGMA table_info(case_cards)")
        columns = {row[1] for row in cur.fetchall()}
        cur.execute(
            """
            SELECT data FROM case_cards
             WHERE owner_user_id = ?
               AND case_id = ?
            """,
            (owner_user_id, case_id),
        )
        row = cur.fetchone()
        current: dict[str, Any] = {}
        if row and row[0]:
            try:
                current = json.loads(row[0])
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse existing case card JSON for case_id={case_id}: {e}")
                current = {}

        current.update(data)

        payload = json.dumps(current, ensure_ascii=False)

        insert_columns = ["owner_user_id", "case_id", "data"]
        placeholders = ["?", "?", "?"]
        values: list[Any] = [owner_user_id, case_id, payload]

        if "created_at" in columns:
            insert_columns.append("created_at")
            placeholders.append("CURRENT_TIMESTAMP")

        if "updated_at" in columns:
            insert_columns.append("updated_at")
            placeholders.append("CURRENT_TIMESTAMP")

        update_set_parts = ["data = excluded.data"]
        if "updated_at" in columns:
            update_set_parts.append("updated_at = CURRENT_TIMESTAMP")

        sql = f"""
            INSERT INTO case_cards ({', '.join(insert_columns)})
            VALUES ({', '.join(placeholders)})
            ON CONFLICT(owner_user_id, case_id) DO UPDATE SET
                {', '.join(update_set_parts)}
        """

        cur.execute(sql, values)
        con.commit()


# =========================
# bot logic
# =========================
from aiogram.fsm.storage.redis import RedisStorage

# Configure Redis storage for FSM
redis_url = os.getenv("REDIS_URL", "redis://localhost:6379/0")
storage = RedisStorage.from_url(redis_url)
dp = Dispatcher(storage=storage)

# =========================
# ROUTER REGISTRATION - PRIORITY ORDER MATTERS!
# =========================
# Router order determines handler priority (first registered = highest priority).
# This prevents message routing conflicts and ensures predictable behavior.
#
# PRIORITY ORDER (highest to lowest):
# 1. COMMANDS - cases_handlers.router (Command handlers with explicit StateFilter)
#    - /newcase, /mycases, /case, /setactive, /editcase, /cancel
#    - These must run FIRST to catch commands before FSM or text handlers
#
# 2. FSM STATES - newcase_fsm.router (FSM handlers with explicit StateFilter)
#    - F.text == "‚ûï –ù–æ–≤–æ–µ –¥–µ–ª–æ" with StateFilter(None) - only when NOT in FSM
#    - NewCase FSM states with StateFilter - only when IN specific FSM state
#    - StateFilter prevents conflicts with commands and other handlers
#
# 3. DIRECT DP HANDLERS - Callbacks and other handlers registered directly on dp
#    - @dp.callback_query() - Callback query handlers
#    - @dp.message(StateFilter(...)) - FSM handlers (AddParty, AddAsset, etc.)
#    - These run LAST as fallback/catch-all handlers
#
# CRITICAL RULES:
# ‚úÖ ALL FSM state handlers MUST use StateFilter to prevent catching commands
# ‚úÖ Command handlers run first (registered routers before direct dp handlers)
# ‚úÖ Text filters (F.text) MUST have StateFilter or be very specific
# ‚úÖ Callback handlers can be registered anywhere (F.data is specific enough)
#
# DO NOT change router order without understanding priority implications!
# =========================

# 1. Commands (highest priority)
dp.include_router(cases_handlers.router)

# 2. FSM states with StateFilter
# Register probability router (Phase 9)
# dp.include_router(probability_handlers.router)
# Register new case FSM router (Phase 12 fix)
dp.include_router(newcase_fsm.router)

# 3. Callback handlers will be registered later after helper functions are defined

# 4. Direct dp handlers (callbacks, FSM, etc.) registered below have lowest priority

USER_FLOW: Dict[int, Dict[str, Any]] = {}
LAST_RESULT: Dict[int, str] = {}


def cancel_flow(uid: int) -> None:
    USER_FLOW.pop(uid, None)


# main_keyboard() definition removed - see line 4336 for active implementation


def export_keyboard() -> InlineKeyboardMarkup:
    """–ö–ª–∞–≤–∏–∞—Ç—É—Ä–∞ —ç–∫—Å–ø–æ—Ä—Ç–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞."""
    kb = InlineKeyboardBuilder()
    kb.button(text="üìÑ –≠–∫—Å–ø–æ—Ä—Ç (–ø–æ–∫–∞–∑–∞—Ç—å —Ç–µ–∫—Å—Ç)", callback_data="export:word")
    kb.adjust(1)
    return kb.as_markup()


def court_type_keyboard() -> InlineKeyboardMarkup:
    """–ö–ª–∞–≤–∏–∞—Ç—É—Ä–∞ –≤—ã–±–æ—Ä–∞ —Ç–∏–ø–∞ —Å—É–¥–∞."""
    kb = InlineKeyboardBuilder()
    kb.button(text="–ê—Ä–±–∏—Ç—Ä–∞–∂–Ω—ã–π —Å—É–¥", callback_data="motion:court:arbitr")
    kb.button(text="–°—É–¥ –æ–±—â–µ–π —é—Ä–∏—Å–¥–∏–∫—Ü–∏–∏", callback_data="motion:court:general")
    kb.adjust(1)
    return kb.as_markup()


def motion_actions_keyboard() -> InlineKeyboardMarkup:
    """–ö–ª–∞–≤–∏–∞—Ç—É—Ä–∞ –¥–µ–π—Å—Ç–≤–∏–π –¥–ª—è —Ö–æ–¥–∞—Ç–∞–π—Å—Ç–≤–∞."""
    kb = InlineKeyboardBuilder()
    kb.button(text="–û—Ç–º–µ–Ω–∞", callback_data="flow:cancel")
    kb.adjust(1)
    return kb.as_markup()


def settlement_actions_keyboard() -> InlineKeyboardMarkup:
    """–ö–ª–∞–≤–∏–∞—Ç—É—Ä–∞ –¥–µ–π—Å—Ç–≤–∏–π –¥–ª—è –º–∏—Ä–æ–≤–æ–≥–æ —Å–æ–≥–ª–∞—à–µ–Ω–∏—è."""
    kb = InlineKeyboardBuilder()
    kb.button(text="–û—Ç–º–µ–Ω–∞", callback_data="flow:cancel")
    kb.adjust(1)
    return kb.as_markup()


MOTION_STEPS = [
    ("fio", "–§–ò–û –∑–∞—è–≤–∏—Ç–µ–ª—è (–¥–æ–ª–∂–Ω–∏–∫–∞):"),
    ("case_number", "–ù–æ–º–µ—Ä –¥–µ–ª–∞ (–µ—Å–ª–∏ –µ—Å—Ç—å) –∏–ª–∏ –Ω–∞–ø–∏—à–∏ ¬´–Ω–µ—Ç¬ª:"),
    ("court", "–°—É–¥ (–ø–æ–ª–Ω–æ–µ –Ω–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ):"),
    ("judge", "–°—É–¥—å—è (–µ—Å–ª–∏ –∏–∑–≤–µ—Å—Ç–Ω–æ) –∏–ª–∏ ¬´–Ω–µ—Ç¬ª:"),
    ("reason", "–ü—Ä–∏—á–∏–Ω–∞ —Ö–æ–¥–∞—Ç–∞–π—Å—Ç–≤–∞ (–∫—Ä–∞—Ç–∫–æ):"),
]

SETTLEMENT_STEPS = [
    ("parties", "–°—Ç–æ—Ä–æ–Ω—ã (–∫—Ç–æ —Å –∫–µ–º –∑–∞–∫–ª—é—á–∞–µ—Ç –º–∏—Ä–æ–≤–æ–µ):"),
    ("dispute", "–°—É—Ç—å —Å–ø–æ—Ä–∞ / —á—Ç–æ —É—Ä–µ–≥—É–ª–∏—Ä—É–µ–º:"),
    ("terms", "–£—Å–ª–æ–≤–∏—è (—á—Ç–æ –∏ –≤ –∫–∞–∫–∏–µ —Å—Ä–æ–∫–∏):"),
    ("expenses", "–†–∞—Å—Ö–æ–¥—ã/–≥–æ—Å–ø–æ—à–ª–∏–Ω–∞ (–µ—Å–ª–∏ –µ—Å—Ç—å) –∏–ª–∏ ¬´–Ω–µ—Ç¬ª:"),
    ("execution", "–ò—Å–ø–æ–ª–Ω–µ–Ω–∏–µ/–æ—Ç–≤–µ—Ç—Å—Ç–≤–µ–Ω–Ω–æ—Å—Ç—å –∑–∞ –Ω–∞—Ä—É—à–µ–Ω–∏–µ:"),
    ("other", "–û—Å–æ–±—ã–µ —É—Å–ª–æ–≤–∏—è (–µ—Å–ª–∏ –µ—Å—Ç—å) –∏–ª–∏ ¬´–Ω–µ—Ç¬ª:"),
]


def system_prompt_for_motion(court_type: str) -> str:
    return (
        "–¢—ã ‚Äî —é—Ä–∏—Å—Ç –ø–æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤—É –≤ –†–æ—Å—Å–∏–∏. –°–æ—Å—Ç–∞–≤—å –ø—Ä–æ–µ–∫—Ç —Ö–æ–¥–∞—Ç–∞–π—Å—Ç–≤–∞ –æ–± —É—á–∞—Å—Ç–∏–∏ –≤ –∑–∞—Å–µ–¥–∞–Ω–∏–∏ –æ–Ω–ª–∞–π–Ω "
        "–∏–ª–∏ –ø–æ—Å—Ä–µ–¥—Å—Ç–≤–æ–º –í–ö–°. –°—Ç–∏–ª—å –æ—Ñ–∏—Ü–∏–∞–ª—å–Ω—ã–π, –∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–π, –±–µ–∑ –≤—ã–¥—É–º—ã–≤–∞–Ω–∏—è —Ñ–∞–∫—Ç–æ–≤."
        f" –¢–∏–ø —Å—É–¥–∞: {court_type}."
    )


def system_prompt_for_settlement() -> str:
    return (
        "–¢—ã ‚Äî —é—Ä–∏—Å—Ç –ø–æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤—É –≤ –†–æ—Å—Å–∏–∏. –°–æ—Å—Ç–∞–≤—å –ø—Ä–æ–µ–∫—Ç –º–∏—Ä–æ–≤–æ–≥–æ —Å–æ–≥–ª–∞—à–µ–Ω–∏—è. "
        "–°—Ç–∏–ª—å –æ—Ñ–∏—Ü–∏–∞–ª—å–Ω—ã–π, –±–µ–∑ –≤—ã–¥—É–º—ã–≤–∞–Ω–∏—è —Ñ–∞–∫—Ç–æ–≤; –µ—Å–ª–∏ –¥–∞–Ω–Ω—ã—Ö –Ω–µ —Ö–≤–∞—Ç–∞–µ—Ç ‚Äî –æ—Å—Ç–∞–≤—å –º–µ—Å—Ç–∞ –¥–ª—è –∑–∞–ø–æ–ª–Ω–µ–Ω–∏—è."
    )


def _val(ans: Dict[str, str], key: str) -> str:
    v = (ans.get(key) or "").strip()
    return v if v else "–Ω–µ —É–∫–∞–∑–∞–Ω–æ"


def build_motion_user_text(ans: Dict[str, str], court_type: str) -> str:
    return (
        f"–§–ò–û: {_val(ans,'fio')}\n"
        f"–ù–æ–º–µ—Ä –¥–µ–ª–∞: {_val(ans,'case_number')}\n"
        f"–°—É–¥: {_val(ans,'court')}\n"
        f"–°—É–¥—å—è: {_val(ans,'judge')}\n"
        f"–ü—Ä–∏—á–∏–Ω–∞: {_val(ans,'reason')}\n"
        f"–¢–∏–ø —Å—É–¥–∞: {court_type}\n"
        "–°—Ñ–æ—Ä–º–∏—Ä—É–π —Ç–µ–∫—Å—Ç —Ö–æ–¥–∞—Ç–∞–π—Å—Ç–≤–∞."
    )


def build_settlement_user_text(ans: Dict[str, str]) -> str:
    return (
        f"–°—Ç–æ—Ä–æ–Ω—ã: {_val(ans,'parties')}\n"
        f"–°—É—Ç—å —É—Ä–µ–≥—É–ª–∏—Ä–æ–≤–∞–Ω–∏—è: {_val(ans,'dispute')}\n"
        f"–£—Å–ª–æ–≤–∏—è: {_val(ans,'terms')}\n"
        f"–†–∞—Å—Ö–æ–¥—ã: {_val(ans,'expenses')}\n"
        f"–ò—Å–ø–æ–ª–Ω–µ–Ω–∏–µ/–æ—Ç–≤–µ—Ç—Å—Ç–≤–µ–Ω–Ω–æ—Å—Ç—å: {_val(ans,'execution')}\n"
        f"–û—Å–æ–±—ã–µ —É—Å–ª–æ–≤–∏—è: {_val(ans,'other')}\n"
        "–°—Ñ–æ—Ä–º–∏—Ä—É–π –ø—Ä–æ–µ–∫—Ç –º–∏—Ä–æ–≤–æ–≥–æ —Å–æ–≥–ª–∞—à–µ–Ω–∏—è."
    )

# =========================
# menu (new)
# =========================

@dp.message(CommandStart())
async def cmd_start(message: Message) -> None:
    """
    Handle /start command.

    Shows welcome message with main menu (both reply keyboard and inline keyboard).
    Reply keyboard provides persistent bottom navigation buttons.

    Args:
        message: Incoming /start command message
    """
    uid = message.from_user.id
    if not is_allowed(uid):
        logger.warning(f"Unauthorized /start attempt by user {uid}")
        return
    cancel_flow(uid)

    # Import refactored main menu (inline) and reply keyboard
    from keyboards import main_menu

    text = (
        "üè† –î–æ–±—Ä–æ –ø–æ–∂–∞–ª–æ–≤–∞—Ç—å!\n\n"
        "–ë–æ—Ç-–ø–æ–º–æ—â–Ω–∏–∫ –ø–æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤—É —Ñ–∏–∑–∏—á–µ—Å–∫–∏—Ö –ª–∏—Ü.\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ —Ä–∞–∑–¥–µ–ª –¥–ª—è —Ä–∞–±–æ—Ç—ã:"
    )

    # Send with reply keyboard (bottom buttons) for persistent navigation
    await message.answer(text, reply_markup=main_menu_kb())

    # Optionally send inline keyboard for additional options
    # await message.answer("–ò–ª–∏ –≤—ã–±–µ—Ä–∏—Ç–µ –∏–∑ –º–µ–Ω—é:", reply_markup=main_menu())


# =========================
# Reply Keyboard Handlers
# =========================

@dp.message(StateFilter(None), F.text == "‚ûï –ù–æ–≤–æ–µ –¥–µ–ª–æ")
async def reply_new_case(message: Message) -> None:
    """
    Handle '‚ûï –ù–æ–≤–æ–µ –¥–µ–ª–æ' reply keyboard button.

    Shows cases menu with options to create new case or view existing cases.
    StateFilter(None) ensures this only fires when user is NOT in FSM.

    Args:
        message: Message with reply keyboard button text
    """
    uid = message.from_user.id
    if not is_allowed(uid):
        logger.warning(f"Unauthorized new case access attempt by user {uid}")
        return

    await message.answer(
        "üìã –ú–æ–∏ –¥–µ–ª–∞ –æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–µ:",
        reply_markup=cases_menu_ikb()
    )


@dp.message(StateFilter(None), F.text == "üë§ –ú–æ–π –ø—Ä–æ—Ñ–∏–ª—å")
async def reply_my_profile(message: Message) -> None:
    """
    Handle 'üë§ –ú–æ–π –ø—Ä–æ—Ñ–∏–ª—å' reply keyboard button.

    Shows user profile with inline keyboard options.
    StateFilter(None) ensures this only fires when user is NOT in FSM.

    Args:
        message: Message with reply keyboard button text
    """
    uid = message.from_user.id
    if not is_allowed(uid):
        logger.warning(f"Unauthorized profile access attempt by user {uid}")
        return

    await message.answer("üë§ –ú–æ–π –ø—Ä–æ—Ñ–∏–ª—å:", reply_markup=profile_ikb())


@dp.message(StateFilter(None), F.text == "üìÑ –î–æ–∫—É–º–µ–Ω—Ç—ã")
async def reply_documents(message: Message) -> None:
    """
    Handle 'üìÑ –î–æ–∫—É–º–µ–Ω—Ç—ã' reply keyboard button.

    Shows public documents catalog with inline keyboard options.
    StateFilter(None) ensures this only fires when user is NOT in FSM.

    Args:
        message: Message with reply keyboard button text
    """
    uid = message.from_user.id
    if not is_allowed(uid):
        logger.warning(f"Unauthorized documents access attempt by user {uid}")
        return

    await message.answer(
        "üìÑ –ü—É–±–ª–∏—á–Ω—ã–π –∫–∞—Ç–∞–ª–æ–≥ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ –∫–∞—Ç–µ–≥–æ—Ä–∏—é:",
        reply_markup=docs_catalog_ikb()
    )


@dp.message(StateFilter(None), F.text == "‚ùì –ü–æ–º–æ—â—å")
async def reply_help(message: Message) -> None:
    """
    Handle '‚ùì –ü–æ–º–æ—â—å' reply keyboard button.

    Shows help menu with inline keyboard options.
    StateFilter(None) ensures this only fires when user is NOT in FSM.

    Args:
        message: Message with reply keyboard button text
    """
    uid = message.from_user.id
    if not is_allowed(uid):
        logger.warning(f"Unauthorized help access attempt by user {uid}")
        return

    await message.answer(
        "‚ùì –ü–æ–º–æ—â—å\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ —Ç–µ–º—É:",
        reply_markup=help_ikb()
    )


@dp.callback_query(F.data == "menu:home")
async def menu_home(call: CallbackQuery) -> None:
    """
    Navigate to home/main menu.

    Args:
        call: Callback query from menu navigation
    """
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return
    await call.message.answer("–ì–ª–∞–≤–Ω–æ–µ –º–µ–Ω—é:", reply_markup=home_ikb())
    await call.answer()


@dp.callback_query(F.data == "menu:profile")
async def menu_profile(call: CallbackQuery) -> None:
    """
    Navigate to user profile menu.

    Args:
        call: Callback query from menu navigation
    """
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return
    await call.message.answer("üë§ –ú–æ–π –ø—Ä–æ—Ñ–∏–ª—å:", reply_markup=profile_ikb())
    await call.answer()


@dp.callback_query(F.data == "menu:docs")
async def menu_docs(call: CallbackQuery) -> None:
    """
    Navigate to public documents catalog.

    –ü—É–±–ª–∏—á–Ω—ã–π –∫–∞—Ç–∞–ª–æ–≥ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ - –¥–æ—Å—Ç—É–ø–µ–Ω –≤—Å–µ–º –∞–≤—Ç–æ—Ä–∏–∑–æ–≤–∞–Ω–Ω—ã–º –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è–º.

    Args:
        call: Callback query from menu navigation
    """
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return
    await call.message.answer(
        "üìÑ –ü—É–±–ª–∏—á–Ω—ã–π –∫–∞—Ç–∞–ª–æ–≥ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤\n\n"
        "–ó–¥–µ—Å—å –≤—ã –Ω–∞–π–¥–µ—Ç–µ —à–∞–±–ª–æ–Ω—ã –∏ –æ–±—Ä–∞–∑—Ü—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –¥–ª—è –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–∞.\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ –∫–∞—Ç–µ–≥–æ—Ä–∏—é:",
        reply_markup=docs_catalog_ikb()
    )
    await call.answer()


@dp.callback_query(F.data == "menu:help")
async def menu_help(call: CallbackQuery) -> None:
    """
    Navigate to help section.

    –ü–æ–¥–º–µ–Ω—é —Ä–∞–∑–¥–µ–ª–∞ –ü–æ–º–æ—â—å.

    Args:
        call: Callback query from menu navigation
    """
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return
    await call.message.answer(
        "‚ùì –†–∞–∑–¥–µ–ª –ø–æ–º–æ—â–∏\n\n"
        "–í—ã–±–µ—Ä–∏—Ç–µ –∏–Ω—Ç–µ—Ä–µ—Å—É—é—â—É—é —Ç–µ–º—É:",
        reply_markup=help_ikb(),
    )
    await call.answer()


# ========== –ù–æ–≤—ã–µ —Ö—ç–Ω–¥–ª–µ—Ä—ã –≥–ª–∞–≤–Ω–æ–≥–æ –º–µ–Ω—é ==========

@dp.callback_query(F.data == "menu:my_cases")
async def menu_my_cases(call: CallbackQuery, state: FSMContext):
    """–†–∞–∑–¥–µ–ª ¬´–ú–æ–∏ –¥–µ–ª–∞¬ª - –∏–Ω—Ç–µ–≥—Ä–∞—Ü–∏—è —Å –º–æ–¥—É–ª–µ–º cases."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    rows = list_cases(uid)

    # –ü–æ–ª—É—á–∏—Ç—å –∞–∫—Ç–∏–≤–Ω–æ–µ –¥–µ–ª–æ –∏–∑ state (–µ—Å–ª–∏ –µ—Å—Ç—å)
    data = await state.get_data()
    active_case_id = data.get("active_case_id")

    text = "üìÇ –ú–æ–∏ –¥–µ–ª–∞\n\n"
    if rows:
        text += f"–£ –≤–∞—Å {len(rows)} –¥–µ–ª(–∞/–æ).\n"
        if active_case_id:
            text += f"–ê–∫—Ç–∏–≤–Ω–æ–µ –¥–µ–ª–æ: #{active_case_id}\n"
        text += "–í—ã–±–µ—Ä–∏—Ç–µ –¥–µ–ª–æ –¥–ª—è —Ä–∞–±–æ—Ç—ã –∏–ª–∏ —Å–æ–∑–¥–∞–π—Ç–µ –Ω–æ–≤–æ–µ."
    else:
        text += "–£ –≤–∞—Å –ø–æ–∫–∞ –Ω–µ—Ç –¥–µ–ª. –°–æ–∑–¥–∞–π—Ç–µ –ø–µ—Ä–≤–æ–µ –¥–µ–ª–æ –¥–ª—è —Ä–∞–±–æ—Ç—ã."

    await call.message.answer(text, reply_markup=my_cases_ikb(rows, active_case_id))
    await call.answer()


@dp.callback_query(F.data == "ai:placeholder")
async def ai_placeholder(call: CallbackQuery):
    """–ó–∞–≥–ª—É—à–∫–∞ –¥–ª—è –ò–ò-–ø–æ–º–æ—â–Ω–∏–∫–∞."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    await call.answer("ü§ñ –ò–ò-–ø–æ–º–æ—â–Ω–∏–∫ –≤ —Ä–∞–∑—Ä–∞–±–æ—Ç–∫–µ. –°–∫–æ—Ä–æ –±—É–¥–µ—Ç –¥–æ—Å—Ç—É–ø–µ–Ω!", show_alert=True)


# ========== –•—ç–Ω–¥–ª–µ—Ä—ã —Ä–∞–∑–¥–µ–ª–∞ ¬´–ü–æ–º–æ—â—å¬ª ==========

@dp.callback_query(F.data == "help:howto")
async def help_howto(call: CallbackQuery):
    """–ö–∞–∫ –ø–æ–ª—å–∑–æ–≤–∞—Ç—å—Å—è –±–æ—Ç–æ–º."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    text = (
        "üìñ –ö–∞–∫ –ø–æ–ª—å–∑–æ–≤–∞—Ç—å—Å—è –±–æ—Ç–æ–º\n\n"
        "1Ô∏è‚É£ –ì–ª–∞–≤–Ω–æ–µ –º–µ–Ω—é\n"
        "–í –≥–ª–∞–≤–Ω–æ–º –º–µ–Ω—é —Ç—Ä–∏ —Ä–∞–∑–¥–µ–ª–∞:\n"
        "‚Ä¢ –ú–æ–∏ –¥–µ–ª–∞ - —É–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –¥–µ–ª–∞–º–∏ –æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–µ\n"
        "‚Ä¢ –î–æ–∫—É–º–µ–Ω—Ç—ã - –ø—É–±–ª–∏—á–Ω—ã–π –∫–∞—Ç–∞–ª–æ–≥ –æ–±—Ä–∞–∑—Ü–æ–≤\n"
        "‚Ä¢ –ü–æ–º–æ—â—å - —Å–ø—Ä–∞–≤–æ—á–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è\n\n"
        "2Ô∏è‚É£ –†–∞–±–æ—Ç–∞ —Å –¥–µ–ª–∞–º–∏\n"
        "‚Ä¢ –°–æ–∑–¥–∞–π—Ç–µ –∫–∞—Ä—Ç–æ—á–∫—É –¥–µ–ª–∞\n"
        "‚Ä¢ –ó–∞–ø–æ–ª–Ω–∏—Ç–µ –¥–∞–Ω–Ω—ã–µ –¥–æ–ª–∂–Ω–∏–∫–∞ –∏ –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤\n"
        "‚Ä¢ –ì–µ–Ω–µ—Ä–∏—Ä—É–π—Ç–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã –ø–æ –¥–µ–ª—É\n\n"
        "3Ô∏è‚É£ –ù–∞–≤–∏–≥–∞—Ü–∏—è\n"
        "–ò—Å–ø–æ–ª—å–∑—É–π—Ç–µ –∫–Ω–æ–ø–∫–∏ –¥–ª—è –ø–µ—Ä–µ–º–µ—â–µ–Ω–∏—è –º–µ–∂–¥—É —Ä–∞–∑–¥–µ–ª–∞–º–∏.\n"
        "–ö–Ω–æ–ø–∫–∞ üè† –≤—Å–µ–≥–¥–∞ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç –≤ –≥–ª–∞–≤–Ω–æ–µ –º–µ–Ω—é."
    )

    await call.message.answer(text, reply_markup=help_item_ikb())
    await call.answer()


@dp.callback_query(F.data == "help:cases")
async def help_cases(call: CallbackQuery):
    """–ß—Ç–æ —Ç–∞–∫–æ–µ –∫–∞—Ä—Ç–æ—á–∫–∏ –¥–µ–ª."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    text = (
        "üìã –ö–∞—Ä—Ç–æ—á–∫–∏ –¥–µ–ª\n\n"
        "–ö–∞—Ä—Ç–æ—á–∫–∞ –¥–µ–ª–∞ - —ç—Ç–æ —Å—Ç—Ä—É–∫—Ç—É—Ä–∏—Ä–æ–≤–∞–Ω–Ω–æ–µ —Ö—Ä–∞–Ω–∏–ª–∏—â–µ –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ "
        "–ø–æ –∫–æ–Ω–∫—Ä–µ—Ç–Ω–æ–º—É –¥–µ–ª—É –æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–µ.\n\n"
        "–ß—Ç–æ —Ö—Ä–∞–Ω–∏—Ç—Å—è:\n"
        "‚Ä¢ –î–∞–Ω–Ω—ã–µ –¥–æ–ª–∂–Ω–∏–∫–∞ (–§–ò–û, –∞–¥—Ä–µ—Å, –ø–∞—Å–ø–æ—Ä—Ç)\n"
        "‚Ä¢ –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞—Ö\n"
        "‚Ä¢ –°—É–º–º–∞ –∑–∞–¥–æ–ª–∂–µ–Ω–Ω–æ—Å—Ç–∏\n"
        "‚Ä¢ –î–æ–∫—É–º–µ–Ω—Ç—ã –ø–æ –¥–µ–ª—É\n"
        "‚Ä¢ –ò—Å—Ç–æ—Ä–∏—è –∏–∑–º–µ–Ω–µ–Ω–∏–π\n\n"
        "–ö–∞—Ä—Ç–æ—á–∫–∞ –ø—Ä–∏–≤—è–∑–∞–Ω–∞ –∫ –≤–∞—à–µ–º—É Telegram-–∞–∫–∫–∞—É–Ω—Ç—É –∏ –¥–æ—Å—Ç—É–ø–Ω–∞ —Ç–æ–ª—å–∫–æ –≤–∞–º.\n\n"
        "–ù–∞ –æ—Å–Ω–æ–≤–µ –¥–∞–Ω–Ω—ã—Ö –∫–∞—Ä—Ç–æ—á–∫–∏ –±–æ—Ç –º–æ–∂–µ—Ç –≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞—Ç—å —é—Ä–∏–¥–∏—á–µ—Å–∫–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã."
    )

    await call.message.answer(text, reply_markup=help_item_ikb())
    await call.answer()


@dp.callback_query(F.data == "help:docs")
async def help_docs(call: CallbackQuery):
    """–û –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ö."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    text = (
        "üìÑ –û –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ö\n\n"
        "–ë–æ—Ç —Ä–∞–±–æ—Ç–∞–µ—Ç —Å –¥–≤—É–º—è —Ç–∏–ø–∞–º–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤:\n\n"
        "1Ô∏è‚É£ –ü—É–±–ª–∏—á–Ω—ã–π –∫–∞—Ç–∞–ª–æ–≥\n"
        "–û–±—Ä–∞–∑—Ü—ã –∏ —à–∞–±–ª–æ–Ω—ã –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤, –¥–æ—Å—Ç—É–ø–Ω—ã–µ –≤—Å–µ–º –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è–º:\n"
        "‚Ä¢ –ó–∞—è–≤–ª–µ–Ω–∏—è\n"
        "‚Ä¢ –•–æ–¥–∞—Ç–∞–π—Å—Ç–≤–∞\n"
        "‚Ä¢ –ü—Ä–æ—á–∏–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã\n\n"
        "2Ô∏è‚É£ –î–æ–∫—É–º–µ–Ω—Ç—ã –ø–æ –¥–µ–ª—É\n"
        "–ì–µ–Ω–µ—Ä–∏—Ä—É—é—Ç—Å—è –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –Ω–∞ –æ—Å–Ω–æ–≤–µ –¥–∞–Ω–Ω—ã—Ö –≤–∞—à–µ–π –∫–∞—Ä—Ç–æ—á–∫–∏ –¥–µ–ª–∞.\n"
        "–ü—Ä–∏–≤—è–∑–∞–Ω—ã –∫ –∫–æ–Ω–∫—Ä–µ—Ç–Ω–æ–º—É –¥–µ–ª—É –∏ —Ö—Ä–∞–Ω—è—Ç—Å—è –≤ –≤–∞—à–µ–º –∞—Ä—Ö–∏–≤–µ.\n\n"
        "–í—Å–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã —Ñ–æ—Ä–º–∏—Ä—É—é—Ç—Å—è –≤ —Ñ–æ—Ä–º–∞—Ç–µ DOCX –∏ –≥–æ—Ç–æ–≤—ã –∫ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—é."
    )

    await call.message.answer(text, reply_markup=help_item_ikb())
    await call.answer()


@dp.callback_query(F.data == "help:contacts")
async def help_contacts(call: CallbackQuery):
    """–ö–æ–Ω—Ç–∞–∫—Ç—ã –∏ –æ–±—Ä–∞—Ç–Ω–∞—è —Å–≤—è–∑—å."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    text = (
        "‚úâÔ∏è –ö–æ–Ω—Ç–∞–∫—Ç—ã –∏ –æ–±—Ä–∞—Ç–Ω–∞—è —Å–≤—è–∑—å\n\n"
        "–ü–æ –≤—Å–µ–º –≤–æ–ø—Ä–æ—Å–∞–º —Ä–∞–±–æ—Ç—ã –±–æ—Ç–∞:\n"
        "‚Ä¢ –°–æ–æ–±—â–∏—Ç–µ –æ–± –æ—à–∏–±–∫–µ\n"
        "‚Ä¢ –ü—Ä–µ–¥–ª–æ–∂–∏—Ç–µ —É–ª—É—á—à–µ–Ω–∏–µ\n"
        "‚Ä¢ –ó–∞–¥–∞–π—Ç–µ –≤–æ–ø—Ä–æ—Å\n\n"
        "üìß Email: support@example.com\n"
        "üí¨ Telegram: @support_username\n\n"
        "–ú—ã –ø–æ—Å—Ç–æ—è–Ω–Ω–æ —Ä–∞–±–æ—Ç–∞–µ–º –Ω–∞–¥ —É–ª—É—á—à–µ–Ω–∏–µ–º —Å–µ—Ä–≤–∏—Å–∞. "
        "–í–∞—à–∏ –æ—Ç–∑—ã–≤—ã –ø–æ–º–æ–≥–∞—é—Ç –¥–µ–ª–∞—Ç—å –±–æ—Ç–∞ –ª—É—á—à–µ!"
    )

    await call.message.answer(text, reply_markup=help_item_ikb())
    await call.answer()


@dp.callback_query(F.data == "help:about")
async def help_about(call: CallbackQuery):
    """–û –±–æ—Ç–µ."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    text = (
        "‚ÑπÔ∏è –û –±–æ—Ç–µ\n\n"
        "Telegram-–±–æ—Ç –ø–æ–º–æ—â–Ω–∏–∫ –ø–æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤—É —Ñ–∏–∑–∏—á–µ—Å–∫–∏—Ö –ª–∏—Ü.\n\n"
        "–í–æ–∑–º–æ–∂–Ω–æ—Å—Ç–∏:\n"
        "‚Ä¢ –£–ø—Ä–∞–≤–ª–µ–Ω–∏–µ –∫–∞—Ä—Ç–æ—á–∫–∞–º–∏ –¥–µ–ª\n"
        "‚Ä¢ –ì–µ–Ω–µ—Ä–∞—Ü–∏—è —é—Ä–∏–¥–∏—á–µ—Å–∫–∏—Ö –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤\n"
        "‚Ä¢ –ö–∞—Ç–∞–ª–æ–≥ –æ–±—Ä–∞–∑—Ü–æ–≤ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤\n"
        "‚Ä¢ –°–ø—Ä–∞–≤–æ—á–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è\n\n"
        "–í–µ—Ä—Å–∏—è: 1.0.0\n"
        "–°—Ç–∞—Ç—É—Å: MVP (–º–∏–Ω–∏–º–∞–ª—å–Ω—ã–π —Ä–∞–±–æ—á–∏–π –ø—Ä–æ–¥—É–∫—Ç)\n\n"
        "–ë–æ—Ç –Ω–∞—Ö–æ–¥–∏—Ç—Å—è –≤ –∞–∫—Ç–∏–≤–Ω–æ–π —Ä–∞–∑—Ä–∞–±–æ—Ç–∫–µ. "
        "–°–ª–µ–¥–∏—Ç–µ –∑–∞ –æ–±–Ω–æ–≤–ª–µ–Ω–∏—è–º–∏!"
    )

    await call.message.answer(text, reply_markup=help_item_ikb())
    await call.answer()


# ========== –•—ç–Ω–¥–ª–µ—Ä—ã –ø—É–±–ª–∏—á–Ω–æ–≥–æ –∫–∞—Ç–∞–ª–æ–≥–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ ==========

@dp.callback_query(F.data.startswith("docs_cat:"))
async def docs_category(call: CallbackQuery):
    """–ü–æ–∫–∞–∑–∞—Ç—å –¥–æ–∫—É–º–µ–Ω—Ç—ã –≤ –∫–∞—Ç–µ–≥–æ—Ä–∏–∏."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    category = call.data.split(":")[-1]
    docs = get_docs_in_category(category)

    if not docs:
        await call.answer("–î–æ–∫—É–º–µ–Ω—Ç—ã –≤ —ç—Ç–æ–π –∫–∞—Ç–µ–≥–æ—Ä–∏–∏ –ø–æ–∫–∞ –æ—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç.", show_alert=True)
        return

    category_title = CATEGORY_TITLES.get(category, "–î–æ–∫—É–º–µ–Ω—Ç—ã")
    text = f"{category_title}\n\n–í—ã–±–µ—Ä–∏—Ç–µ –¥–æ–∫—É–º–µ–Ω—Ç –¥–ª—è –ø—Ä–æ—Å–º–æ—Ç—Ä–∞:"

    await call.message.answer(text, reply_markup=docs_category_ikb(category, docs))
    await call.answer()


@dp.callback_query(F.data.startswith("docs_item:"))
async def docs_item(call: CallbackQuery):
    """–ü–æ–∫–∞–∑–∞—Ç—å –∫–∞—Ä—Ç–æ—á–∫—É –¥–æ–∫—É–º–µ–Ω—Ç–∞."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    if len(parts) < 3:
        await call.answer("–û—à–∏–±–∫–∞ –≤ –¥–∞–Ω–Ω—ã—Ö.", show_alert=True)
        return

    category = parts[1]
    doc_id = parts[2]

    doc = get_document(category, doc_id)
    if not doc:
        await call.answer("–î–æ–∫—É–º–µ–Ω—Ç –Ω–µ –Ω–∞–π–¥–µ–Ω.", show_alert=True)
        return

    text = f"üìÑ {doc['title']}\n\n{doc['description']}"

    await call.message.answer(text, reply_markup=docs_item_ikb(category))
    await call.answer()


# ========== –°—Ç–∞—Ä—ã–µ —Ö—ç–Ω–¥–ª–µ—Ä—ã (–¥–ª—è —Å–æ–≤–º–µ—Å—Ç–∏–º–æ—Å—Ç–∏) ==========

@dp.callback_query(F.data == "profile:cases")
async def profile_cases(call: CallbackQuery):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    rows = list_cases(uid)
    if not rows:
        await call.message.answer("–ü–æ–∫–∞ –Ω–µ—Ç –¥–µ–ª.", reply_markup=profile_ikb())
        await call.answer()
        return

    await call.message.answer("üìÇ –í–∞—à–∏ –¥–µ–ª–∞:", reply_markup=cases_list_ikb(rows))
    await call.answer()


# case_open() DUPLICATE REMOVED - see line ~2700 for active implementation

@dp.callback_query(F.data.startswith("case:docs:"))
async def case_docs(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    case_id = int(call.data.split(":")[-1])

    # —Å–æ—Ö—Ä–∞–Ω–∏–º –≤—ã–±—Ä–∞–Ω–Ω–æ–µ –¥–µ–ª–æ (–Ω–∞ –±—É–¥—É—â–µ–µ)
    await state.update_data(docs_case_id=case_id)

    # –ø–æ–∫–∞–∑—ã–≤–∞–µ–º —É–∂–µ —Å–æ–∑–¥–∞–Ω–Ω—ã–µ —Ñ–∞–π–ª—ã –ø–æ –¥–µ–ª—É (–¢–û–õ–¨–ö–û –Ω–æ–≤–∞—è —Å—Ç—Ä—É–∫—Ç—É—Ä–∞)
    case_dir = GENERATED_DIR / "cases" / str(case_id)
    files = []
    if case_dir.is_dir():
        files = sorted(
            [p.name for p in case_dir.iterdir() if p.is_file() and p.suffix.lower() == ".docx"],
            reverse=True,
        )

    # –∫–ª–∞–≤–∏–∞—Ç—É—Ä–∞: –≥–µ–Ω–µ—Ä–∞—Ü–∏—è + –ø–æ—Å–ª–µ–¥–Ω–∏–π –¥–æ–∫—É–º–µ–Ω—Ç + –∞—Ä—Ö–∏–≤
    kb = InlineKeyboardBuilder()
    kb.button(text="üßæ –°—Ñ–æ—Ä–º–∏—Ä–æ–≤–∞—Ç—å –∑–∞—è–≤–ª–µ–Ω–∏–µ –æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–µ (–Ω–æ–≤–æ–µ)", callback_data=f"case:gen:{case_id}:petition")
    if files:
        latest = files[0]
        kb.button(text="üìé –ü–æ—Å–ª–µ–¥–Ω–∏–π –¥–æ–∫—É–º–µ–Ω—Ç", callback_data=f"case:lastdoc:{case_id}")
        kb.button(text="üìö –ê—Ä—Ö–∏–≤ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤", callback_data=f"case:archive:{case_id}:1")
    kb.button(text="üîô –ù–∞–∑–∞–¥ –∫ –¥–µ–ª—É", callback_data=f"case:open:{case_id}")
    kb.adjust(1)

    if not files:
        await call.message.answer(
            f"üìé –î–æ–∫—É–º–µ–Ω—Ç—ã –ø–æ –¥–µ–ª—É #{case_id} –ø–æ–∫–∞ –æ—Ç—Å—É—Ç—Å—Ç–≤—É—é—Ç.\n"
            "–ù–∞–∂–º–∏ –∫–Ω–æ–ø–∫—É –Ω–∏–∂–µ, —á—Ç–æ–±—ã —Å—Ñ–æ—Ä–º–∏—Ä–æ–≤–∞—Ç—å –Ω–æ–≤—ã–π –¥–æ–∫—É–º–µ–Ω—Ç (–æ–Ω —Å–æ—Ö—Ä–∞–Ω–∏—Ç—Å—è –≤ –∞—Ä—Ö–∏–≤).",
            reply_markup=kb.as_markup(),
        )
        if hasattr(call, "answer"):
            await call.answer()
        return

    await call.message.answer(
        f"üìé –î–æ–∫—É–º–µ–Ω—Ç—ã –ø–æ –¥–µ–ª—É #{case_id} (–ø–æ—Å–ª–µ–¥–Ω–∏–µ —Å–≤–µ—Ä—Ö—É):",
        reply_markup=kb.as_markup(),
    )
    if hasattr(call, "answer"):
        await call.answer()

@dp.callback_query(F.data.startswith("case:lastdoc:"))
async def case_lastdoc_send(call: CallbackQuery):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    case_id = int(call.data.split(":")[-1])
    case_dir = GENERATED_DIR / "cases" / str(case_id)
    if not case_dir.is_dir():
        await call.message.answer("–î–æ–∫—É–º–µ–Ω—Ç—ã –Ω–µ –Ω–∞–π–¥–µ–Ω—ã.")
        await call.answer()
        return

    files = sorted(
        [p.name for p in case_dir.iterdir() if p.is_file() and p.suffix.lower() == ".docx"],
        reverse=True,
    )
    if not files:
        await call.message.answer("–î–æ–∫—É–º–µ–Ω—Ç—ã –Ω–µ –Ω–∞–π–¥–µ–Ω—ã.")
        await call.answer()
        return

    path = case_dir / files[0]
    if not path.is_file():
        await call.message.answer("–§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω (–≤–æ–∑–º–æ–∂–Ω–æ, —É–¥–∞–ª—ë–Ω).")
        await call.answer()
        return

    await call.message.answer_document(FSInputFile(path), caption=f"–ü–æ—Å–ª–µ–¥–Ω–∏–π –¥–æ–∫—É–º–µ–Ω—Ç –ø–æ –¥–µ–ª—É #{case_id}")
    await call.answer()


@dp.callback_query(F.data.startswith("case:archive:"))
async def case_archive(call: CallbackQuery):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    if len(parts) < 4:
        await call.answer()
        return

    case_id = int(parts[2])
    try:
        page = int(parts[3])
    except ValueError:
        page = 1
    if page < 1:
        page = 1

    case_dir = GENERATED_DIR / "cases" / str(case_id)
    files_all = []
    if case_dir.is_dir():
        files_all = sorted(
            [p.name for p in case_dir.iterdir() if p.is_file() and p.suffix.lower() == ".docx"],
            reverse=True,
        )

    archive_files = files_all[1:] if len(files_all) > 1 else []
    per_page = 10
    total = len(archive_files)
    max_page = max(1, (total + per_page - 1) // per_page)
    if page > max_page:
        page = max_page

    start = (page - 1) * per_page
    end = min(start + per_page, total)
    chunk = archive_files[start:end]

    kb = InlineKeyboardBuilder()
    if not chunk:
        kb.button(text="(–∞—Ä—Ö–∏–≤ –ø—É—Å—Ç)", callback_data="noop")
    else:
        for i, name in enumerate(chunk, start=start):
            kb.button(text=f"üìé {name}", callback_data=f"case:fileidx:{case_id}:{i}")

    if page > 1:
        kb.button(text="‚¨ÖÔ∏è –ù–∞–∑–∞–¥", callback_data=f"case:archive:{case_id}:{page-1}")
    if page < max_page:
        kb.button(text="‚û°Ô∏è –î–∞–ª–µ–µ", callback_data=f"case:archive:{case_id}:{page+1}")

    kb.button(text="üîô –ù–∞–∑–∞–¥ –∫ –¥–æ–∫—É–º–µ–Ω—Ç–∞–º", callback_data=f"case:docs:{case_id}")
    kb.adjust(1)

    await call.message.answer(
        f"üìö –ê—Ä—Ö–∏–≤ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –ø–æ –¥–µ–ª—É #{case_id} (—Å—Ç—Ä. {page}/{max_page})",
        reply_markup=kb.as_markup(),
    )
    await call.answer()


@dp.callback_query(F.data.startswith("case:fileidx:"))
async def case_file_send_by_index(call: CallbackQuery):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    if len(parts) < 4:
        await call.answer()
        return

    case_id = int(parts[2])
    try:
        idx = int(parts[3])
    except ValueError:
        await call.answer()
        return

    case_dir = GENERATED_DIR / "cases" / str(case_id)
    files_all = []
    if case_dir.is_dir():
        files_all = sorted(
            [p.name for p in case_dir.iterdir() if p.is_file() and p.suffix.lower() == ".docx"],
            reverse=True,
        )

    archive_files = files_all[1:] if len(files_all) > 1 else []
    if idx < 0 or idx >= len(archive_files):
        await call.message.answer("–§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω (–≤–æ–∑–º–æ–∂–Ω–æ, –∞—Ä—Ö–∏–≤ –∏–∑–º–µ–Ω–∏–ª—Å—è). –û—Ç–∫—Ä–æ–π –∞—Ä—Ö–∏–≤ –∑–∞–Ω–æ–≤–æ.")
        await call.answer()
        return

    filename = archive_files[idx]
    path = case_dir / filename
    if not path.is_file():
        await call.message.answer("–§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω (–≤–æ–∑–º–æ–∂–Ω–æ, —É–¥–∞–ª—ë–Ω).")
        await call.answer()
        return

    await call.message.answer_document(FSInputFile(path))
    await call.answer()

# case_file_send() DUPLICATE REMOVED - see line ~2350 for active implementation

@dp.callback_query(F.data == "noop")
async def noop(call: CallbackQuery):
    await call.answer()


@dp.callback_query(F.data.startswith("case:gen:"))
async def case_generate_from_case_docs(call: CallbackQuery, state: FSMContext):
    """
    –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –Ω–æ–≤–æ–≥–æ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –ø—Ä—è–º–æ –∏–∑ "–î–æ–∫—É–º–µ–Ω—Ç—ã –ø–æ –¥–µ–ª—É"
    callback_data: case:gen:<case_id>:petition|online
    """
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    if len(parts) != 4:
        await call.message.answer("–ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω–∞—è –∫–æ–º–∞–Ω–¥–∞.")
        await call.answer()
        return

    case_id = int(parts[2])
    doc_kind = parts[3]

    case_row = get_case(uid, case_id)
    if not case_row:
        await call.message.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ.")
        await call.answer()
        return

    # —Å–æ—Ö—Ä–∞–Ω—è–µ–º –≤—ã–±—Ä–∞–Ω–Ω–æ–µ –¥–µ–ª–æ –≤ state
    await state.update_data(docs_case_id=case_id)

    if doc_kind == "petition":
        card = get_case_card(uid, case_id)
        if not card:
            await call.message.answer("–ö–∞—Ä—Ç–æ—á–∫–∞ –¥–µ–ª–∞ –µ—â—ë –Ω–µ –∑–∞–ø–æ–ª–Ω–µ–Ω–∞. –°–Ω–∞—á–∞–ª–∞ –∑–∞–ø–æ–ª–Ω–∏ –∫–∞—Ä—Ç–æ—á–∫—É –¥–µ–ª–∞.")
            await call.answer()
            return

        validation = validate_case_card(card)
        missing = validation.get("missing", []) if isinstance(validation, dict) else (validation or [])

        if missing:
            await call.message.answer(
                "–ù–µ —Ö–≤–∞—Ç–∞–µ—Ç –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã—Ö –¥–∞–Ω–Ω—ã—Ö –≤ –∫–∞—Ä—Ç–æ—á–∫–µ –¥–µ–ª–∞:\n"
                + "- " + _humanize_missing(missing).replace(", ", "\n- ")
                + "\n\n–ù–∞–∂–º–∏ ¬´–†–µ–¥–∞–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –∫–∞—Ä—Ç–æ—á–∫–∏¬ª –∏ –∑–∞–ø–æ–ª–Ω–∏ –ø–æ–ª—è –ø–æ —à–∞–≥–∞–º."
            )
            await call.answer()
            return

        path = await build_bankruptcy_petition_doc(case_row, card)
        await call.message.answer_document(
            FSInputFile(path),
            caption=f"–ì–æ—Ç–æ–≤–æ ‚úÖ –ó–∞—è–≤–ª–µ–Ω–∏–µ –æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–µ (–¥–µ–ª–æ #{case_id})",
        )

    else:
        await call.message.answer("–ù–µ–∏–∑–≤–µ—Å—Ç–Ω—ã–π —Ç–∏–ø –¥–æ–∫—É–º–µ–Ω—Ç–∞.")
        await call.answer()
        return

    # –ø–æ—Å–ª–µ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ ‚Äî —Å—Ä–∞–∑—É –ø–æ–∫–∞–∑–∞—Ç—å –æ–±–Ω–æ–≤–ª—ë–Ω–Ω—ã–π –∞—Ä—Ö–∏–≤
    fake = type("X", (), {})()
    fake.from_user = call.from_user
    fake.data = f"case:docs:{case_id}"
    fake.message = call.message

    await case_docs(fake, state)
    await call.answer()

@dp.callback_query(lambda c: c.data.startswith("case:edit:") and c.data.count(":") == 2)
async def case_edit_menu(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    case_id = int(call.data.split(":")[-1])

    await state.clear()


    

    # --- EDIT MENU SHELL (no docs, no CaseCardFill) ---

    row = get_case(uid, case_id)

    if not row:

        await call.message.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ.")

        await call.answer()

        return

    

    try:

        case_number = row[2] if len(row) > 2 else ""

        stage = row[3] if len(row) > 3 else ""

        court = row[5] if len(row) > 5 else ""

        judge = row[6] if len(row) > 6 else ""

        fin_manager = row[7] if len(row) > 7 else ""

        notes = row[8] if len(row) > 8 else ""

    except (IndexError, TypeError) as e:
        logger.warning(f"Failed to parse case row: {e}")
        case_number = stage = court = judge = fin_manager = notes = ""

    

    text = (

        f"‚úèÔ∏è –†–µ–¥–∞–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –∫–∞—Ä—Ç–æ—á–∫–∏ –¥–µ–ª–∞ #{case_id}\n\n"

        f"–ù–æ–º–µ—Ä –¥–µ–ª–∞: {case_number or '‚Äî'}\n"

        f"–°—É–¥: {court or '‚Äî'}\n"

        f"–°—É–¥—å—è: {judge or '‚Äî'}\n"

        f"–§–£: {fin_manager or '‚Äî'}\n"

        f"–°—Ç–∞–¥–∏—è: {stage or '‚Äî'}\n"

        f"–ó–∞–º–µ—Ç–∫–∏: {notes or '‚Äî'}"

    )

    

    kb = InlineKeyboardBuilder()

    kb.button(text="üìã –ü–æ–∫–∞–∑–∞—Ç—å –∫–∞—Ä—Ç–æ—á–∫—É –¥–µ–ª–∞", callback_data=f"case:card:{case_id}")

    kb.button(text="‚úèÔ∏è –ù–æ–º–µ—Ä –¥–µ–ª–∞", callback_data=f"case:edit:{case_id}:case_number")

    kb.button(text="‚úèÔ∏è –°—É–¥", callback_data=f"case:edit:{case_id}:court")

    kb.button(text="‚úèÔ∏è –°—É–¥—å—è", callback_data=f"case:edit:{case_id}:judge")

    kb.button(text="‚úèÔ∏è –§–£", callback_data=f"case:edit:{case_id}:fin_manager")

    kb.button(text="‚úèÔ∏è –°—Ç–∞–¥–∏—è", callback_data=f"case:edit:{case_id}:stage")

    kb.button(text="üóí –ó–∞–º–µ—Ç–∫–∏", callback_data=f"case:edit:{case_id}:notes")

    kb.button(text="üîô –ù–∞–∑–∞–¥ –∫ –¥–µ–ª—É", callback_data=f"case:open:{case_id}")

    kb.adjust(1, 2, 2, 2, 1)

    

    await call.message.answer(text, reply_markup=kb.as_markup())

    await call.answer()

    return

    # --- /EDIT MENU SHELL ---

    
    card = get_case_card(uid, case_id) or {}
    next_field = None
    for key, _meta in CASE_CARD_FIELDS:
        val = card.get(key)
        if val is None or (isinstance(val, str) and not val.strip()):
            next_field = key
            break

    if not next_field:
        await state.update_data(card_case_id=case_id)
        await send_card_fill_menu(call.message, uid, case_id)
        await call.answer()
        return

    await state.update_data(card_case_id=case_id, card_field_key=next_field)
    await state.set_state(CaseCardFill.waiting_value)

    filled, total = _card_completion_status(card)
    prompt = CASE_CARD_FIELD_META[next_field]["prompt"] + "\n–û—Ç–ø—Ä–∞–≤—å '-' —á—Ç–æ–±—ã –æ—Å—Ç–∞–≤–∏—Ç—å –ø—É—Å—Ç—ã–º."
    await call.message.answer(
        f"‚úçÔ∏è –ó–∞–ø–æ–ª–Ω—è–µ–º –∫–∞—Ä—Ç–æ—á–∫—É –¥–µ–ª–∞ #{case_id}. –ó–∞–ø–æ–ª–Ω–µ–Ω–æ {filled}/{total}.\n"
        f"–°–µ–π—á–∞—Å: {CASE_CARD_FIELD_META[next_field]['title']}.\n"
        f"{prompt}"
    )
    await call.answer()


@dp.callback_query(lambda c: c.data == "profile:menu")
async def profile_menu(call: CallbackQuery):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    row = get_profile(uid)

    if not row:
        text = "–ü—Ä–æ—Ñ–∏–ª—å –ø–æ–∫–∞ –Ω–µ –∑–∞–ø–æ–ª–Ω–µ–Ω.\n\n–ù–∞–∂–º–∏ ¬´‚úèÔ∏è –ó–∞–ø–æ–ª–Ω–∏—Ç—å –ø—Ä–æ—Ñ–∏–ª—å¬ª."
    else:
        _, full_name, role, address, phone, email, *_ = row
        text = (
            "üë§ –ú–æ–π –ø—Ä–æ—Ñ–∏–ª—å:\n"
            f"–§–ò–û/–û—Ä–≥: {full_name or '-'}\n"
            f"–°—Ç–∞—Ç—É—Å: {role or '-'}\n"
            f"–ê–¥—Ä–µ—Å: {address or '-'}\n"
            f"–¢–µ–ª–µ—Ñ–æ–Ω: {phone or '-'}\n"
            f"Email: {email or '-'}\n\n"
            "–ù–∞–∂–º–∏ ¬´‚úèÔ∏è –ó–∞–ø–æ–ª–Ω–∏—Ç—å –ø—Ä–æ—Ñ–∏–ª—å¬ª, —á—Ç–æ–±—ã –∏–∑–º–µ–Ω–∏—Ç—å."
        )

    kb = InlineKeyboardBuilder()
    kb.button(text="‚úèÔ∏è –ó–∞–ø–æ–ª–Ω–∏—Ç—å –ø—Ä–æ—Ñ–∏–ª—å", callback_data="profile:edit")
    kb.button(text="üîô –ù–∞–∑–∞–¥", callback_data="docs:back_menu")
    kb.adjust(1)

    await call.message.answer(text, reply_markup=kb.as_markup())
    await call.answer()
@dp.callback_query(lambda c: c.data == "profile:edit")
async def profile_edit_start(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    await state.clear()
    await state.set_state(ProfileFill.full_name)
    await call.message.answer("–í–≤–µ–¥–∏ –§–ò–û –∏–ª–∏ –Ω–∞–∑–≤–∞–Ω–∏–µ –æ—Ä–≥–∞–Ω–∏–∑–∞—Ü–∏–∏ (–∫–∞–∫ –±—É–¥–µ—Ç –≤ –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ö).")
    await call.answer()


@dp.callback_query(lambda c: c.data == "docs:choose_case")
async def docs_choose_case(call: CallbackQuery):
    """–ü–æ–∫–∞–∑—ã–≤–∞–µ—Ç —Å–ø–∏—Å–æ–∫ –¥–µ–ª –¥–ª—è –≤—ã–±–æ—Ä–∞ –¥–µ–ª–∞ –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    rows = list_cases(uid)
    if not rows:
        await call.message.answer("–ü–æ–∫–∞ –Ω–µ—Ç –¥–µ–ª. –°–æ–∑–¥–∞–π –¥–µ–ª–æ —á–µ—Ä–µ–∑ ¬´üìÇ –î–µ–ª–∞¬ª.")
        await call.answer()
        return

    kb = InlineKeyboardBuilder()
    lines = ["üìÑ –í—ã–±–µ—Ä–∏ –¥–µ–ª–æ –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤:"]
    for (cid, code_name, case_number, stage, updated_at) in rows:
        num = case_number or "-"
        lines.append(f"#{cid} | {code_name} | ‚Ññ {num}")
        kb.button(text=f"–î–µ–ª–æ #{cid}: {code_name}", callback_data=f"docs:case:{cid}")

    kb.button(text="üîô –ù–∞–∑–∞–¥", callback_data="docs:back_menu")
    kb.adjust(1)

    await call.message.answer("\n".join(lines), reply_markup=kb.as_markup())
    await call.answer()


@dp.callback_query(lambda c: c.data.startswith("docs:case:"))
async def docs_case_selected(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    cid = int(call.data.split(":")[2])
    row = get_case(uid, cid)
    if not row:
        await call.message.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ.")
        await call.answer()
        return

    await state.update_data(docs_case_id=cid)
    await call.message.answer(
        f"‚úÖ –í—ã–±—Ä–∞–Ω–æ –¥–µ–ª–æ #{cid}. –¢–µ–ø–µ—Ä—å –≤—ã–±–µ—Ä–∏ –¥–æ–∫—É–º–µ–Ω—Ç üëá",
        reply_markup=docs_menu_ikb(cid),
    )
    await call.answer()

@dp.callback_query(lambda c: c.data.startswith("docs:petition:"))
async def docs_petition(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":", 2)
    doc_key = parts[2] if len(parts) == 3 else ""

    # –ë–µ—Ä—ë–º –≤—ã–±—Ä–∞–Ω–Ω–æ–µ –¥–µ–ª–æ –∏–∑ state (–º—ã –µ–≥–æ —Å–æ—Ö—Ä–∞–Ω—è–µ–º –≤ case:docs:<id>)
    cid = await _selected_case_id(state)
    if cid is None:
        await call.message.answer("–°–Ω–∞—á–∞–ª–∞ –≤—ã–±–µ—Ä–∏ –¥–µ–ª–æ‚Ä¶")
        await docs_choose_case(call)
        await call.answer()
        return

    case_row = get_case(uid, cid)
    if not case_row:
        await state.update_data(docs_case_id=None)
        await call.message.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ. –í—ã–±–µ—Ä–∏ –µ–≥–æ –∑–∞–Ω–æ–≤–æ.")
        await docs_choose_case(call)
        await call.answer()
        return

    card = get_case_card(uid, cid)
    if not card:
        await call.message.answer(
            "–ö–∞—Ä—Ç–æ—á–∫–∞ –¥–µ–ª–∞ –µ—â—ë –Ω–µ –∑–∞–ø–æ–ª–Ω–µ–Ω–∞.\n"
            "–î–æ–±–∞–≤—å –¥–∞–Ω–Ω—ã–µ –¥–µ–ª–∞ (–ø–æ–ª, –ø–∞—Å–ø–æ—Ä—Ç, –¥–æ–ª–≥–∏ –∏ —Ç.–¥.)."
        )
        await call.answer()
        return

    validation = validate_case_card(card)
    missing = validation.get("missing", [])
    if missing:
        await call.message.answer(
            "–ù–µ —Ö–≤–∞—Ç–∞–µ—Ç –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã—Ö –¥–∞–Ω–Ω—ã—Ö –≤ –∫–∞—Ä—Ç–æ—á–∫–µ –¥–µ–ª–∞:\n"
            + "\n".join(f"- {m}" for m in missing)
        )
        await call.answer()
        return

    if doc_key != "bankruptcy_petition":
        await call.message.answer("–î–æ–∫—É–º–µ–Ω—Ç –Ω–µ –Ω–∞–π–¥–µ–Ω")
        await call.answer()
        return

    path = await build_bankruptcy_petition_doc(case_row, card)
    await call.message.answer_document(
        FSInputFile(path),
        caption=f"–ì–æ—Ç–æ–≤–æ ‚úÖ –ó–∞—è–≤–ª–µ–Ω–∏–µ –æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–µ –¥–ª—è –¥–µ–ª–∞ #{cid}",
    )
    await call.answer()

@dp.callback_query(F.data.startswith("case:file:"))
async def case_file_send(call: CallbackQuery) -> None:
    """
    Send a document file from case directory to user.

    Callback format: case:file:<case_id>:<filename>

    Args:
        call: Telegram callback query with file request

    Security: Validates filename to prevent directory traversal attacks
    """
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":", maxsplit=3)
    if len(parts) < 4:
        logger.warning(f"Invalid case:file callback format: {call.data}")
        await call.answer("–ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –∫–æ–º–∞–Ω–¥—ã")
        return

    cid_str, filename = parts[2], parts[3]

    # Security: Prevent directory traversal attacks
    if any(bad in filename for bad in ("/", "\\", "..")):
        logger.warning(f"Directory traversal attempt by user {uid}: {filename}")
        await call.message.answer("–ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ–µ –∏–º—è —Ñ–∞–π–ª–∞")
        await call.answer()
        return

    path = GENERATED_DIR / "cases" / cid_str / filename
    if not path.is_file():
        logger.info(f"File not found: {path}")
        await call.message.answer("–§–∞–π–ª –Ω–µ –Ω–∞–π–¥–µ–Ω")
        await call.answer()
        return

    try:
        await call.message.answer_document(FSInputFile(path))
        await call.answer()
    except Exception as e:
        logger.error(f"Failed to send file {path}: {e}")
        await call.message.answer("–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ—Ç–ø—Ä–∞–≤–∫–µ —Ñ–∞–π–ª–∞")
        await call.answer()


@dp.callback_query(lambda c: c.data == "docs:back_menu")
async def docs_back_menu(call: CallbackQuery, state: FSMContext):
    cid = await _selected_case_id(state)
    await call.message.answer("–î–æ–∫—É–º–µ–Ω—Ç—ã: –≤—ã–±–µ—Ä–∏ –¥–µ–π—Å—Ç–≤–∏–µ üëá", reply_markup=docs_menu_ikb(cid))
    await call.answer()

@dp.message(lambda m: m.text == "‚ÑπÔ∏è –ü–æ–º–æ—â—å")
async def help_entry(message: Message):
    await message.answer("–ü–æ–º–æ—â—å: –≤—ã–±–µ—Ä–∏ —Ä–∞–∑–¥–µ–ª –∫–Ω–æ–ø–∫–∞–º–∏. –ï—Å–ª–∏ —á—Ç–æ-—Ç–æ —Å–ª–æ–º–∞–ª–æ—Å—å ‚Äî –Ω–∞–ø–∏—à–∏ /start")
@dp.callback_query(lambda c: c.data == "back:main")
async def back_to_main(call: CallbackQuery):
    await call.message.answer("–ì–ª–∞–≤–Ω–æ–µ –º–µ–Ω—é üëá", reply_markup=main_menu_kb())
    await call.answer()

@dp.message(Command("card_set"))
async def card_set(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    cid = await _selected_case_id(state)
    if cid is None:
        await message.answer("–°–Ω–∞—á–∞–ª–∞ –≤—ã–±–µ—Ä–∏ –¥–µ–ª–æ —á–µ—Ä–µ–∑ ¬´üìÇ –î–µ–ª–∞¬ª, –∑–∞—Ç–µ–º –ø–æ–≤—Ç–æ—Ä–∏ /card_set –∏ –æ—Ç–ø—Ä–∞–≤—å JSON.")
        return

    text = (message.text or "").strip()
    parts = text.split(maxsplit=1)
    if len(parts) < 2:
        await message.answer(
            "–ü—Ä–∏—à–ª–∏ –∫–æ–º–∞–Ω–¥—É —Ç–∞–∫:\n"
            "/card_set {JSON}\n\n"
            "–ü—Ä–∏–º–µ—Ä:\n"
            "/card_set {\"debtor_gender\":\"male\"}"
        )
        return

    raw_json = parts[1].strip()
    try:
        data = json.loads(raw_json)
        if not isinstance(data, dict):
            raise ValueError("JSON –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –æ–±—ä–µ–∫—Ç–æ–º (—Å–ª–æ–≤–∞—Ä—ë–º)")
    except Exception as e:
        await message.answer(f"–û—à–∏–±–∫–∞ JSON: {e}\n\n–ü—Ä–æ–≤–µ—Ä—å –∫–∞–≤—ã—á–∫–∏ –∏ –∑–∞–ø—è—Ç—ã–µ –∏ –ø—Ä–∏—à–ª–∏ —Å–Ω–æ–≤–∞.")
        return

    # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∫–∞—Ä—Ç–æ—á–∫—É
    upsert_case_card(uid, cid, data)

    validation = validate_case_card(data)
    missing = validation.get("missing", [])
    if missing:
        await message.answer(
            "–ö–∞—Ä—Ç–æ—á–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∞ ‚úÖ\n"
            "–ù–æ –ø–æ–∫–∞ –Ω–µ —Ö–≤–∞—Ç–∞–µ—Ç –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã—Ö –ø–æ–ª–µ–π:\n"
            + "\n".join(f"- {m}" for m in missing)
        )
        return

    await message.answer(
        "–ö–∞—Ä—Ç–æ—á–∫–∞ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∞ ‚úÖ\n"
        "–í—Å–µ –æ–±—è–∑–∞—Ç–µ–ª—å–Ω—ã–µ –ø–æ–ª—è –∑–∞–ø–æ–ª–Ω–µ–Ω—ã. –¢–µ–ø–µ—Ä—å –º–æ–∂–Ω–æ –Ω–∞–∂–∞—Ç—å ¬´üìÑ –ó–∞—è–≤–ª–µ–Ω–∏–µ –æ –±–∞–Ω–∫—Ä–æ—Ç—Å—Ç–≤–µ¬ª."
    )


@dp.message(Command("doc_test"))
async def doc_test(message: Message):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    rows = list_cases(uid)
    if not rows:
        await message.answer("–ù–µ—Ç –¥–µ–ª. –°–Ω–∞—á–∞–ª–∞ —Å–æ–∑–¥–∞–π –¥–µ–ª–æ –≤ ¬´üìÇ –î–µ–ª–∞¬ª.")
        return

    # –≤–æ–∑—å–º—ë–º —Å–∞–º–æ–µ —Å–≤–µ–∂–µ–µ –¥–µ–ª–æ
    cid = rows[0][0]
    case_row = get_case(uid, cid)
    if not case_row:
        await message.answer("–ù–µ –Ω–∞—à—ë–ª –¥–µ–ª–æ –¥–ª—è —Ç–µ—Å—Ç–∞.")
        return

@dp.callback_query(lambda c: c.data == "case:new")
async def case_new(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    await state.clear()
    await state.set_state(CaseCreate.code_name)
    await call.message.answer("–í–≤–µ–¥–∏ –∫–æ–¥–æ–≤–æ–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –¥–µ–ª–∞ (–Ω–∞–ø—Ä–∏–º–µ—Ä: –ò–í–ê–ù–û–í_2025).")
    await call.answer()

@dp.message(StateFilter(CaseCreate.code_name))
async def case_step_code_name(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –í–≤–µ–¥–∏ –∫–æ–¥–æ–≤–æ–µ –Ω–∞–∑–≤–∞–Ω–∏–µ –¥–µ–ª–∞.")
        return

    await state.update_data(code_name=text)
    await state.set_state(CaseCreate.case_number)
    await message.answer("–¢–µ–ø–µ—Ä—å –≤–≤–µ–¥–∏ –Ω–æ–º–µ—Ä –¥–µ–ª–∞ (–º–æ–∂–Ω–æ '-' –µ—Å–ª–∏ –ø–æ–∫–∞ –Ω–µ—Ç).")

@dp.message(StateFilter(CaseCreate.case_number))
async def case_step_case_number(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –í–≤–µ–¥–∏ –Ω–æ–º–µ—Ä –¥–µ–ª–∞ –∏–ª–∏ '-'.")
        return

    await state.update_data(case_number=None if text == "-" else text)
    await state.set_state(CaseCreate.court)
    await message.answer("–£–∫–∞–∂–∏ —Å—É–¥ (–Ω–∞–ø—Ä–∏–º–µ—Ä: –ê–° –≥. –ú–æ—Å–∫–≤—ã) –∏–ª–∏ '-'.")

@dp.message(StateFilter(ProfileFill.full_name))
async def profile_step_full_name(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –í–≤–µ–¥–∏ –§–ò–û/–æ—Ä–≥–∞–Ω–∏–∑–∞—Ü–∏—é.")
        return

    await state.update_data(full_name=text)
    await state.set_state(ProfileFill.role)
    await message.answer("–°—Ç–∞—Ç—É—Å –≤ –¥–µ–ª–µ (–Ω–∞–ø—Ä–∏–º–µ—Ä: –¥–æ–ª–∂–Ω–∏–∫ / –ø—Ä–µ–¥—Å—Ç–∞–≤–∏—Ç–µ–ª—å / –∫—Ä–µ–¥–∏—Ç–æ—Ä).")

@dp.message(StateFilter(ProfileFill.role))
async def profile_step_role(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –í–≤–µ–¥–∏ —Å—Ç–∞—Ç—É—Å (–¥–æ–ª–∂–Ω–∏–∫/–ø—Ä–µ–¥—Å—Ç–∞–≤–∏—Ç–µ–ª—å/–∫—Ä–µ–¥–∏—Ç–æ—Ä).")
        return

    await state.update_data(role=text)
    await state.set_state(ProfileFill.address)
    await message.answer("–ê–¥—Ä–µ—Å (–¥–ª—è —à–∞–ø–∫–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞). –ú–æ–∂–Ω–æ '-' –µ—Å–ª–∏ –Ω–µ –Ω—É–∂–Ω–æ.")

@dp.message(StateFilter(ProfileFill.address))
async def profile_step_address(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –í–≤–µ–¥–∏ –∞–¥—Ä–µ—Å –∏–ª–∏ '-'.")
        return

    await state.update_data(address=None if text == "-" else text)
    await state.set_state(ProfileFill.phone)
    await message.answer("–¢–µ–ª–µ—Ñ–æ–Ω. –ú–æ–∂–Ω–æ '-' –µ—Å–ª–∏ –Ω–µ –Ω—É–∂–Ω–æ.")

@dp.message(StateFilter(ProfileFill.phone))
async def profile_step_phone(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –í–≤–µ–¥–∏ —Ç–µ–ª–µ—Ñ–æ–Ω –∏–ª–∏ '-'.")
        return

    await state.update_data(phone=None if text == "-" else text)
    await state.set_state(ProfileFill.email)
    await message.answer("Email. –ú–æ–∂–Ω–æ '-' –µ—Å–ª–∏ –Ω–µ –Ω—É–∂–Ω–æ.")
@dp.message(StateFilter(ProfileFill.email))
async def profile_step_email(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –í–≤–µ–¥–∏ email –∏–ª–∏ '-'.")
        return

    data = await state.get_data()

    upsert_profile(
        uid,
        full_name=data.get("full_name"),
        role=data.get("role"),
        address=data.get("address"),
        phone=data.get("phone"),
        email=None if text == "-" else text,
    )

    await state.clear()

    await message.answer(
        "‚úÖ –ü—Ä–æ—Ñ–∏–ª—å —Å–æ—Ö—Ä–∞–Ω—ë–Ω.\n\n"
        "–¢–µ–ø–µ—Ä—å —ç—Ç–∏ –¥–∞–Ω–Ω—ã–µ –±—É–¥—É—Ç –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –ø–æ–¥—Å—Ç–∞–≤–ª—è—Ç—å—Å—è –≤ –¥–æ–∫—É–º–µ–Ω—Ç—ã.\n"
        "–ú–æ–∂–µ—à—å –æ—Ç–∫—Ä—ã—Ç—å ¬´üë§ –ú–æ–π –ø—Ä–æ—Ñ–∏–ª—å¬ª, —á—Ç–æ–±—ã –ø—Ä–æ–≤–µ—Ä–∏—Ç—å."
    )

@dp.message(StateFilter(CaseCreate.court))
async def case_step_court(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –£–∫–∞–∂–∏ —Å—É–¥ –∏–ª–∏ '-'.")
        return

    await state.update_data(court=None if text == "-" else text)
    await state.set_state(CaseCreate.judge)
    await message.answer("–£–∫–∞–∂–∏ —Å—É–¥—å—é (–§–ò–û) –∏–ª–∏ '-'.")


@dp.message(StateFilter(CaseCreate.judge))
async def case_step_judge(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –£–∫–∞–∂–∏ —Å—É–¥—å—é –∏–ª–∏ '-'.")
        return

    await state.update_data(judge=None if text == "-" else text)
    await state.set_state(CaseCreate.fin_manager)
    await message.answer("–£–∫–∞–∂–∏ —Ñ–∏–Ω–∞–Ω—Å–æ–≤–æ–≥–æ —É–ø—Ä–∞–≤–ª—è—é—â–µ–≥–æ –∏–ª–∏ '-'.")


@dp.message(StateFilter(CaseCreate.fin_manager))
async def case_step_fin_manager(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –£–∫–∞–∂–∏ –§–£ –∏–ª–∏ '-'.")
        return

    await state.update_data(fin_manager=None if text == "-" else text)
    data = await state.get_data()

    code_name = data.get("code_name")
    case_number = data.get("case_number")
    court = data.get("court")
    judge = data.get("judge")
    fin_manager = data.get("fin_manager")

    # —Å–æ–∑–¥–∞—ë–º –¥–µ–ª–æ –∏ –∑–∞–ø–æ–ª–Ω—è–µ–º –ø–æ–ª—è
    cid = create_case(uid, code_name)
    update_case_fields(uid, cid, case_number=case_number, court=court, judge=judge, fin_manager=fin_manager)

    await state.clear()

    await message.answer(
        "‚úÖ –î–µ–ª–æ —Å–æ–∑–¥–∞–Ω–æ.\n"
        f"ID: {cid}\n"
        f"–ö–æ–¥: {code_name}\n"
        f"–ù–æ–º–µ—Ä: {case_number or '-'}\n"
        f"–°—É–¥: {court or '-'}\n"
        f"–°—É–¥—å—è: {judge or '-'}\n"
        f"–§–£: {fin_manager or '-'}"
    )

@dp.callback_query(lambda c: c.data == "case:list")
async def case_list(call: CallbackQuery):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    rows = list_cases(uid)  # –±–µ—Ä—ë–º –ø–æ—Å–ª–µ–¥–Ω–∏–µ 20 –¥–µ–ª
    if not rows:
        await call.message.answer("–ü–æ–∫–∞ –Ω–µ—Ç –¥–µ–ª. –ù–∞–∂–º–∏ ¬´‚ûï –°–æ–∑–¥–∞—Ç—å –¥–µ–ª–æ¬ª.")
        await call.answer()
        return

    kb = InlineKeyboardBuilder()
    lines = ["üìÑ –í–∞—à–∏ –¥–µ–ª–∞ (–ø–æ—Å–ª–µ–¥–Ω–∏–µ 20):"]

    for (cid, code_name, case_number, stage, updated_at) in rows:
        num = case_number or "-"
        st = stage or "-"
        lines.append(f"#{cid} | {code_name} | ‚Ññ {num} | —Å—Ç–∞–¥–∏—è: {st}")
        kb.button(text=f"–û—Ç–∫—Ä—ã—Ç—å #{cid}", callback_data=f"case:open:{cid}")
        kb.button(text="üóÇ –ó–∞–ø–æ–ª–Ω–∏—Ç—å –∫–∞—Ä—Ç–æ—á–∫—É –¥–µ–ª–∞", callback_data = f"case:card:{cid}")

    kb.button(text="üîô –ù–∞–∑–∞–¥", callback_data="back:cases")
    kb.adjust(1)

    await call.message.answer("\n".join(lines), reply_markup=kb.as_markup())
    await call.answer()

@dp.callback_query(lambda c: c.data == "back:cases")
async def back_to_cases(call: CallbackQuery):
    await call.message.answer(
        "–†–∞–∑–¥–µ–ª ¬´–î–µ–ª–∞¬ª. –í—ã–±–µ—Ä–∏ –¥–µ–π—Å—Ç–≤–∏–µ:",
        reply_markup=cases_menu_ikb()
    )
    await call.answer()

@dp.callback_query(F.data.startswith("case:open:"))
async def case_open(call: CallbackQuery) -> None:
    """
    Display full case card with all details and editing options.

    Callback format: case:open:<case_id>
    Shows: code_name, case_number, court, judge, manager, stage, notes, timestamps

    Args:
        call: Callback query from cases list
    """
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    try:
        cid = int(call.data.split(":")[2])
    except (ValueError, IndexError) as e:
        logger.error(f"Invalid callback data format: {call.data}, error: {e}")
        await call.answer("–û—à–∏–±–∫–∞ —Ñ–æ—Ä–º–∞—Ç–∞ –¥–∞–Ω–Ω—ã—Ö")
        return

    row = get_case(uid, cid)
    if not row:
        logger.info(f"Case {cid} not found or access denied for user {uid}")
        await call.message.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ –∏–ª–∏ —É –≤–∞—Å –Ω–µ—Ç –¥–æ—Å—Ç—É–ø–∞.")
        await call.answer()
        return

    (cid, _owner_user_id, code_name, case_number, court, judge,
     fin_manager, stage, notes, created_at, updated_at) = row

    text = (
        f"üìå –î–µ–ª–æ #{cid}\n"
        f"üìã –ö–æ–¥: {code_name}\n"
        f"üìÑ –ù–æ–º–µ—Ä: {case_number or '‚Äî'}\n"
        f"üèõ –°—É–¥: {court or '‚Äî'}\n"
        f"‚öñÔ∏è –°—É–¥—å—è: {judge or '‚Äî'}\n"
        f"üë§ –§–∏–Ω. —É–ø—Ä–∞–≤–ª—è—é—â–∏–π: {fin_manager or '‚Äî'}\n"
        f"üìä –°—Ç–∞–¥–∏—è: {stage or '‚Äî'}\n"
        f"üìù –ó–∞–º–µ—Ç–∫–∏: {notes or '‚Äî'}\n"
        f"üìÖ –°–æ–∑–¥–∞–Ω–æ: {created_at}\n"
        f"üîÑ –û–±–Ω–æ–≤–ª–µ–Ω–æ: {updated_at}"
    )

    kb = InlineKeyboardBuilder()
    kb.button(text="üìÅ –ö–∞—Ä—Ç–æ—á–∫–∞ –¥–µ–ª–∞", callback_data=f"case:card:{cid}")
    kb.button(text="‚úèÔ∏è –ù–æ–º–µ—Ä –¥–µ–ª–∞", callback_data=f"case:edit:{cid}:case_number")
    kb.button(text="‚úèÔ∏è –°—É–¥", callback_data=f"case:edit:{cid}:court")
    kb.button(text="‚úèÔ∏è –°—É–¥—å—è", callback_data=f"case:edit:{cid}:judge")
    kb.button(text="‚úèÔ∏è –§–£", callback_data=f"case:edit:{cid}:fin_manager")
    kb.button(text="‚úèÔ∏è –°—Ç–∞–¥–∏—è", callback_data=f"case:edit:{cid}:stage")
    kb.button(text="üóí –ó–∞–º–µ—Ç–∫–∏", callback_data=f"case:edit:{cid}:notes")
    kb.button(text="üîô –ö —Å–ø–∏—Å–∫—É –¥–µ–ª", callback_data="case:list")
    kb.adjust(1, 2, 2, 2)

    await call.message.answer(text, reply_markup=kb.as_markup())
    await call.answer()


@dp.callback_query(lambda c: c.data.startswith("case:card:"))
async def case_card_open(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    cid = int(call.data.split(":")[2])
    await state.update_data(card_case_id=cid)
    card = get_case_card(uid, cid) or {}

    lines = [f"üìÅ –ö–∞—Ä—Ç–æ—á–∫–∞ –¥–µ–ª–∞ #{cid}"]
    for key, title in [
        ("court_name", "–°—É–¥"),
        ("court_address", "–ê–¥—Ä–µ—Å —Å—É–¥–∞"),
        ("debtor_full_name", "–î–æ–ª–∂–Ω–∏–∫"),
        ("debtor_gender", "–ü–æ–ª"),
        ("debtor_birth_date", "–î–∞—Ç–∞ —Ä–æ–∂–¥–µ–Ω–∏—è"),
        ("debtor_address", "–ê–¥—Ä–µ—Å –¥–æ–ª–∂–Ω–∏–∫–∞"),
        ("passport_series", "–ü–∞—Å–ø–æ—Ä—Ç —Å–µ—Ä–∏—è"),
        ("passport_number", "–ü–∞—Å–ø–æ—Ä—Ç –Ω–æ–º–µ—Ä"),
        ("passport_issued_by", "–ö–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç"),
        ("passport_date", "–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞"),
        ("passport_code", "–ö–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è"),
        ("total_debt_rubles", "–°—É–º–º–∞ –¥–æ–ª–≥–∞ (—Ä—É–±–ª–∏)"),
        ("total_debt_kopeks", "–°—É–º–º–∞ –¥–æ–ª–≥–∞ (–∫–æ–ø–µ–π–∫–∏)"),
    ]:
        lines.append(f"{title}: {card.get(key) or '‚Äî'}")

    kb = InlineKeyboardBuilder()
    kb.button(text="‚úèÔ∏è –ó–∞–ø–æ–ª–Ω–∏—Ç—å", callback_data=f"card:fill:{cid}")
    kb.button(text="üîô –ù–∞–∑–∞–¥", callback_data=f"case:open:{cid}")
    kb.adjust(1)

    await call.message.answer("\n".join(lines), reply_markup=kb.as_markup())
    await call.answer()

CASE_CARD_FIELDS = [
    (
        "court_name",
        {
            "title": "–ù–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ —Å—É–¥–∞",
            "prompt": "–£–∫–∞–∂–∏ –Ω–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ —Å—É–¥–∞.",
        },
    ),
    (
        "court_address",
        {
            "title": "–ê–¥—Ä–µ—Å —Å—É–¥–∞",
            "prompt": "–£–∫–∞–∂–∏ –∞–¥—Ä–µ—Å —Å—É–¥–∞.",
        },
    ),
    (
        "debtor_last_name",
        {
            "title": "–§–∞–º–∏–ª–∏—è –¥–æ–ª–∂–Ω–∏–∫–∞",
            "prompt": "–£–∫–∞–∂–∏ —Ñ–∞–º–∏–ª–∏—é –¥–æ–ª–∂–Ω–∏–∫–∞.",
        },
    ),
    (
        "debtor_first_name",
        {
            "title": "–ò–º—è –¥–æ–ª–∂–Ω–∏–∫–∞",
            "prompt": "–£–∫–∞–∂–∏ –∏–º—è –¥–æ–ª–∂–Ω–∏–∫–∞.",
        },
    ),
    (
        "debtor_middle_name",
        {
            "title": "–û—Ç—á–µ—Å—Ç–≤–æ –¥–æ–ª–∂–Ω–∏–∫–∞",
            "prompt": "–£–∫–∞–∂–∏ –æ—Ç—á–µ—Å—Ç–≤–æ –¥–æ–ª–∂–Ω–∏–∫–∞ –∏–ª–∏ '-' –µ—Å–ª–∏ –Ω–µ—Ç.",
        },
    ),
    (
        "debtor_gender",
        {
            "title": "–ü–æ–ª –¥–æ–ª–∂–Ω–∏–∫–∞",
            "prompt": "–£–∫–∞–∂–∏ –ø–æ–ª –¥–æ–ª–∂–Ω–∏–∫–∞: –º/–∂.",
        },
    ),
    (
        "debtor_birth_date",
        {
            "title": "–î–∞—Ç–∞ —Ä–æ–∂–¥–µ–Ω–∏—è",
            "prompt": "–£–∫–∞–∂–∏ –¥–∞—Ç—É —Ä–æ–∂–¥–µ–Ω–∏—è: –î–î.–ú–ú.–ì–ì–ì–ì.",
        },
    ),
    (
        "debtor_address",
        {
            "title": "–ê–¥—Ä–µ—Å –¥–æ–ª–∂–Ω–∏–∫–∞",
            "prompt": "–£–∫–∞–∂–∏ –∞–¥—Ä–µ—Å –¥–æ–ª–∂–Ω–∏–∫–∞.",
        },
    ),
    (
        "debtor_phone",
        {
            "title": "–¢–µ–ª–µ—Ñ–æ–Ω –¥–æ–ª–∂–Ω–∏–∫–∞",
            "prompt": "–£–∫–∞–∂–∏ —Ç–µ–ª–µ—Ñ–æ–Ω –¥–æ–ª–∂–Ω–∏–∫–∞ (–º–æ–∂–Ω–æ –≤ —Ñ–æ—Ä–º–∞—Ç–µ +7...) –∏–ª–∏ '-' –µ—Å–ª–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç.",
        },
    ),
    (
        "debtor_inn",
        {
            "title": "–ò–ù–ù –¥–æ–ª–∂–Ω–∏–∫–∞",
            "prompt": "–£–∫–∞–∂–∏ –ò–ù–ù –∏–ª–∏ '-' –µ—Å–ª–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç.",
        },
    ),
    (
        "debtor_snils",
        {
            "title": "–°–ù–ò–õ–° –¥–æ–ª–∂–Ω–∏–∫–∞",
            "prompt": "–£–∫–∞–∂–∏ –°–ù–ò–õ–° –∏–ª–∏ '-' –µ—Å–ª–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç.",
        },
    ),
    (
        "passport_series",
        {
            "title": "–ü–∞—Å–ø–æ—Ä—Ç —Å–µ—Ä–∏—è",
            "prompt": "–£–∫–∞–∂–∏ —Å–µ—Ä–∏—é –ø–∞—Å–ø–æ—Ä—Ç–∞ (4 —Ü–∏—Ñ—Ä—ã) –∏–ª–∏ '-' –µ—Å–ª–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç.",
        },
    ),
    (
        "passport_number",
        {
            "title": "–ü–∞—Å–ø–æ—Ä—Ç –Ω–æ–º–µ—Ä",
            "prompt": "–£–∫–∞–∂–∏ –Ω–æ–º–µ—Ä –ø–∞—Å–ø–æ—Ä—Ç–∞ (6 —Ü–∏—Ñ—Ä) –∏–ª–∏ '-' –µ—Å–ª–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç.",
        },
    ),
    (
        "passport_issued_by",
        {
            "title": "–ö–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç",
            "prompt": "–£–∫–∞–∂–∏ –∫–µ–º –≤—ã–¥–∞–Ω –ø–∞—Å–ø–æ—Ä—Ç –∏–ª–∏ '-' –µ—Å–ª–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç.",
        },
    ),
    (
        "passport_date",
        {
            "title": "–î–∞—Ç–∞ –≤—ã–¥–∞—á–∏ –ø–∞—Å–ø–æ—Ä—Ç–∞",
            "prompt": "–£–∫–∞–∂–∏ –¥–∞—Ç—É –≤—ã–¥–∞—á–∏: –î–î.–ú–ú.–ì–ì–ì–ì –∏–ª–∏ '-' –µ—Å–ª–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç.",
        },
    ),
    (
        "passport_code",
        {
            "title": "–ö–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è",
            "prompt": "–£–∫–∞–∂–∏ –∫–æ–¥ –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏—è (XXX-XXX) –∏–ª–∏ '-' –µ—Å–ª–∏ –æ—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç.",
        },
    ),
    (
        "marital_status",
        {
            "title": "–°–µ–º–µ–π–Ω–æ–µ –ø–æ–ª–æ–∂–µ–Ω–∏–µ",
            "prompt": "–£–∫–∞–∂–∏ —Å–µ–º–µ–π–Ω–æ–µ –ø–æ–ª–æ–∂–µ–Ω–∏–µ (–∂–µ–Ω–∞—Ç/–∑–∞–º—É–∂–µ–º/–Ω–µ —Å–æ—Å—Ç–æ–∏—Ç –∏ —Ç.–ø.) –∏–ª–∏ '-' –µ—Å–ª–∏ –Ω–µ–∏–∑–≤–µ—Å—Ç–Ω–æ.",
        },
    ),
    (
        "certificate_number",
        {
            "title": "–°–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–æ (–Ω–æ–º–µ—Ä)",
            "prompt": "–£–∫–∞–∂–∏ –Ω–æ–º–µ—Ä —Å–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–∞ (–æ –±—Ä–∞–∫–µ/—Ä–∞–∑–≤–æ–¥–µ) –∏–ª–∏ '-' –µ—Å–ª–∏ –Ω–µ –ø—Ä–∏–º–µ–Ω–∏–º–æ.",
        },
    ),
    (
        "certificate_date",
        {
            "title": "–°–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–æ (–¥–∞—Ç–∞)",
            "prompt": "–£–∫–∞–∂–∏ –¥–∞—Ç—É —Å–≤–∏–¥–µ—Ç–µ–ª—å—Å—Ç–≤–∞: –î–î.–ú–ú.–ì–ì–ì–ì –∏–ª–∏ '-' –µ—Å–ª–∏ –Ω–µ –ø—Ä–∏–º–µ–Ω–∏–º–æ.",
        },
    ),
    (
        "total_debt_rubles",
        {
            "title": "–°—É–º–º–∞ –¥–æ–ª–≥–∞ (—Ä—É–±–ª–∏)",
            "prompt": "–£–∫–∞–∂–∏ —Å—É–º–º—É –¥–æ–ª–≥–∞ –≤ —Ä—É–±–ª—è—Ö (—Ü–µ–ª–æ–µ —á–∏—Å–ª–æ).",
        },
    ),
    (
        "total_debt_kopeks",
        {
            "title": "–°—É–º–º–∞ –¥–æ–ª–≥–∞ (–∫–æ–ø–µ–π–∫–∏)",
            "prompt": "–£–∫–∞–∂–∏ –∫–æ–ø–µ–π–∫–∏ (0-99).",
        },
    ),

    # –í–ê–ñ–ù–û: —ç—Ç–æ –Ω–µ –æ–±—ã—á–Ω–æ–µ –ø–æ–ª–µ –≤–≤–æ–¥–∞, –∞ –æ—Ç–¥–µ–ª—å–Ω–æ–µ –º–µ–Ω—é –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤.
    # –ú—ã –±—É–¥–µ–º –ø–µ—Ä–µ—Ö–≤–∞—Ç—ã–≤–∞—Ç—å —ç—Ç–æ—Ç key –≤ –æ–±—Ä–∞–±–æ—Ç—á–∏–∫–µ –∫–ª–∏–∫–∞ –ø–æ –ø–æ–ª—é.
    (
        "creditors",
        {
            "title": "üè¶ –ö—Ä–µ–¥–∏—Ç–æ—Ä—ã (—Å–ø–∏—Å–æ–∫)",
            "prompt": "–û—Ç–∫—Ä—ã–≤–∞—é –º–µ–Ω—é –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤‚Ä¶",
        },
    ),
]

CASE_CARD_FIELD_META = {k: v for k, v in CASE_CARD_FIELDS}


def _format_case_card(card: dict[str, Any]) -> list[str]:
    lines = []
    for key, meta in CASE_CARD_FIELDS:
        val = card.get(key)
        show_val = "‚Äî"
        if val is None or str(val).strip() == "":
            show_val = "‚Äî"
        elif isinstance(val, (int, float)):
            show_val = str(val)
        else:
            show_val = str(val)
        lines.append(f"{meta['title']}: {show_val}")
    return lines


def _humanize_missing(missing: list[str]) -> str:
    titles = [CASE_CARD_FIELD_META.get(key, {}).get("title", key) for key in missing]
    return ", ".join(titles)


def _card_completion_status(card: dict[str, Any]) -> tuple[int, int]:
    validation = validate_case_card(card)
    missing = validation.get("missing") or []
    total = len(CASE_CARD_REQUIRED_FIELDS)
    return total - len(missing), total


async def send_card_fill_menu(message_target, uid: int, cid: int) -> None:
    row = get_case(uid, cid)
    if not row:
        await message_target.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ.")
        return

    _, _owner_user_id, code_name, *_ = row
    card = get_case_card(uid, cid)
    validation = validate_case_card(card)

    filled, total = _card_completion_status(card)
    text_lines = ["üìÅ –ö–∞—Ä—Ç–æ—á–∫–∞ –¥–µ–ª–∞", f"–î–µ–ª–æ #{cid} | {code_name}"]
    text_lines.append("")
    text_lines.extend(_format_case_card(card))
    text_lines.append("")
    text_lines.append(f"–ó–∞–ø–æ–ª–Ω–µ–Ω–æ: {filled}/{total}")

    if validation.get("missing"):
        text_lines.append("–ù–µ –∑–∞–ø–æ–ª–Ω–µ–Ω–æ: " + _humanize_missing(validation["missing"]))
    else:
        text_lines.append("–ö–∞—Ä—Ç–æ—á–∫–∞ –∑–∞–ø–æ–ª–Ω–µ–Ω–∞ ‚úÖ")

    kb = InlineKeyboardBuilder()
    for key, meta in CASE_CARD_FIELDS:
        kb.button(text=f"‚úèÔ∏è {meta['title']}", callback_data=f"case:cardfield:{cid}:{key}")

    # –æ—Ç–¥–µ–ª—å–Ω—ã–π —Ä–∞–∑–¥–µ–ª –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤
    creditors_count = 0
    try:
        creditors_val = card.get("creditors")
        if isinstance(creditors_val, list):
            creditors_count = len(creditors_val)
    except (TypeError, AttributeError):
        creditors_count = 0

    kb.button(text=f"üë• –ö—Ä–µ–¥–∏—Ç–æ—Ä—ã ({creditors_count})", callback_data=f"case:creditors:{cid}")
    kb.button(text="üîô –ù–∞–∑–∞–¥ –∫ –¥–µ–ª–∞–º", callback_data="case:list")
    kb.adjust(1)

    await message_target.answer("\n".join(text_lines), reply_markup=kb.as_markup())


async def send_case_card_menu(message_target, uid: int, cid: int) -> None:
    row = get_case(uid, cid)
    if not row:
        await message_target.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ.")
        return

    _, _owner_user_id, code_name, *_ = row
    card = get_case_card(uid, cid)
    validation = validate_case_card(card)

    text_lines = ["üìÅ –ö–∞—Ä—Ç–æ—á–∫–∞ –¥–µ–ª–∞", f"–î–µ–ª–æ #{cid} | {code_name}"]
    text_lines.append("")
    text_lines.extend(_format_case_card(card))

    if validation.get("missing"):
        text_lines.append("")
        text_lines.append("–ù–µ –∑–∞–ø–æ–ª–Ω–µ–Ω–æ: " + _humanize_missing(validation["missing"]))
    else:
        text_lines.append("")
        text_lines.append("–ö–∞—Ä—Ç–æ—á–∫–∞ –∑–∞–ø–æ–ª–Ω–µ–Ω–∞ ‚úÖ")

    kb = InlineKeyboardBuilder()
    for key, meta in CASE_CARD_FIELDS:
        kb.button(text=f"‚úèÔ∏è {meta['title']}", callback_data=f"case:card_edit:{cid}:{key}")
    kb.button(text="üîô –ì–æ—Ç–æ–≤–æ", callback_data=f"case:open:{cid}")
    kb.adjust(1)

    await message_target.answer("\n".join(text_lines), reply_markup=kb.as_markup())


@dp.callback_query(lambda c: c.data.startswith("case:card:"))
async def case_card_menu(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    cid = int(call.data.split(":")[2])
    await state.clear()
    await state.update_data(card_case_id=cid)
    await send_card_fill_menu(call.message, uid, cid)
    await call.answer()

@dp.callback_query(lambda c: c.data.startswith("case:card_edit:"))
async def case_card_edit(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    _, _, cid_str, field = call.data.split(":", maxsplit=3)
    cid = int(cid_str)

    if field not in CASE_CARD_FIELD_META:
        await call.answer()
        return

    row = get_case(uid, cid)
    if not row:
        await call.message.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ.")
        await call.answer()
        return

    # ‚úÖ –í–ê–ñ–ù–û: creditors ‚Äî —ç—Ç–æ –ù–ï —Ç–µ–∫—Å—Ç–æ–≤–æ–µ –ø–æ–ª–µ, –∞ –æ—Ç–¥–µ–ª—å–Ω–æ–µ –º–µ–Ω—é
    if field == "creditors":
        await state.clear()
        await state.update_data(card_case_id=cid)
        await send_creditors_menu(call.message, uid, cid)
        await call.answer()
        return

    await state.clear()
    await state.update_data(card_cid=cid, card_field=field)
    await state.set_state(CaseCardFill.waiting_value)

    prompt = CASE_CARD_FIELD_META[field]["prompt"] + "\n–û—Ç–ø—Ä–∞–≤—å '-' —á—Ç–æ–±—ã –æ—Å—Ç–∞–≤–∏—Ç—å –ø—É—Å—Ç—ã–º."
    await call.message.answer(prompt)
    await call.answer()

@dp.callback_query(lambda c: c.data.startswith("card:fill:"))
async def card_fill_start(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    _, _, cid_str = call.data.split(":", maxsplit=2)
    cid = int(cid_str)

    await state.clear()

    # –ë–µ—Ä—ë–º —Ç–µ–∫—É—â—É—é –∫–∞—Ä—Ç–æ—á–∫—É –∏ –Ω–∞—Ö–æ–¥–∏–º –ø–µ—Ä–≤–æ–µ –Ω–µ–∑–∞–ø–æ–ª–Ω–µ–Ω–Ω–æ–µ –ø–æ–ª–µ
    card = get_case_card(uid, cid) or {}
    next_field = None
    for key, _meta in CASE_CARD_FIELDS:
        val = card.get(key)
        if val is None or (isinstance(val, str) and not val.strip()):
            next_field = key
            break

    # –ï—Å–ª–∏ –≤—Å—ë –∑–∞–ø–æ–ª–Ω–µ–Ω–æ ‚Äî –ø—Ä–æ—Å—Ç–æ –ø–æ–∫–∞–∂–µ–º –º–µ–Ω—é –∫–∞—Ä—Ç–æ—á–∫–∏
    if not next_field:
        await state.update_data(card_case_id=cid)
        await send_card_fill_menu(call.message, uid, cid)
        await call.answer()
        return

    # –ò–Ω–∞—á–µ ‚Äî —Å—Ä–∞–∑—É —Å—Ç–∞—Ä—Ç—É–µ–º –≤–≤–æ–¥ –ø–µ—Ä–≤–æ–≥–æ –Ω–µ–∑–∞–ø–æ–ª–Ω–µ–Ω–Ω–æ–≥–æ –ø–æ–ª—è
    await state.update_data(card_case_id=cid, card_field_key=next_field)
    await state.set_state(CaseCardFill.waiting_value)

    filled, total = _card_completion_status(card)
    prompt = CASE_CARD_FIELD_META[next_field]["prompt"] + "\n–û—Ç–ø—Ä–∞–≤—å '-' —á—Ç–æ–±—ã –æ—Å—Ç–∞–≤–∏—Ç—å –ø—É—Å—Ç—ã–º."
    await call.message.answer(
        f"‚úçÔ∏è –ó–∞–ø–æ–ª–Ω—è–µ–º –∫–∞—Ä—Ç–æ—á–∫—É –¥–µ–ª–∞ #{cid}. –ó–∞–ø–æ–ª–Ω–µ–Ω–æ {filled}/{total}.\n"
        f"–°–µ–π—á–∞—Å: {CASE_CARD_FIELD_META[next_field]['title']}.\n"
        f"{prompt}"
    )
    await call.answer()

@dp.callback_query(lambda c: c.data.startswith("case:cardfield:"))
async def card_field_start(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    _, _, cid_str, field = call.data.split(":", maxsplit=3)
    cid = int(cid_str)

    if field not in CASE_CARD_FIELD_META:
        await call.answer()
        return

    # –ö—Ä–µ–¥–∏—Ç–æ—Ä—ã ‚Äî –æ—Ç–¥–µ–ª—å–Ω–æ–µ –º–µ–Ω—é, –Ω–µ –æ–±—ã—á–Ω—ã–π –≤–≤–æ–¥ —Ç–µ–∫—Å—Ç–∞
    if field == "creditors":
        await state.clear()
        await state.update_data(card_case_id=cid)
        await send_creditors_menu(call.message, uid, cid)
        await call.answer()
        return

    await state.clear()
    await state.update_data(card_case_id=cid, card_field_key=field)
    await state.set_state(CaseCardFill.waiting_value)

    prompt = CASE_CARD_FIELD_META[field]["prompt"] + "\n–û—Ç–ø—Ä–∞–≤—å '-' —á—Ç–æ–±—ã –æ—Å—Ç–∞–≤–∏—Ç—å –ø—É—Å—Ç—ã–º."
    await call.message.answer(prompt)
    await call.answer()

def _normalize_card_input(field: str, text: str) -> tuple[bool, str | int | None, str | None]:
    cleaned = text.strip()
    if not cleaned:
        return False, None, "–ü—É—Å—Ç–æ. –ü–æ–≤—Ç–æ—Ä–∏ –≤–≤–æ–¥."

    if field == "debtor_gender":
        gender = cleaned.lower()
        if gender in ("–º", "male", "m"):
            return True, "male", None
        if gender in ("–∂", "female", "f", "–∂–µ–Ω", "–∂–µ–Ω—â–∏–Ω–∞"):
            return True, "female", None
        return False, None, "–£–∫–∞–∂–∏ –ø–æ–ª –∫–∞–∫ –º/–∂ –∏–ª–∏ male/female."

    if field == "passport_date":
        try:
            datetime.strptime(cleaned, "%d.%m.%Y")
        except ValueError:
            return False, None, "–§–æ—Ä–º–∞—Ç –¥–∞—Ç—ã: –î–î.–ú–ú.–ì–ì–ì–ì. –ü–æ–ø—Ä–æ–±—É–π –µ—â—ë —Ä–∞–∑."
        return True, cleaned, None

    if field == "total_debt_rubles":
        try:
            val = int(cleaned)
        except ValueError:
            return False, None, "–ù—É–∂–Ω–æ —Ü–µ–ª–æ–µ —á–∏—Å–ª–æ –≤ —Ä—É–±–ª—è—Ö."
        if val < 0:
            return False, None, "–ó–Ω–∞—á–µ–Ω–∏–µ –Ω–µ –º–æ–∂–µ—Ç –±—ã—Ç—å –æ—Ç—Ä–∏—Ü–∞—Ç–µ–ª—å–Ω—ã–º."
        return True, val, None

    if field == "total_debt_kopeks":
        try:
            val = int(cleaned)
        except ValueError:
            return False, None, "–ù—É–∂–Ω–æ —Ü–µ–ª–æ–µ —á–∏—Å–ª–æ (0-99)."
        if val < 0 or val > 99:
            return False, None, "–ö–æ–ø–µ–π–∫–∏ –¥–æ–ª–∂–Ω—ã –±—ã—Ç—å –æ—Ç 0 –¥–æ 99."
        return True, val, None

    return True, cleaned, None


@dp.message(StateFilter(CaseCardFill.waiting_value))
async def case_card_value_set(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    data = await state.get_data()
    cid = data.get("card_cid") or data.get("card_case_id")
    field = data.get("card_field") or data.get("card_field_key")

    if not cid or not field:
        await state.clear()
        await message.answer("–ß—Ç–æ-—Ç–æ –ø–æ—à–ª–æ –Ω–µ —Ç–∞–∫. –û—Ç–∫—Ä–æ–π –∫–∞—Ä—Ç–æ—á–∫—É –∑–∞–Ω–æ–≤–æ —á–µ—Ä–µ–∑ –¥–µ–ª–æ.")
        return

    card = get_case_card(uid, int(cid))
    raw_text = message.text or ""
    if raw_text.strip() == "-":
        ok, value, error_msg = True, None, None
    else:
        ok, value, error_msg = _normalize_card_input(field, raw_text)
        if not ok:
            await message.answer(error_msg)
            return

    card[field] = value
    if field in {"debtor_last_name", "debtor_first_name", "debtor_middle_name"}:
        composed = _compose_debtor_full_name(card)
        if composed:
            card["debtor_full_name"] = composed

    upsert_case_card(uid, int(cid), card)
    next_field = None
    for key, _meta in CASE_CARD_FIELDS:
        val = card.get(key)
        if val is None or (isinstance(val, str) and not val.strip()):
            next_field = key
            break

    if next_field:
        await state.update_data(card_case_id=int(cid), card_field_key=next_field)
        await state.set_state(CaseCardFill.waiting_value)
        prompt = CASE_CARD_FIELD_META[next_field]["prompt"]
        filled, total = _card_completion_status(card)
        await message.answer(
            f"‚úÖ –°–æ—Ö—Ä–∞–Ω–µ–Ω–æ. –ó–∞–ø–æ–ª–Ω–µ–Ω–æ {filled}/{total}.\n"
            f"–î–∞–ª–µ–µ: {CASE_CARD_FIELD_META[next_field]['title']}.\n"
            f"{prompt}\n–û—Ç–ø—Ä–∞–≤—å '-' —á—Ç–æ–±—ã –æ—Å—Ç–∞–≤–∏—Ç—å –ø—É—Å—Ç—ã–º."
        )
        return

    await state.clear()
    filled, total = _card_completion_status(card)
    await message.answer(f"‚úÖ –ö–∞—Ä—Ç–æ—á–∫–∞ –∑–∞–ø–æ–ª–Ω–µ–Ω–∞. –ó–∞–ø–æ–ª–Ω–µ–Ω–æ {filled}/{total}.")

def _format_creditor_line(i: int, c: dict) -> str:
    name = (c.get("name") or "‚Äî").strip()
    inn = (c.get("inn") or "").strip()
    ogrn = (c.get("ogrn") or "").strip()
    debt_r = (c.get("debt_rubles") or "").strip()
    debt_k = (c.get("debt_kopeks") or "").strip()

    parts = [f"{i}) {name}"]
    ids = []
    if inn:
        ids.append(f"–ò–ù–ù {inn}")
    if ogrn:
        ids.append(f"–û–ì–†–ù {ogrn}")
    if ids:
        parts.append(" (" + ", ".join(ids) + ")")
    if debt_r or debt_k:
        dk = debt_k if debt_k else "00"
        dr = debt_r if debt_r else "0"
        parts.append(f" ‚Äî {dr} —Ä—É–±. {dk} –∫–æ–ø.")
    return "".join(parts)


def _safe_digits(s: str) -> str:
    return "".join(ch for ch in s if ch.isdigit())


async def send_creditors_menu(message_target, uid: int, cid: int) -> None:
    """Helper —Ñ—É–Ω–∫—Ü–∏—è –¥–ª—è –æ—Ç–ø—Ä–∞–≤–∫–∏ –º–µ–Ω—é –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤."""
    card = get_case_card(uid, cid) or {}
    creditors = card.get("creditors")
    if not isinstance(creditors, list):
        creditors = []

    creditors_text = (card.get("creditors_text") or "").strip()

    lines = [f"üë• –ö—Ä–µ–¥–∏—Ç–æ—Ä—ã –¥–ª—è –¥–µ–ª–∞ #{cid}"]
    lines.append(f"–°–ø–∏—Å–æ–∫: {len(creditors)}")
    if creditors_text:
        lines.append("–ï—Å—Ç—å —Ä—É—á–Ω–æ–π —Ç–µ–∫—Å—Ç creditors_text: ‚úÖ (–æ–Ω –∏–º–µ–µ—Ç –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç –ø—Ä–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏)")
    else:
        lines.append("–†—É—á–Ω–æ–π —Ç–µ–∫—Å—Ç creditors_text: ‚Äî")

    if creditors:
        lines.append("")
        lines.append("–¢–µ–∫—É—â–∏–π —Å–ø–∏—Å–æ–∫:")
        for i, c in enumerate(creditors, 1):
            lines.append(_format_creditor_line(i, c))

    kb = InlineKeyboardBuilder()
    kb.button(text="‚ûï –î–æ–±–∞–≤–∏—Ç—å –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞", callback_data=f"creditors:add:{cid}")
    kb.button(text="üóë –£–¥–∞–ª–∏—Ç—å –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞", callback_data=f"creditors:del:{cid}")
    kb.button(text="üßæ –í–≤–µ—Å—Ç–∏ –æ–¥–Ω–∏–º —Ç–µ–∫—Å—Ç–æ–º", callback_data=f"creditors:text:{cid}")
    kb.button(text="üßπ –û—á–∏—Å—Ç–∏—Ç—å creditors_text", callback_data=f"creditors:text_clear:{cid}")
    kb.button(text="üîô –ù–∞–∑–∞–¥ –≤ –∫–∞—Ä—Ç–æ—á–∫—É", callback_data=f"case:card:{cid}")
    kb.adjust(1)

    await message_target.answer("\n".join(lines), reply_markup=kb.as_markup())


# ============================================================================
# REGISTER CALLBACK ROUTER (handlers/callbacks.py)
# Must be done AFTER all helper functions are defined for circular import
# ============================================================================
from handlers.callbacks import callback_router
dp.include_router(callback_router)


@dp.callback_query(lambda c: c.data.startswith("case:creditors:"))
async def creditors_menu(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    cid = int(call.data.split(":")[2])
    await state.clear()
    await state.update_data(card_case_id=cid)

    await send_creditors_menu(call.message, uid, cid)
    await call.answer()


@dp.callback_query(lambda c: c.data.startswith("creditors:add:"))
async def creditors_add_start(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return
    cid = int(call.data.split(":")[2])

    await state.clear()
    await state.update_data(card_case_id=cid, creditor_tmp={})
    await state.set_state(CreditorsFill.name)
    await call.message.answer("–í–≤–µ–¥–∏ –Ω–∞–∑–≤–∞–Ω–∏–µ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞ (–æ–±—è–∑–∞—Ç–µ–ª—å–Ω–æ).")
    await call.answer()


@dp.callback_query(lambda c: c.data.startswith("creditors:del:"))
async def creditors_delete_menu(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return
    cid = int(call.data.split(":")[2])

    card = get_case_card(uid, cid) or {}
    creditors = card.get("creditors")
    if not isinstance(creditors, list) or not creditors:
        await call.message.answer("–°–ø–∏—Å–æ–∫ –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤ –ø—É—Å—Ç.")
        await call.answer()
        return

    kb = InlineKeyboardBuilder()
    lines = [f"üóë –£–¥–∞–ª–µ–Ω–∏–µ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞ (–¥–µ–ª–æ #{cid})", "–í—ã–±–µ—Ä–∏ –Ω–æ–º–µ—Ä:"]
    for i, c in enumerate(creditors, 1):
        lines.append(_format_creditor_line(i, c))
        kb.button(text=f"–£–¥–∞–ª–∏—Ç—å #{i}", callback_data=f"creditors:delone:{cid}:{i}")
    kb.button(text="üîô –ù–∞–∑–∞–¥", callback_data=f"case:creditors:{cid}")
    kb.adjust(1)

    await call.message.answer("\n".join(lines), reply_markup=kb.as_markup())
    await call.answer()


@dp.callback_query(lambda c: c.data.startswith("creditors:delone:"))
async def creditors_delete_one(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    _, _, cid_str, idx_str = call.data.split(":")
    cid = int(cid_str)
    idx = int(idx_str)

    card = get_case_card(uid, cid) or {}
    creditors = card.get("creditors")
    if not isinstance(creditors, list):
        creditors = []
    if idx < 1 or idx > len(creditors):
        await call.message.answer("–ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–π –Ω–æ–º–µ—Ä.")
        await call.answer()
        return

    removed = creditors.pop(idx - 1)
    card["creditors"] = creditors
    upsert_case_card(uid, cid, card)

    name = (removed.get("name") or "‚Äî").strip()
    await call.message.answer(f"‚úÖ –£–¥–∞–ª–µ–Ω–æ: {name}")
    # –≤–µ—Ä–Ω—ë–º –º–µ–Ω—é –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤
    await creditors_menu(call, state)


@dp.callback_query(lambda c: c.data.startswith("creditors:text_clear:"))
async def creditors_text_clear(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return
    cid = int(call.data.split(":")[2])

    card = get_case_card(uid, cid) or {}
    card["creditors_text"] = None
    upsert_case_card(uid, cid, card)

    await call.message.answer("‚úÖ creditors_text –æ—á–∏—â–µ–Ω.")
    await creditors_menu(call, state)


@dp.callback_query(lambda c: c.data.startswith("creditors:text:"))
async def creditors_text_start(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return
    cid = int(call.data.split(":")[2])

    await state.clear()
    await state.update_data(card_case_id=cid)
    await state.set_state(CreditorsFill.creditors_text)

    await call.message.answer(
        "–í—Å—Ç–∞–≤—å —Ç–µ–∫—Å—Ç –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤ –æ–¥–Ω–∏–º –±–ª–æ–∫–æ–º.\n"
        "–û–Ω –±—É–¥–µ—Ç –∏–º–µ—Ç—å –ø—Ä–∏–æ—Ä–∏—Ç–µ—Ç –Ω–∞–¥ —Å–ø–∏—Å–∫–æ–º creditors –ø—Ä–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏.\n"
        "–û—Ç–ø—Ä–∞–≤—å '-' —á—Ç–æ–±—ã –æ—á–∏—Å—Ç–∏—Ç—å."
    )
    await call.answer()


@dp.message(StateFilter(CreditorsFill.creditors_text))
async def creditors_text_set(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    data = await state.get_data()
    cid = int(data.get("card_case_id"))

    text = (message.text or "").strip()
    card = get_case_card(uid, cid) or {}

    if text == "-":
        card["creditors_text"] = None
        upsert_case_card(uid, cid, card)
        await state.clear()
        await message.answer("‚úÖ creditors_text –æ—á–∏—â–µ–Ω.")
        # –ø–æ–∫–∞–∑–∞—Ç—å –º–µ–Ω—é –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤
        await send_creditors_menu(message, uid, cid)
        return

    card["creditors_text"] = text
    upsert_case_card(uid, cid, card)

    await state.clear()
    await message.answer("‚úÖ –°–æ—Ö—Ä–∞–Ω–µ–Ω–æ creditors_text.")
    await send_creditors_menu(message, uid, cid)


@dp.message(StateFilter(CreditorsFill.name))
async def creditors_step_name(message: Message, state: FSMContext):
    txt = (message.text or "").strip()
    if not txt or txt == "-":
        await message.answer("–ù–∞–∑–≤–∞–Ω–∏–µ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞ –æ–±—è–∑–∞—Ç–µ–ª—å–Ω–æ. –í–≤–µ–¥–∏ –µ—â—ë —Ä–∞–∑.")
        return

    data = await state.get_data()
    tmp = data.get("creditor_tmp") or {}
    tmp["name"] = txt
    await state.update_data(creditor_tmp=tmp)
    await state.set_state(CreditorsFill.inn)
    await message.answer("–ò–ù–ù (–º–æ–∂–Ω–æ '-' —á—Ç–æ–±—ã –ø—Ä–æ–ø—É—Å—Ç–∏—Ç—å).")


@dp.message(StateFilter(CreditorsFill.inn))
async def creditors_step_inn(message: Message, state: FSMContext):
    txt = (message.text or "").strip()
    data = await state.get_data()
    tmp = data.get("creditor_tmp") or {}

    if txt != "-" and txt:
        tmp["inn"] = _safe_digits(txt)
    await state.update_data(creditor_tmp=tmp)
    await state.set_state(CreditorsFill.ogrn)
    await message.answer("–û–ì–†–ù (–º–æ–∂–Ω–æ '-' —á—Ç–æ–±—ã –ø—Ä–æ–ø—É—Å—Ç–∏—Ç—å).")


@dp.message(StateFilter(CreditorsFill.ogrn))
async def creditors_step_ogrn(message: Message, state: FSMContext):
    txt = (message.text or "").strip()
    data = await state.get_data()
    tmp = data.get("creditor_tmp") or {}

    if txt != "-" and txt:
        tmp["ogrn"] = _safe_digits(txt)
    await state.update_data(creditor_tmp=tmp)
    await state.set_state(CreditorsFill.address)
    await message.answer("–ê–¥—Ä–µ—Å –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞ (–º–æ–∂–Ω–æ '-' —á—Ç–æ–±—ã –ø—Ä–æ–ø—É—Å—Ç–∏—Ç—å).")


@dp.message(StateFilter(CreditorsFill.address))
async def creditors_step_address(message: Message, state: FSMContext):
    txt = (message.text or "").strip()
    data = await state.get_data()
    tmp = data.get("creditor_tmp") or {}

    if txt != "-" and txt:
        tmp["address"] = txt
    await state.update_data(creditor_tmp=tmp)
    await state.set_state(CreditorsFill.debt_rubles)
    await message.answer("–°—É–º–º–∞ –¥–æ–ª–≥–∞ (—Ä—É–±–ª–∏) (–º–æ–∂–Ω–æ '-' —á—Ç–æ–±—ã –ø—Ä–æ–ø—É—Å—Ç–∏—Ç—å).")


@dp.message(StateFilter(CreditorsFill.debt_rubles))
async def creditors_step_debt_rubles(message: Message, state: FSMContext):
    txt = (message.text or "").strip()
    data = await state.get_data()
    tmp = data.get("creditor_tmp") or {}

    if txt != "-" and txt:
        digits = _safe_digits(txt)
        if digits == "":
            await message.answer("–ù—É–∂–Ω–æ —á–∏—Å–ª–æ (–∏–ª–∏ '-' —á—Ç–æ–±—ã –ø—Ä–æ–ø—É—Å—Ç–∏—Ç—å).")
            return
        tmp["debt_rubles"] = digits
    await state.update_data(creditor_tmp=tmp)
    await state.set_state(CreditorsFill.debt_kopeks)
    await message.answer("–°—É–º–º–∞ –¥–æ–ª–≥–∞ (–∫–æ–ø–µ–π–∫–∏ 0-99) (–º–æ–∂–Ω–æ '-' —á—Ç–æ–±—ã –ø—Ä–æ–ø—É—Å—Ç–∏—Ç—å).")


@dp.message(StateFilter(CreditorsFill.debt_kopeks))
async def creditors_step_debt_kopeks(message: Message, state: FSMContext):
    txt = (message.text or "").strip()
    data = await state.get_data()
    tmp = data.get("creditor_tmp") or {}

    if txt != "-" and txt:
        digits = _safe_digits(txt)
        if digits == "":
            await message.answer("–ù—É–∂–Ω–æ —á–∏—Å–ª–æ 0-99 (–∏–ª–∏ '-' —á—Ç–æ–±—ã –ø—Ä–æ–ø—É—Å—Ç–∏—Ç—å).")
            return
        try:
            val = int(digits)
        except ValueError:
            await message.answer("–ù—É–∂–Ω–æ —á–∏—Å–ª–æ 0-99.")
            return
        if val < 0 or val > 99:
            await message.answer("–ö–æ–ø–µ–π–∫–∏ –¥–æ–ª–∂–Ω—ã –±—ã—Ç—å 0-99.")
            return
        tmp["debt_kopeks"] = f"{val:02d}"
    await state.update_data(creditor_tmp=tmp)
    await state.set_state(CreditorsFill.note)
    await message.answer("–û—Å–Ω–æ–≤–∞–Ω–∏–µ/–∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π (–Ω–∞–ø—Ä–∏–º–µ—Ä: –≤—ã–ø–∏—Å–∫–∞ –û–ö–ë) (–º–æ–∂–Ω–æ '-' —á—Ç–æ–±—ã –ø—Ä–æ–ø—É—Å—Ç–∏—Ç—å).")


@dp.message(StateFilter(CreditorsFill.note))
async def creditors_step_note(message: Message, state: FSMContext):
    txt = (message.text or "").strip()
    data = await state.get_data()
    cid = int(data.get("card_case_id"))
    tmp = data.get("creditor_tmp") or {}

    if txt != "-" and txt:
        tmp["note"] = txt

    # —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å –≤ –∫–∞—Ä—Ç–æ—á–∫—É
    card = get_case_card(message.from_user.id, cid) or {}
    creditors = card.get("creditors")
    if not isinstance(creditors, list):
        creditors = []
    creditors.append(tmp)
    card["creditors"] = creditors
    upsert_case_card(message.from_user.id, cid, card)

    await state.clear()

    kb = InlineKeyboardBuilder()
    kb.button(text="‚ûï –î–æ–±–∞–≤–∏—Ç—å –µ—â—ë", callback_data=f"creditors:add:{cid}")
    kb.button(text="‚úÖ –ì–æ—Ç–æ–≤–æ", callback_data=f"case:creditors:{cid}")
    kb.adjust(1)

    await message.answer(
        f"‚úÖ –ö—Ä–µ–¥–∏—Ç–æ—Ä –¥–æ–±–∞–≤–ª–µ–Ω. –°–µ–π—á–∞—Å –≤ —Å–ø–∏—Å–∫–µ: {len(creditors)}",
        reply_markup=kb.as_markup(),
    )

@dp.callback_query(lambda c: c.data.startswith("case:edit:") and c.data.count(":") == 3)
async def case_edit_start(call: CallbackQuery, state: FSMContext):
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    _, _, cid_str, field = call.data.split(":")
    cid = int(cid_str)

    # –ø—Ä–æ–≤–µ—Ä–∏–º, —á—Ç–æ –¥–µ–ª–æ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç –∏ —Ç–≤–æ—ë
    row = get_case(uid, cid)
    if not row:
        await call.message.answer("–î–µ–ª–æ –Ω–µ –Ω–∞–π–¥–µ–Ω–æ.")
        await call.answer()
        return

    await state.clear()
    await state.update_data(edit_cid=cid, edit_field=field)
    await state.set_state(CaseEdit.value)

    field_titles = {
        "case_number": "–Ω–æ–º–µ—Ä –¥–µ–ª–∞",
        "court": "—Å—É–¥",
        "judge": "—Å—É–¥—å—é",
        "fin_manager": "—Ñ–∏–Ω–∞–Ω—Å–æ–≤–æ–≥–æ —É–ø—Ä–∞–≤–ª—è—é—â–µ–≥–æ",
        "stage": "—Å—Ç–∞–¥–∏—é",
        "notes": "–∑–∞–º–µ—Ç–∫–∏",
    }
    title = field_titles.get(field, field)

    kb = InlineKeyboardBuilder()
    kb.button(text="‚ùå –û—Ç–º–µ–Ω–∞", callback_data=f"case:edit:{cid}")
    kb.adjust(1)

    await call.message.answer(
        f"–í–≤–µ–¥–∏ –Ω–æ–≤–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ –¥–ª—è ¬´{title}¬ª.\n–ï—Å–ª–∏ –Ω—É–∂–Ω–æ –æ—á–∏—Å—Ç–∏—Ç—å –ø–æ–ª–µ ‚Äî –æ—Ç–ø—Ä–∞–≤—å `-`.",
        reply_markup=kb.as_markup(),
    )

    await call.answer()
@dp.message(StateFilter(CaseEdit.value))
async def case_edit_apply(message: Message, state: FSMContext):
    uid = message.from_user.id
    if not is_allowed(uid):
        return

    data = await state.get_data()
    cid = data.get("edit_cid")
    field = data.get("edit_field")

    if not cid or not field:
        await state.clear()
        await message.answer("–ß—Ç–æ-—Ç–æ –ø–æ—à–ª–æ –Ω–µ —Ç–∞–∫. –ù–∞—á–Ω–∏ –∑–∞–Ω–æ–≤–æ —á–µ—Ä–µ–∑ –∫–∞—Ä—Ç–æ—á–∫—É –¥–µ–ª–∞.")
        return

    text = (message.text or "").strip()
    if not text:
        await message.answer("–ü—É—Å—Ç–æ. –í–≤–µ–¥–∏ –∑–Ω–∞—á–µ–Ω–∏–µ –∏–ª–∏ '-' —á—Ç–æ–±—ã –æ—á–∏—Å—Ç–∏—Ç—å.")
        return

    value = None if text == "-" else text

    if field in ("case_number", "court", "judge", "fin_manager"):
        update_case_fields(
            uid,
            cid,
            case_number=value if field == "case_number" else None,
            court=value if field == "court" else None,
            judge=value if field == "judge" else None,
            fin_manager=value if field == "fin_manager" else None,
        )
    elif field in ("stage", "notes"):
        update_case_meta(
            uid,
            cid,
            stage=value if field == "stage" else None,
            notes=value if field == "notes" else None,
        )
    else:
        await message.answer("–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–æ–µ –ø–æ–ª–µ –¥–ª—è —Ä–µ–¥–∞–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏—è.")
        await state.clear()
        return
    await state.clear()

    # –ø–æ—Å–ª–µ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è ‚Äî –≤–µ—Ä–Ω—É—Ç—å –≤ –º–µ–Ω—é —Ä–µ–¥–∞–∫—Ç–∏—Ä–æ–≤–∞–Ω–∏—è –∫–∞—Ä—Ç–æ—á–∫–∏
    fake = type("X", (), {})()
    fake.from_user = message.from_user
    fake.data = f"case:edit:{cid}"
    fake.message = message
    await case_edit_menu(fake, state)

# =========================
# ROUTER CONFLICT FIX: Duplicate commands removed
# =========================
# These command handlers are now handled by bankrot_bot/handlers/cases.py
# which uses the newer PostgreSQL implementation instead of SQLite.
# Commands available in handlers/cases.py:
#   - /newcase (was /case_new here)
#   - /mycases (was /cases here)
#   - /case <id> (same)
#   - /setactive <id>
#   - /editcase
#   - /cancel
# =========================
# @dp.message(Command("case_new"))  # REMOVED - use /newcase from handlers/cases.py
# @dp.message(Command("cases"))     # REMOVED - use /mycases from handlers/cases.py
# @dp.message(Command("case"))      # REMOVED - use /case from handlers/cases.py


# =========================
# FSM handlers for AddParty and AddAsset (must be BEFORE catch-all!)
# =========================

@dp.message(StateFilter(AddParty.name))
async def process_party_name(message: Message, state: FSMContext):
    """–û–±—Ä–∞–±–æ—Ç–∞—Ç—å –≤–≤–æ–¥ –∏–º–µ–Ω–∏ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞/–¥–æ–ª–∂–Ω–∏–∫–∞. ONLY active in AddParty.name state."""
    logger.info(f"User {message.from_user.id} entered party name in AddParty FSM")
    await state.update_data(name=message.text.strip())
    await message.answer("–í–≤–µ–¥–∏—Ç–µ —Å—É–º–º—É (–Ω–∞–ø—Ä–∏–º–µ—Ä: 100000 –∏–ª–∏ 100 000.50):")
    await state.set_state(AddParty.amount)


@dp.message(StateFilter(AddParty.amount))
async def process_party_amount(message: Message, state: FSMContext):
    """–û–±—Ä–∞–±–æ—Ç–∞—Ç—å –≤–≤–æ–¥ —Å—É–º–º—ã. ONLY active in AddParty.amount state."""
    logger.info(f"User {message.from_user.id} entered party amount in AddParty FSM")
    # Normalize to JSON-safe string for Redis FSM storage
    amount_str = normalize_amount_to_string(message.text)
    if amount_str is None:
        await message.answer(
            "‚ùå –ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç —Å—É–º–º—ã. –í–≤–µ–¥–∏—Ç–µ —á–∏—Å–ª–æ (–Ω–∞–ø—Ä–∏–º–µ—Ä: 100000 –∏–ª–∏ 100 000.50):"
        )
        return
    await state.update_data(amount=amount_str)
    await message.answer("–í–≤–µ–¥–∏—Ç–µ –æ—Å–Ω–æ–≤–∞–Ω–∏–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è/–¥–æ–ª–≥–∞ (–∏–ª–∏ '-' –¥–ª—è –ø—Ä–æ–ø—É—Å–∫–∞):")
    await state.set_state(AddParty.basis)


@dp.message(StateFilter(AddParty.basis))
async def process_party_basis(message: Message, state: FSMContext):
    """–ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞/–¥–æ–ª–∂–Ω–∏–∫–∞. ONLY active in AddParty.basis state."""
    logger.info(f"User {message.from_user.id} completing AddParty FSM")
    basis = message.text.strip() if message.text.strip() != "-" else None

    data = await state.get_data()
    case_id = data["case_id"]
    role = data["role"]
    name = data["name"]
    amount_str = data["amount"]

    # Convert JSON-safe string back to Decimal for DB insertion
    amount = string_to_decimal(amount_str)

    from bankrot_bot.database import get_session
    async with get_session() as session:
        await add_case_party(session, case_id, role, name, amount, basis=basis)
        await session.commit()
        logger.info(f"User {message.from_user.id} added party to case {case_id}: {name}, {amount}")

    role_text = "–ö—Ä–µ–¥–∏—Ç–æ—Ä" if role == "creditor" else "–î–æ–ª–∂–Ω–∏–∫"
    await message.answer(f"‚úÖ {role_text} –¥–æ–±–∞–≤–ª–µ–Ω: {name}, —Å—É–º–º–∞: {amount:.2f} ‚ÇΩ")
    await state.clear()


@dp.message(StateFilter(AddAsset.kind))
async def process_asset_kind(message: Message, state: FSMContext):
    """–û–±—Ä–∞–±–æ—Ç–∞—Ç—å –≤–≤–æ–¥ –≤–∏–¥–∞ –∏–º—É—â–µ—Å—Ç–≤–∞. ONLY active in AddAsset.kind state."""
    logger.info(f"User {message.from_user.id} entered asset kind in AddAsset FSM")
    await state.update_data(kind=message.text.strip())
    await message.answer("–í–≤–µ–¥–∏—Ç–µ –æ–ø–∏—Å–∞–Ω–∏–µ (–∞–¥—Ä–µ—Å, –º–∞—Ä–∫–∞, —Ä–µ–∫–≤–∏–∑–∏—Ç—ã –∏ —Ç.–ø.):")
    await state.set_state(AddAsset.description)


@dp.message(StateFilter(AddAsset.description))
async def process_asset_description(message: Message, state: FSMContext):
    """–û–±—Ä–∞–±–æ—Ç–∞—Ç—å –≤–≤–æ–¥ –æ–ø–∏—Å–∞–Ω–∏—è. ONLY active in AddAsset.description state."""
    logger.info(f"User {message.from_user.id} entered asset description in AddAsset FSM")
    await state.update_data(description=message.text.strip())
    await message.answer("–í–≤–µ–¥–∏—Ç–µ —Å—Ç–æ–∏–º–æ—Å—Ç—å (–∏–ª–∏ '-' –¥–ª—è –ø—Ä–æ–ø—É—Å–∫–∞):")
    await state.set_state(AddAsset.value)


@dp.message(StateFilter(AddAsset.value))
async def process_asset_value(message: Message, state: FSMContext):
    """–ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ –∏–º—É—â–µ—Å—Ç–≤–∞. ONLY active in AddAsset.value state."""
    logger.info(f"User {message.from_user.id} completing AddAsset FSM")
    value_text = message.text.strip()

    # Parse value with validation
    value = None
    if value_text != "-":
        # Use normalize_amount_to_string for consistent validation
        normalized = normalize_amount_to_string(value_text)
        if normalized is None:
            await message.answer(
                "‚ùå –ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç —Å—Ç–æ–∏–º–æ—Å—Ç–∏. –í–≤–µ–¥–∏—Ç–µ —á–∏—Å–ª–æ (–Ω–∞–ø—Ä–∏–º–µ—Ä: 100000 –∏–ª–∏ 100 000.50) –∏–ª–∏ '-' –¥–ª—è –ø—Ä–æ–ø—É—Å–∫–∞:"
            )
            return
        value = string_to_decimal(normalized)

    data = await state.get_data()
    case_id = data["case_id"]
    kind = data["kind"]
    description = data["description"]

    from bankrot_bot.database import get_session
    async with get_session() as session:
        await add_case_asset(session, case_id, kind, description, value=value)
        await session.commit()
        logger.info(f"User {message.from_user.id} added asset to case {case_id}: {kind}, {value}")

    value_str = f"{float(value):.2f} ‚ÇΩ" if value else "–Ω–µ —É–∫–∞–∑–∞–Ω–∞"
    await message.answer(f"‚úÖ –ò–º—É—â–µ—Å—Ç–≤–æ –¥–æ–±–∞–≤–ª–µ–Ω–æ:\n{kind}\n–°—Ç–æ–∏–º–æ—Å—Ç—å: {value_str}")
    await state.clear()


# =========================
# Catch-all message handler (must be AFTER FSM handlers!)
# =========================
@dp.message()
async def main_text_router(message: Message, state: FSMContext):
    # –ï—Å–ª–∏ –∏–¥—ë—Ç FSM (—Å–æ–∑–¥–∞–Ω–∏–µ –¥–µ–ª–∞ –∏ —Ç.–ø.) ‚Äî –Ω–µ –º–µ—à–∞–µ–º
    if await state.get_state() is not None:
        return
    uid = message.from_user.id
    if not is_allowed(uid):
        return
    
    # ‚úÖ –ö–æ–º–∞–Ω–¥—ã –≤—Å–µ–≥–¥–∞ —Ä–∞–∑—Ä–µ—à–µ–Ω—ã
    if message.text and message.text.startswith("/"):
        return
    
    if uid not in USER_FLOW:
        await message.answer("–°–Ω–∞—á–∞–ª–∞ –≤—ã–±–µ—Ä–∏ –∑–∞–¥–∞—á—É —á–µ—Ä–µ–∑ /start.")
        return

    # –¥–∞–ª—å—à–µ ‚Äî —Ç–≤–æ—è —Å—Ç–∞—Ä–∞—è –ª–æ–≥–∏–∫–∞ USER_FLOW (motion / settlement)
    flow = USER_FLOW[uid]
    text = (message.text or "").strip()

    if flow.get("flow") == "settlement":
        step = int(flow.get("step", 0))
        if step >= len(SETTLEMENT_STEPS):
            cancel_flow(uid)
            await message.answer("–ê–Ω–∫–µ—Ç–∞ –∑–∞–≤–µ—Ä—à–µ–Ω–∞. –ú–µ–Ω—é üëá", reply_markup=main_keyboard())
            return

        key = SETTLEMENT_STEPS[step][0]
        flow["answers"][key] = text
        step += 1
        flow["step"] = step

        if step < len(SETTLEMENT_STEPS):
            await message.answer(SETTLEMENT_STEPS[step][1], reply_markup=settlement_actions_keyboard())
            return

        await message.answer("–ü—Ä–∏–Ω—è–ª –¥–∞–Ω–Ω—ã–µ. –ì–æ—Ç–æ–≤–ª—é –ø—Ä–æ–µ–∫—Ç –º–∏—Ä–æ–≤–æ–≥–æ‚Ä¶")
        try:
            user_text = build_settlement_user_text(flow.get("answers", {}))
            result = await gigachat_chat(
                auth_key=AUTH_KEY,
                scope=SCOPE,
                model=MODEL,
                system_prompt=system_prompt_for_settlement(),
                user_text=user_text,
            )
            LAST_RESULT[uid] = result
            await message.answer(result)
            await message.answer("–≠–∫—Å–ø–æ—Ä—Ç üëá", reply_markup=export_keyboard())
        except Exception as e:
            await message.answer(f"–û—à–∏–±–∫–∞ GigaChat:\n{e}")

        cancel_flow(uid)
        return


# ========== –•—ç–Ω–¥–ª–µ—Ä—ã –¥–ª—è –ö—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤/–î–æ–ª–∂–Ω–∏–∫–æ–≤ ==========

@dp.callback_query(F.data.startswith("case:parties:"))
async def show_case_parties(call: CallbackQuery):
    """–ü–æ–∫–∞–∑–∞—Ç—å —Å–ø–∏—Å–æ–∫ –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤/–¥–æ–ª–∂–Ω–∏–∫–æ–≤ –ø–æ –¥–µ–ª—É."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    case_id = int(call.data.split(":")[-1])

    from bankrot_bot.database import get_session
    async with get_session() as session:
        parties = await get_case_parties(session, case_id)
        totals = calculate_parties_totals(parties)

        text = f"üí∞ –ö—Ä–µ–¥–∏—Ç–æ—Ä—ã –∏ –¥–æ–ª–∂–Ω–∏–∫–∏ –ø–æ –¥–µ–ª—É #{case_id}\n\n"
        text += f"–ö—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤: {totals['creditors_count']}, —Å—É–º–º–∞: {totals['total_creditors']:.2f} ‚ÇΩ\n"
        text += f"–î–æ–ª–∂–Ω–∏–∫–æ–≤: {totals['debtors_count']}, —Å—É–º–º–∞: {totals['total_debtors']:.2f} ‚ÇΩ"

        await call.message.answer(
            text,
            reply_markup=case_parties_ikb(case_id, parties, totals['creditors_count'], totals['debtors_count'])
        )
    await call.answer()


@dp.callback_query(F.data.startswith("party:add_creditor:") | F.data.startswith("party:add_debtor:"))
async def start_add_party(call: CallbackQuery, state: FSMContext):
    """–ù–∞—á–∞—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞/–¥–æ–ª–∂–Ω–∏–∫–∞."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    role = "creditor" if "creditor" in call.data else "debtor"
    case_id = int(parts[-1])

    await state.update_data(case_id=case_id, role=role)
    await state.set_state(AddParty.name)

    role_text = "–∫—Ä–µ–¥–∏—Ç–æ—Ä–∞" if role == "creditor" else "–¥–æ–ª–∂–Ω–∏–∫–∞"
    await call.message.answer(f"–î–æ–±–∞–≤–ª–µ–Ω–∏–µ {role_text}\n\n–í–≤–µ–¥–∏—Ç–µ –Ω–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ/–§–ò–û:")
    await call.answer()


@dp.callback_query(F.data.startswith("party:view:"))
async def view_party(call: CallbackQuery):
    """–ü—Ä–æ—Å–º–æ—Ç—Ä –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞/–¥–æ–ª–∂–Ω–∏–∫–∞."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    party_id = int(call.data.split(":")[-1])

    from bankrot_bot.database import get_session
    from bankrot_bot.models.case_party import CaseParty
    from sqlalchemy import select

    async with get_session() as session:
        stmt = select(CaseParty).where(CaseParty.id == party_id)
        result = await session.execute(stmt)
        party = result.scalar_one_or_none()

        if not party:
            await call.answer("–ó–∞–ø–∏—Å—å –Ω–µ –Ω–∞–π–¥–µ–Ω–∞", show_alert=True)
            return

        role_text = "–ö—Ä–µ–¥–∏—Ç–æ—Ä" if party.role == "creditor" else "–î–æ–ª–∂–Ω–∏–∫"
        text = f"{role_text}\n\n"
        text += f"–ù–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ: {party.name}\n"
        text += f"–°—É–º–º–∞: {float(party.amount):.2f} {party.currency}\n"
        if party.basis:
            text += f"–û—Å–Ω–æ–≤–∞–Ω–∏–µ: {party.basis}\n"
        if party.notes:
            text += f"–ü—Ä–∏–º–µ—á–∞–Ω–∏—è: {party.notes}\n"

        await call.message.answer(text, reply_markup=party_view_ikb(party_id, party.case_id))
    await call.answer()


@dp.callback_query(F.data.startswith("party:delete:"))
async def delete_party(call: CallbackQuery):
    """–£–¥–∞–ª–∏—Ç—å –∫—Ä–µ–¥–∏—Ç–æ—Ä–∞/–¥–æ–ª–∂–Ω–∏–∫–∞."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    party_id = int(parts[2])
    case_id = int(parts[3])

    from bankrot_bot.database import get_session
    async with get_session() as session:
        success = await delete_case_party(session, party_id, case_id)
        await session.commit()

        if success:
            await call.message.answer("‚úÖ –ó–∞–ø–∏—Å—å —É–¥–∞–ª–µ–Ω–∞")
        else:
            await call.answer("–û—à–∏–±–∫–∞ —É–¥–∞–ª–µ–Ω–∏—è", show_alert=True)
    await call.answer()


# ========== –•—ç–Ω–¥–ª–µ—Ä—ã –¥–ª—è –û–ø–∏—Å–∏ –∏–º—É—â–µ—Å—Ç–≤–∞ ==========

@dp.callback_query(F.data.startswith("case:assets:"))
async def show_case_assets(call: CallbackQuery):
    """–ü–æ–∫–∞–∑–∞—Ç—å –æ–ø–∏—Å—å –∏–º—É—â–µ—Å—Ç–≤–∞ –ø–æ –¥–µ–ª—É."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    case_id = int(call.data.split(":")[-1])

    from bankrot_bot.database import get_session
    async with get_session() as session:
        assets = await get_case_assets(session, case_id)
        total = calculate_assets_total(assets)

        text = f"üè† –û–ø–∏—Å—å –∏–º—É—â–µ—Å—Ç–≤–∞ –ø–æ –¥–µ–ª—É #{case_id}\n\n"
        text += f"–ó–∞–ø–∏—Å–µ–π: {len(assets)}\n"
        text += f"–û–±—â–∞—è —Å—Ç–æ–∏–º–æ—Å—Ç—å: {float(total):.2f} ‚ÇΩ"

        await call.message.answer(
            text,
            reply_markup=case_assets_ikb(case_id, assets, float(total))
        )
    await call.answer()


@dp.callback_query(F.data.startswith("asset:add:"))
async def start_add_asset(call: CallbackQuery, state: FSMContext):
    """–ù–∞—á–∞—Ç—å –¥–æ–±–∞–≤–ª–µ–Ω–∏–µ –∏–º—É—â–µ—Å—Ç–≤–∞."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    case_id = int(call.data.split(":")[-1])
    await state.update_data(case_id=case_id)
    await state.set_state(AddAsset.kind)

    await call.message.answer("–î–æ–±–∞–≤–ª–µ–Ω–∏–µ –∏–º—É—â–µ—Å—Ç–≤–∞\n\n–í–≤–µ–¥–∏—Ç–µ –≤–∏–¥ –∏–º—É—â–µ—Å—Ç–≤–∞ (–Ω–∞–ø—Ä–∏–º–µ—Ä: –∫–≤–∞—Ä—Ç–∏—Ä–∞, –∞–≤—Ç–æ–º–æ–±–∏–ª—å, –∞–∫—Ü–∏–∏):")
    await call.answer()


@dp.callback_query(F.data.startswith("asset:view:"))
async def view_asset(call: CallbackQuery):
    """–ü—Ä–æ—Å–º–æ—Ç—Ä –∏–º—É—â–µ—Å—Ç–≤–∞."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    asset_id = int(call.data.split(":")[-1])

    from bankrot_bot.database import get_session
    from bankrot_bot.models.case_asset import CaseAsset
    from sqlalchemy import select

    async with get_session() as session:
        stmt = select(CaseAsset).where(CaseAsset.id == asset_id)
        result = await session.execute(stmt)
        asset = result.scalar_one_or_none()

        if not asset:
            await call.answer("–ó–∞–ø–∏—Å—å –Ω–µ –Ω–∞–π–¥–µ–Ω–∞", show_alert=True)
            return

        text = f"üè† {asset.kind}\n\n"
        text += f"–û–ø–∏—Å–∞–Ω–∏–µ: {asset.description}\n"
        if asset.qty_or_area:
            text += f"–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ/–ø–ª–æ—â–∞–¥—å: {asset.qty_or_area}\n"
        if asset.value:
            text += f"–°—Ç–æ–∏–º–æ—Å—Ç—å: {float(asset.value):.2f} ‚ÇΩ\n"
        if asset.notes:
            text += f"–ü—Ä–∏–º–µ—á–∞–Ω–∏—è: {asset.notes}\n"

        await call.message.answer(text, reply_markup=asset_view_ikb(asset_id, asset.case_id))
    await call.answer()


@dp.callback_query(F.data.startswith("asset:delete:"))
async def delete_asset(call: CallbackQuery):
    """–£–¥–∞–ª–∏—Ç—å –∏–º—É—â–µ—Å—Ç–≤–æ."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    asset_id = int(parts[2])
    case_id = int(parts[3])

    from bankrot_bot.database import get_session
    async with get_session() as session:
        success = await delete_case_asset(session, asset_id, case_id)
        await session.commit()

        if success:
            await call.message.answer("‚úÖ –ó–∞–ø–∏—Å—å —É–¥–∞–ª–µ–Ω–∞")
        else:
            await call.answer("–û—à–∏–±–∫–∞ —É–¥–∞–ª–µ–Ω–∏—è", show_alert=True)
    await call.answer()


# ========== DOCX Generation Handlers ==========

@dp.callback_query(F.data.startswith("party:generate_doc:"))
async def generate_creditors_doc(call: CallbackQuery):
    """–ì–µ–Ω–µ—Ä–∞—Ü–∏—è —Å–ø–∏—Å–∫–∞ –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤ –∏ –¥–æ–ª–∂–Ω–∏–∫–æ–≤ –≤ DOCX."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    case_id = int(parts[2])

    await call.answer("–ì–µ–Ω–µ—Ä–∏—Ä—É—é –¥–æ–∫—É–º–µ–Ω—Ç...")

    try:
        # –ì–µ–Ω–µ—Ä–∞—Ü–∏—è DOCX –∏–∑ —à–∞–±–ª–æ–Ω–∞
        doc_bytes = await render_creditors_list(case_id)

        # –°–æ–∑–¥–∞–Ω–∏–µ —Ñ–∞–π–ª–∞ –¥–ª—è –æ—Ç–ø—Ä–∞–≤–∫–∏
        filename = f"creditors_list_case_{case_id}.docx"
        input_file = BufferedInputFile(doc_bytes, filename=filename)

        # –û—Ç–ø—Ä–∞–≤–∫–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é
        await call.message.answer_document(
            input_file,
            caption="üìÑ –°–ø–∏—Å–æ–∫ –∫—Ä–µ–¥–∏—Ç–æ—Ä–æ–≤ –∏ –¥–æ–ª–∂–Ω–∏–∫–æ–≤"
        )
    except Exception as e:
        logger.error(f"Error generating creditors list: {e}", exc_info=True)
        await call.message.answer("‚ùå –û—à–∏–±–∫–∞ –ø—Ä–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞. –ü—Ä–æ–≤–µ—Ä—å—Ç–µ, —á—Ç–æ –≤—ã –∑–∞–ø–æ–ª–Ω–∏–ª–∏ –ø—Ä–æ—Ñ–∏–ª—å –¥–æ–ª–∂–Ω–∏–∫–∞.")


@dp.callback_query(F.data.startswith("asset:generate_doc:"))
async def generate_inventory_doc(call: CallbackQuery):
    """–ì–µ–Ω–µ—Ä–∞—Ü–∏—è –æ–ø–∏—Å–∏ –∏–º—É—â–µ—Å—Ç–≤–∞ –≤ DOCX."""
    uid = call.from_user.id
    if not is_allowed(uid):
        await call.answer()
        return

    parts = call.data.split(":")
    case_id = int(parts[2])

    await call.answer("–ì–µ–Ω–µ—Ä–∏—Ä—É—é –¥–æ–∫—É–º–µ–Ω—Ç...")

    try:
        # –ì–µ–Ω–µ—Ä–∞—Ü–∏—è DOCX –∏–∑ —à–∞–±–ª–æ–Ω–∞
        doc_bytes = await render_inventory(case_id)

        # –°–æ–∑–¥–∞–Ω–∏–µ —Ñ–∞–π–ª–∞ –¥–ª—è –æ—Ç–ø—Ä–∞–≤–∫–∏
        filename = f"inventory_case_{case_id}.docx"
        input_file = BufferedInputFile(doc_bytes, filename=filename)

        # –û—Ç–ø—Ä–∞–≤–∫–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é
        await call.message.answer_document(
            input_file,
            caption="üìÑ –û–ø–∏—Å—å –∏–º—É—â–µ—Å—Ç–≤–∞ –≥—Ä–∞–∂–¥–∞–Ω–∏–Ω–∞"
        )
    except Exception as e:
        logger.error(f"Error generating inventory: {e}", exc_info=True)
        await call.message.answer("‚ùå –û—à–∏–±–∫–∞ –ø—Ä–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–∞. –ü—Ä–æ–≤–µ—Ä—å—Ç–µ, —á—Ç–æ –≤—ã –∑–∞–ø–æ–ª–Ω–∏–ª–∏ –ø—Ä–æ—Ñ–∏–ª—å –¥–æ–ª–∂–Ω–∏–∫–∞.")


# =========================
# Catch-all callback handler (must be AFTER all specific handlers)
# =========================
@dp.callback_query()
async def on_callback(call: CallbackQuery):
    uid = call.from_user.id
    data = call.data or ""
    flow = USER_FLOW.get(uid) or {}

    if data.startswith(("docs:", "case:", "profile:", "back:")):
        await call.answer()
        return

    if not is_allowed(uid):
        await call.answer("–ù–µ—Ç –¥–æ—Å—Ç—É–ø–∞", show_alert=True)
        return

    is_flow_callback = (
        data == "export:word"
        or data == "flow:cancel"
        or data == "flow:motion"
        or data == "flow:settlement"
        or data.startswith("motion:court:")
    )

    if not is_flow_callback:
        await call.answer()
        return

    if data == "export:word":
        await call.answer()
        text = LAST_RESULT.get(uid)
        if text:
            await call.message.answer(text)
        else:
            await call.message.answer("–ü–æ–∫–∞ –Ω–µ—á–µ–≥–æ —ç–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å.")
        return

    if data == "flow:cancel":
        await call.answer()
        cancel_flow(uid)
        await call.message.answer("–û–∫, –æ—Ç–º–µ–Ω–∏–ª. –ú–µ–Ω—é üëá", reply_markup=main_keyboard())
        return

    if data == "flow:motion":
        await call.answer()
        USER_FLOW[uid] = {"flow": "motion", "stage": "choose_court", "court_type": None, "step": 0, "answers": {}}
        await call.message.answer("–í—ã–±–µ—Ä–∏ —Ç–∏–ø —Å—É–¥–∞:", reply_markup=court_type_keyboard())
        return

    if data.startswith("motion:court:"):
        await call.answer()
        ct = data.split(":")[-1]
        if uid not in USER_FLOW or USER_FLOW[uid].get("flow") != "motion":
            USER_FLOW[uid] = {"flow": "motion", "stage": "fill", "court_type": ct, "step": 0, "answers": {}}
        else:
            USER_FLOW[uid]["stage"] = "fill"
            USER_FLOW[uid]["court_type"] = ct
            USER_FLOW[uid]["step"] = 0
            USER_FLOW[uid]["answers"] = {}
        await call.message.answer(MOTION_STEPS[0][1], reply_markup=motion_actions_keyboard())
        return

    if data == "flow:settlement":
        await call.answer()
        USER_FLOW[uid] = {"flow": "settlement", "step": 0, "answers": {}}
        await call.message.answer(SETTLEMENT_STEPS[0][1], reply_markup=settlement_actions_keyboard())
        return

    await call.answer()
    return


async def main():
    import logging
    logger = logging.getLogger(__name__)

    # Initialize authorization module FIRST (breaks circular imports)
    init_allowed_users(ALLOWED_USERS, ADMIN_USERS)
    logger.info(f"Authorization initialized: {len(ALLOWED_USERS)} allowed users, {len(ADMIN_USERS)} admins")

    # Initialize old SQLite database for existing functionality
    init_db()
    database_url = os.getenv("DATABASE_URL", "postgresql+asyncpg://bankrot:bankrot_password@localhost:5432/bankrot")
    # –ï—Å–ª–∏ –µ—Å—Ç—å engine = ... ‚Üí engine = create_async_engine(database_url)
 
    # Initialize new PostgreSQL database for cases module
    await init_pg_db()
    logger.info("PostgreSQL database initialized")

    bot = Bot(token=BOT_TOKEN)

    # Execution mode: polling (default) or webhook
    mode = os.getenv('TELEGRAM_MODE', 'polling').strip().lower()

    if mode == 'polling':
        # POLLING MODE: Clear any existing webhook and start polling
        logger.info("=" * 60)
        logger.info("Bot starting in POLLING mode")
        logger.info("=" * 60)

        try:
            # Robustly delete any existing webhook
            logger.info("Deleting any existing webhook...")
            webhook_info = await bot.get_webhook_info()
            if webhook_info.url:
                logger.warning(f"Found active webhook: {webhook_info.url}")
                logger.info("Removing webhook to enable polling...")

            await bot.delete_webhook(drop_pending_updates=True)
            logger.info("Webhook deleted successfully. Starting polling...")

            await dp.start_polling(bot)
        except Exception as e:
            logger.error(f"Failed to start bot: {e}")
            if "conflict" in str(e).lower():
                logger.error("TelegramConflictError detected!")
                logger.error("This usually means another bot instance is running or webhook is still active.")
                logger.error("Solutions:")
                logger.error("  1. Stop any other running bot instances")
                logger.error("  2. Manually delete webhook via: curl https://api.telegram.org/bot<TOKEN>/deleteWebhook")
                logger.error("  3. Wait a few minutes and try again")
            raise
        return

    if mode == 'webhook':
        # WEBHOOK MODE: FastAPI (web.py) receives updates and calls dp.feed_update(...)
        logger.info("=" * 60)
        logger.info("Bot starting in WEBHOOK mode")
        logger.info("=" * 60)

        # Initialize web app with dependencies (breaks circular import)
        from web import init_web_app
        init_web_app(BOT_TOKEN, dp)
        logger.info("Web app dependencies initialized")

        logger.info("Webhook updates handled by web.py (FastAPI)")
        logger.info("Start FastAPI server separately: uvicorn web:app --host 0.0.0.0 --port 8000")

        # Keep process alive if someone runs bot.py directly by mistake
        while True:
            await asyncio.sleep(3600)

    raise RuntimeError(f'Unknown TELEGRAM_MODE={mode!r}. Use polling|webhook')

if __name__ == "__main__":
    asyncio.run(main())
# =========================
# HOTFIX: unify main menu
# =========================
try:
    from bankrot_bot.keyboards.menus import main_menu_kb
except (ImportError, ModuleNotFoundError) as e:
    logger.warning(f"Failed to import main_menu_kb: {e}")
    main_menu_kb = None

def main_keyboard() -> InlineKeyboardMarkup:
    """
    Override legacy main_keyboard().
    Always return new unified menu with '‚ûï –°–æ–∑–¥–∞—Ç—å –¥–µ–ª–æ'.
    """
    if main_menu_kb:
        return main_menu_kb()
    raise RuntimeError("main_menu_kb not available")
