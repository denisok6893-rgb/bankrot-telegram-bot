from __future__ import annotations
import logging
import os

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Tuple

from docx import Document

from bankrot_bot.config import load_settings

logger = logging.getLogger(__name__)
from bankrot_bot.services.blocks import (
    build_creditors_block,
    build_creditors_header_block,
    sum_creditors_total,
    build_attachments_list,
    build_vehicle_block,
)


settings = load_settings()
GENERATED_DIR = Path(settings.get('GENERATED_DIR') or 'generated')

def _replace_placeholders_strong(doc: Document, mapping: Dict[str, Any]) -> None:
    """
    Замена плейсхолдеров формата {{key}} по полному тексту параграфов и ячеек таблиц.
    mapping: ключи БЕЗ фигурных скобок, например: {"court_name": "..." }
    """
    def apply_to_paragraph(p):
        text = p.text
        if not text or "{{" not in text:
            return
        changed = False
        for k, v in mapping.items():
            placeholder = f"{{{{{k}}}}}"
            if placeholder in text:
                text = text.replace(placeholder, "" if v is None else str(v))
                changed = True
        if changed:
            _set_paragraph_text_keep_style(p, text)

    for p in doc.paragraphs:
        apply_to_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_to_paragraph(p)
                for nested in cell.tables:
                    # рекурсивно для вложенных таблиц
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            for np in ncell.paragraphs:
                                apply_to_paragraph(np)


def build_gender_forms(gender: str | None) -> dict:
    """
    Возвращает слова в нужном роде для плейсхолдеров шаблона:
    {{debtor_having_word}}, {{debtor_registered_word}}, {{debtor_living_word}},
    {{debtor_not_registered_word}}, {{debtor_insolvent_word}}
    """
    g = (gender or "").strip().lower()
    if g == "female":
        return {
            "debtor_having_word": "имеющая",
            "debtor_registered_word": "зарегистрированная",
            "debtor_living_word": "проживающая",
            "debtor_not_registered_word": "не зарегистрирована",
            "debtor_insolvent_word": "несостоятельной",
        }
    # по умолчанию male
    return {
        "debtor_having_word": "имеющий",
        "debtor_registered_word": "зарегистрированный",
        "debtor_living_word": "проживающий",
        "debtor_not_registered_word": "не зарегистрирован",
        "debtor_insolvent_word": "несостоятельным",
    }



def build_debtor_last_name_initials(card: dict) -> str:
    """
    Из 'Иванов Иван Иванович' делает 'Иванов И. И.'
    Если ФИО пустое/неполное — возвращает как есть.
    """
    full_name = (card.get("debtor_full_name") or "").strip()
    parts = [p for p in full_name.split() if p]
    if len(parts) >= 2:
        last = parts[0]
        first_i = parts[1][0].upper() + "."
        patro_i = (parts[2][0].upper() + ".") if len(parts) >= 3 and parts[2] else ""
        return (last + " " + first_i + (" " + patro_i if patro_i else "")).strip()
    return full_name



def build_family_status_block(card: dict) -> str:
    """
    Возвращает текстовый блок о семейном положении/детях для {{family_status_block}}.
    Поля ожидаются: marital_status, spouse_full_name, has_minor_children, children_count,
    marriage_certificate_number, marriage_certificate_date
    """
    marital_status = (card.get("marital_status") or "").strip()
    spouse_full_name = (card.get("spouse_full_name") or "").strip()
    has_minor_children = card.get("has_minor_children")
    children_count = card.get("children_count")
    cert_no = (card.get("marriage_certificate_number") or "").strip()
    cert_date = (card.get("marriage_certificate_date") or "").strip()

    lines: list[str] = []

    if marital_status == "married":
        line = "Состоит в браке"
        if spouse_full_name:
            line += f" с {spouse_full_name}"
        line += "."
        lines.append(line)

        if cert_no:
            cert_line = f"Свидетельство о заключении брака № {cert_no}"
            if cert_date:
                cert_line += f" от {cert_date}"
            cert_line += "."
            lines.append(cert_line)

    elif marital_status == "single":
        lines.append("В браке не состоит.")

    if has_minor_children is True:
        cnt = ""
        if children_count not in (None, ""):
            cnt = f" ({children_count} ребёнок(детей))"
        lines.append(f"Имеет несовершеннолетних детей{cnt}.")
    elif has_minor_children is False:
        lines.append("Несовершеннолетних детей нет.")

    return "\n".join(lines)



async def build_bankruptcy_petition_doc(case_row: Tuple, card: dict) -> Path:
    """
    Генерация заявления о банкротстве по шаблону.
    Подстановка строго по 23 плейсхолдерам шаблона + дефолты для пустых данных.

    НОВОЕ: приоритетно используем данные из case_parties/case_assets (если есть).
    """
    cid = case_row[0]

    template_path = Path("templates/petitions/bankruptcy_petition.docx")
    doc = Document(template_path)

    # Попытка загрузить кредиторов из новых таблиц
    creditors_from_db = []
    try:
        from bankrot_bot.database import get_session
        from bankrot_bot.services.case_financials import get_case_parties, format_parties_for_doc

        async with get_session() as session:
            parties = await get_case_parties(session, cid, role="creditor")
            if parties:
                creditors_from_db = format_parties_for_doc(parties, role="creditor")
    except Exception as e:
        logger.warning(f"Failed to load creditors from DB for case {cid}: {e}")

    # --- дефолты ---
    def _txt(v: Any) -> str:
        v = "" if v is None else str(v).strip()
        return v if v else "не указано"

    def _money_rubles(v: Any) -> str:
        v = "" if v is None else str(v).strip()
        return v if v else "0"

    def _money_kopeks(v: Any) -> str:
        if v is None or str(v).strip() == "":
            return "00"
        try:
            return f"{int(str(v).strip()):02d}"
        except (ValueError, TypeError):
            s = str(v).strip()
            digits = "".join(ch for ch in s if ch.isdigit())
            if digits == "":
                return "00"
            try:
                return f"{int(digits):02d}"
            except (ValueError, TypeError):
                return "00"

    # --- исходные данные ---
    court_name = _txt(card.get("court_name") or (case_row[4] if len(case_row) > 4 else None))
    court_address = _txt(card.get("court_address"))

    financial_manager_info = _txt(
        card.get("financial_manager_info") or (case_row[6] if len(case_row) > 6 else None)
    )

    debtor_full_name = _txt(card.get("debtor_full_name"))
    debtor_address = _txt(card.get("debtor_address"))
    debtor_birth_date = _txt(card.get("debtor_birth_date"))
    debtor_inn = _txt(card.get("debtor_inn"))
    debtor_snils = _txt(card.get("debtor_snils"))
    debtor_phone = _txt(card.get("debtor_phone"))

    passport_series = (card.get("passport_series") or "").strip()
    passport_number = (card.get("passport_number") or "").strip()
    debtor_passport = _txt(f"{passport_series} {passport_number}".strip())

    debtor_passport_issued_by = _txt(card.get("passport_issued_by"))
    debtor_passport_date = _txt(card.get("passport_date"))
    debtor_passport_code = _txt(card.get("passport_code"))

    raw_marital = card.get("marital_status")
    raw_marital = ("" if raw_marital is None else str(raw_marital)).strip().lower()

    debtor_address = (debtor_address or "").strip()
    while ",," in debtor_address:
        debtor_address = debtor_address.replace(",,", ",")
    debtor_address = debtor_address.rstrip(" ,")

    marital_map = {
        "married": "Состоит в зарегистрированном браке.",
        "single": "В браке не состоит.",
        "divorced": "Брак расторгнут.",
        "widowed": "Вдовец/вдова.",
    }

    # Если уже введён нормальный русский текст — оставляем как есть
    if raw_marital in marital_map:
        marital_status = marital_map[raw_marital]
    else:
        # если строка пустая -> дефолт "не указано"
        # если строка не пустая (в т.ч. русский текст) -> используем как есть
        marital_status = _txt(raw_marital)

    certificate_number = card.get("certificate_number") or card.get("marriage_certificate_number")
    certificate_date = card.get("certificate_date") or card.get("marriage_certificate_date")
    certificate_number = _txt(certificate_number)
    certificate_date = _txt(certificate_date)

    # НОВАЯ ЛОГИКА: приоритет - creditors_from_db, потом card creditors
    if creditors_from_db:
        creditors = creditors_from_db
    else:
        creditors = card.get("creditors") if isinstance(card.get("creditors"), list) else []

    auto_r, auto_k = sum_creditors_total(creditors)
    if auto_r or auto_k:
        total_debt_rubles = str(auto_r)
        total_debt_kopeks = f"{auto_k:02d}"
    else:
        total_debt_rubles = _money_rubles(card.get("total_debt_rubles"))
        total_debt_kopeks = _money_kopeks(card.get("total_debt_kopeks"))

    deposit_deferral_request = card.get("deposit_deferral_request") or ""
    # attachments_list по утверждённым дефолтам должен быть пустым, если нет данных
    attachments_list = ""
    try:
        built_attachments = build_attachments_list(card)
        if built_attachments and str(built_attachments).strip():
            attachments_list = str(built_attachments)
    except (KeyError, TypeError, AttributeError) as e:
        logger.warning(f"Failed to build attachments list: {e}")
        attachments_list = ""

    # creditors_block: creditors_text приоритетно, иначе список из БД/card, иначе нейтральный текст
    creditors_text = card.get("creditors_text")
    creditors_text = str(creditors_text).strip() if creditors_text is not None else ""

    if creditors_text:
        creditors_block = creditors_text
    elif creditors:
        creditors_block = build_creditors_block(creditors)
    else:
        creditors_block = "Сведения о кредиторах не представлены."

    # creditors_header_block: короткий список для шапки (из того же источника, что и creditors_block)
    if creditors:
        creditors_header_block = build_creditors_header_block(creditors)
    else:
        creditors_header_block = "Сведения о кредиторах не представлены."

    # vehicle_block: дефолт при отсутствии
    vehicle_block = ""
    try:
        vehicle_block = build_vehicle_block(card) or ""
    except (KeyError, TypeError, AttributeError) as e:
        logger.warning(f"Failed to build vehicle block: {e}")
        vehicle_block = ""
    if not str(vehicle_block).strip():
        vehicle_block = "Транспортные средства: отсутствуют."

    # --- статус ИП (умная логика: справка или ЕГРИП) ---
    ip_cert_number = (card.get("ip_certificate_number") or "").strip()
    ip_cert_date = (card.get("ip_certificate_date") or "").strip()

    if ip_cert_number and ip_cert_date:
        ip_status_text = (
            "не зарегистрирован в качестве индивидуального предпринимателя, "
            f"что подтверждается справкой № {ip_cert_number} от {ip_cert_date}."
        )
    else:
        ip_status_text = (
            "не зарегистрирован в качестве индивидуального предпринимателя, "
            "что подтверждается сведениями из ЕГРИП"
        )

    # нормализация, чтобы не было 'ЕГРИП..'
    ip_status_text = (ip_status_text or "").strip()
    while ".." in ip_status_text:
        ip_status_text = ip_status_text.replace("..", ".")

    mapping = {
        "attachments_list": attachments_list,
        "certificate_date": certificate_date,
        "certificate_number": certificate_number,
        "court_address": court_address,
        "court_name": court_name,

        # Кредиторы: шапка + основной блок
        "creditors_block": creditors_block,
        "creditors_header_block": creditors_header_block,

        "date": datetime.now().strftime("%d.%m.%Y"),

        "debtor_address": debtor_address,
        "debtor_birth_date": debtor_birth_date,
        "debtor_full_name": debtor_full_name,

        # В шаблоне есть и обычные, и *_or_absent
        "debtor_inn": debtor_inn if debtor_inn != "не указано" else "",
        "debtor_inn_or_absent": debtor_inn if debtor_inn != "не указано" else "отсутствует",

        "debtor_snils": debtor_snils if debtor_snils != "не указано" else "",
        "debtor_snils_or_absent": debtor_snils if debtor_snils != "не указано" else "отсутствует",

        "debtor_phone_or_absent": debtor_phone if debtor_phone != "не указано" else "отсутствует",

        # Паспорт: ключи должны совпадать с плейсхолдерами шаблона
        "passport_series": passport_series or "",
        "passport_number": passport_number or "",
        "passport_issued_by": debtor_passport_issued_by if debtor_passport_issued_by != "не указано" else "",
        "passport_date": debtor_passport_date if debtor_passport_date != "не указано" else "",
        "passport_code": debtor_passport_code if debtor_passport_code != "не указано" else "",

        # Эти плейсхолдеры есть в шаблоне (ты их показывал в списке)
        "debtor_last_name_initials": build_debtor_last_name_initials(card),

        "financial_manager_info": financial_manager_info,
        "family_status_block": build_family_status_block(card),
        "ip_status_text": ip_status_text,

        "marital_status": marital_status,

        "total_debt_kopeks": total_debt_kopeks,
        "total_debt_rubles": total_debt_rubles,

        "vehicle_block": vehicle_block,

        "deposit_deferral_request": deposit_deferral_request,
    }

    # гендерные формы (debtor_having_word, debtor_registered_word, debtor_living_word,
    # debtor_not_registered_word, debtor_insolvent_word)
    try:
        gender_forms = build_gender_forms(card.get("debtor_gender"))
        if isinstance(gender_forms, dict):
            mapping.update(gender_forms)
    except (KeyError, TypeError, AttributeError) as e:
        # если пол не заполнен или функция упала — ставим нейтральные значения
        logger.warning(f"Failed to build gender forms: {e}")
        mapping.update(
            {
                "debtor_having_word": "имеющий(ая)",
                "debtor_registered_word": "зарегистрированный(ая)",
                "debtor_living_word": "проживающий(ая)",
                "debtor_not_registered_word": "не зарегистрирован(а)",
                "debtor_insolvent_word": "неплатёжеспособный(ая)",
            }
        )


    _replace_placeholders_strong(doc, mapping)
    # второй проход — добиваем плейсхолдеры, разорванные Word по runs
    for p in doc.paragraphs:
        for run in p.runs:
            if "{{" in run.text:
                for k, v in mapping.items():
                    run.text = run.text.replace(f"{{{{{k}}}}}", "" if v is None else str(v))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if "{{" in run.text:
                            for k, v in mapping.items():
                                run.text = run.text.replace(f"{{{{{k}}}}}", "" if v is None else str(v))


    # Контроль: не должно остаться {{...}}
    def _has_unreplaced_placeholders(d: Document) -> bool:
        for p in d.paragraphs:
            if "{{" in (p.text or ""):
                return True
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{{" in (p.text or ""):
                            return True
        return False

    if _has_unreplaced_placeholders(doc):
        import re

        # диагностируем, что именно осталось в документе
        left = set()

        def _scan_paragraph_for_left(p):
            txt = p.text or ""
            for m in re.findall(r"\{\{[^}]+\}\}", txt):
                left.add(m)

        for p in doc.paragraphs:
            _scan_paragraph_for_left(p)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        _scan_paragraph_for_left(p)
                    for nested in cell.tables:
                        for nrow in nested.rows:
                            for ncell in nrow.cells:
                                for p in ncell.paragraphs:
                                    _scan_paragraph_for_left(p)

        import logging
        logging.exception("UNREPLACED_PLACEHOLDERS: %s", sorted(left))

        raise ValueError("В документе остались не заменённые плейсхолдеры вида {{...}}")

    fname = f"bankruptcy_petition_case_{cid}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    case_dir = GENERATED_DIR / "cases" / str(cid)
    case_dir.mkdir(parents=True, exist_ok=True)
    out_path = case_dir / fname
    doc.save(out_path)
    return out_path


async def _selected_case_id(state: FSMContext) -> int | None:
    data = await state.get_data()
    try:
        return int(data.get("docs_case_id"))
    except (TypeError, ValueError):
        return None
