"""
Генерация DOCX-форм по шаблонам для дел о банкротстве.

Модуль предоставляет функции для автоматического заполнения шаблонов:
- Список кредиторов и должников гражданина
- Опись имущества гражданина
"""
from __future__ import annotations

import io
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, Tuple

from docx import Document
from docx.shared import Pt
from docx.table import Table, _Cell

logger = logging.getLogger(__name__)


# ========== Утилиты для работы с DOCX ==========

def find_table_by_text(doc: Document, search_text: str) -> Optional[Table]:
    """
    Найти таблицу по тексту в первой строке или ячейке.

    Args:
        doc: Document объект
        search_text: Текст для поиска (например, "Сведения о кредиторах")

    Returns:
        Table или None
    """
    for table in doc.tables:
        # Проверяем первую строку
        if table.rows:
            first_row_text = " ".join(cell.text.strip() for cell in table.rows[0].cells)
            if search_text.lower() in first_row_text.lower():
                return table
    return None


def find_cell_with_text(table: Table, search_text: str) -> Optional[Tuple[int, int, _Cell]]:
    """
    Найти ячейку в таблице по тексту.

    Returns:
        (row_idx, col_idx, cell) или None
    """
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if search_text.lower() in cell.text.lower():
                return (row_idx, col_idx, cell)
    return None


def set_cell_text(cell: _Cell, text: str, preserve_style: bool = True):
    """
    Установить текст в ячейку с сохранением стиля.

    Args:
        cell: Ячейка таблицы
        text: Текст для записи
        preserve_style: Сохранить стиль первого run
    """
    if not cell.paragraphs:
        return

    para = cell.paragraphs[0]

    # Сохраняем стиль первого run
    original_style = None
    if preserve_style and para.runs:
        original_style = para.runs[0].font.size

    # Очищаем содержимое
    para.clear()

    # Добавляем новый текст
    run = para.add_run(str(text))
    if original_style and preserve_style:
        run.font.size = original_style


def add_table_row(table: Table, values: list[str], template_row_idx: int = 1) -> None:
    """
    Добавить строку в таблицу, копируя стиль из template_row_idx.

    Args:
        table: Таблица
        values: Список значений для ячеек
        template_row_idx: Индекс строки-шаблона для копирования стиля
    """
    # Добавляем новую строку
    new_row = table.add_row()

    # Копируем стиль из template_row
    if template_row_idx < len(table.rows) - 1:  # -1 т.к. мы только что добавили строку
        template_row = table.rows[template_row_idx]
        for col_idx, cell in enumerate(new_row.cells):
            if col_idx < len(template_row.cells):
                # Копируем стиль параграфа
                template_cell = template_row.cells[col_idx]
                if template_cell.paragraphs and cell.paragraphs:
                    template_para = template_cell.paragraphs[0]
                    new_para = cell.paragraphs[0]

                    if template_para.runs:
                        template_run = template_para.runs[0]
                        new_run = new_para.add_run()
                        if hasattr(template_run.font, 'size') and template_run.font.size:
                            new_run.font.size = template_run.font.size
                        if hasattr(template_run.font, 'name') and template_run.font.name:
                            new_run.font.name = template_run.font.name

    # Заполняем значениями
    for col_idx, value in enumerate(values):
        if col_idx < len(new_row.cells):
            set_cell_text(new_row.cells[col_idx], value)


def fill_debtor_info_table(doc: Document, debtor_data: dict) -> None:
    """
    Заполнить таблицу с данными должника (общая для обоих шаблонов).

    Ожидаемая структура таблицы:
    - Поиск по лейблам: "Фамилия", "Имя", "Отчество", "Дата рождения", "Адрес", "Паспорт"
    - Заполнение соседней ячейки справа
    """
    # Ищем первую таблицу (обычно это данные должника)
    if not doc.tables:
        logger.warning("No tables found in document")
        return

    first_table = doc.tables[0]

    # Маппинг лейблов на данные
    field_mapping = {
        "фамилия": debtor_data.get("last_name", "-"),
        "имя": debtor_data.get("first_name", "-"),
        "отчество": debtor_data.get("middle_name", "-"),
        "дата рождения": debtor_data.get("birth_date", "-"),
        "место рождения": debtor_data.get("birth_place", "-"),
        "адрес": debtor_data.get("address", "-"),
        "паспорт": debtor_data.get("passport", "-"),
        "снилс": debtor_data.get("snils", "-"),
        "инн": debtor_data.get("inn", "-"),
    }

    # Заполняем ячейки
    for label, value in field_mapping.items():
        result = find_cell_with_text(first_table, label)
        if result:
            row_idx, col_idx, cell = result
            # Записываем значение в следующую ячейку (справа)
            if col_idx + 1 < len(first_table.rows[row_idx].cells):
                target_cell = first_table.rows[row_idx].cells[col_idx + 1]
                set_cell_text(target_cell, value)


# ========== Генерация документов ==========

async def render_creditors_list(case_id: int) -> bytes:
    """
    Сгенерировать "Список кредиторов и должников гражданина".

    Args:
        case_id: ID дела

    Returns:
        bytes: содержимое DOCX-файла
    """
    from bankrot_bot.database import get_session
    from bankrot_bot.services.case_financials import (
        get_case_parties,
        calculate_parties_totals,
    )

    template_path = Path("templates/forms/creditors_list_template.docx")

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Загружаем шаблон
    doc = Document(str(template_path))

    # Получаем данные
    async with get_session() as session:
        # Кредиторы
        creditors = await get_case_parties(session, case_id, role="creditor")
        # Должники (дебиторы)
        debtors = await get_case_parties(session, case_id, role="debtor")

        # Данные должника (из карточки дела)
        # TODO: загрузить из Case/card если нужно
        debtor_data = {
            "last_name": "-",
            "first_name": "-",
            "middle_name": "-",
            "birth_date": "-",
            "birth_place": "-",
            "address": "-",
            "passport": "-",
            "snils": "-",
            "inn": "-",
        }

    # Заполняем данные должника
    fill_debtor_info_table(doc, debtor_data)

    # Заполняем кредиторов
    creditors_table = find_table_by_text(doc, "Сведения о кредиторах")
    if creditors_table and creditors:
        # Удаляем пустые строки-шаблоны (если есть)
        # Предполагаем, что первая строка - заголовок, вторая - шаблон
        # Оставляем только заголовок
        while len(creditors_table.rows) > 1:
            creditors_table._element.remove(creditors_table.rows[-1]._element)

        # Добавляем кредиторов
        for idx, creditor in enumerate(creditors, start=1):
            amount = f"{float(creditor.amount):.2f}" if creditor.amount else "0.00"
            basis = creditor.basis or "-"

            add_table_row(creditors_table, [
                str(idx),
                creditor.name,
                basis,
                amount,
                creditor.currency or "RUB",
            ], template_row_idx=0)

    # Заполняем должников (дебиторов)
    debtors_table = find_table_by_text(doc, "Сведения о должниках")
    if debtors_table and debtors:
        # Аналогично для должников
        while len(debtors_table.rows) > 1:
            debtors_table._element.remove(debtors_table.rows[-1]._element)

        for idx, debtor in enumerate(debtors, start=1):
            amount = f"{float(debtor.amount):.2f}" if debtor.amount else "0.00"
            basis = debtor.basis or "-"

            add_table_row(debtors_table, [
                str(idx),
                debtor.name,
                basis,
                amount,
                debtor.currency or "RUB",
            ], template_row_idx=0)

    # Считаем итоги
    all_parties = creditors + debtors
    totals = calculate_parties_totals(all_parties)

    # Заполняем итоги (ищем таблицу с "Итого")
    for table in doc.tables:
        result = find_cell_with_text(table, "Итого")
        if result:
            row_idx, col_idx, cell = result
            # Записываем суммы в следующие ячейки
            if col_idx + 1 < len(table.rows[row_idx].cells):
                creditors_total = f"{float(totals['total_creditors']):.2f}"
                set_cell_text(table.rows[row_idx].cells[col_idx + 1], creditors_total)

            # Сумма должников (если есть колонка)
            if col_idx + 2 < len(table.rows[row_idx].cells):
                debtors_total = f"{float(totals['total_debtors']):.2f}"
                set_cell_text(table.rows[row_idx].cells[col_idx + 2], debtors_total)

    # Сохраняем в память
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


async def render_inventory(case_id: int) -> bytes:
    """
    Сгенерировать "Опись имущества гражданина".

    Args:
        case_id: ID дела

    Returns:
        bytes: содержимое DOCX-файла
    """
    from bankrot_bot.database import get_session
    from bankrot_bot.services.case_financials import (
        get_case_assets,
        calculate_assets_total,
    )

    template_path = Path("templates/forms/inventory_template.docx")

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Загружаем шаблон
    doc = Document(str(template_path))

    # Получаем данные
    async with get_session() as session:
        assets = await get_case_assets(session, case_id)

        # Данные должника
        debtor_data = {
            "last_name": "-",
            "first_name": "-",
            "middle_name": "-",
            "birth_date": "-",
            "birth_place": "-",
            "address": "-",
            "passport": "-",
            "snils": "-",
            "inn": "-",
        }

    # Заполняем данные должника
    fill_debtor_info_table(doc, debtor_data)

    # Заполняем имущество
    # Обычно в описи имущества несколько таблиц по типам:
    # - Недвижимое имущество
    # - Движимое имущество
    # - Ценные бумаги
    # - Иное имущество

    # Поиск основной таблицы имущества
    inventory_table = find_table_by_text(doc, "имущество") or find_table_by_text(doc, "движимое")

    if inventory_table and assets:
        # Удаляем пустые строки-шаблоны
        while len(inventory_table.rows) > 1:
            inventory_table._element.remove(inventory_table.rows[-1]._element)

        # Добавляем имущество
        for idx, asset in enumerate(assets, start=1):
            value = f"{float(asset.value):.2f}" if asset.value else "-"
            qty = asset.qty_or_area or "-"
            description = asset.description or "-"

            add_table_row(inventory_table, [
                str(idx),
                asset.kind,
                description,
                qty,
                value,
            ], template_row_idx=0)

    # Считаем итоговую стоимость
    total = calculate_assets_total(assets)

    # Заполняем итоги
    for table in doc.tables:
        result = find_cell_with_text(table, "Итого") or find_cell_with_text(table, "Общая стоимость")
        if result:
            row_idx, col_idx, cell = result
            # Записываем сумму в следующую ячейку
            if col_idx + 1 < len(table.rows[row_idx].cells):
                total_str = f"{float(total):.2f}"
                set_cell_text(table.rows[row_idx].cells[col_idx + 1], total_str)

    # Сохраняем в память
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()
